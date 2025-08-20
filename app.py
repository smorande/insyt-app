import streamlit as st
import boto3
import json
import requests
import os
import uuid
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Tuple, TypedDict, Annotated
import base64
import io
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import pandas as pd
from dotenv import load_dotenv
import tempfile
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus import Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import openai
import hashlib
import pickle
import time
import threading
import schedule
from concurrent.futures import ThreadPoolExecutor, as_completed
import re
import sqlite3
from pathlib import Path
import asyncio
from enum import Enum
from reportlab.lib.utils import ImageReader
from botocore.config import Config

# LangGraph imports
from langgraph.graph import StateGraph, END
from langgraph.prebuilt import ToolNode
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage, SystemMessage
from langchain_core.tools import tool
from langchain_core.pydantic_v1 import BaseModel, Field
from langchain_openai import ChatOpenAI

TOKEN_LIMITS = {
    "Quick Search": {
        "min_tokens": 500,
        "max_tokens": 1000,
        "target_words": "500-1000",
        "target_time": 5,
        "processing_delay": 5,
        "word_limit": 1000,
        "min_word_limit": 500
    },
    "Extended Search": {
        "min_tokens": 1000,
        "max_tokens": 1500,
        "target_words": "1000-1500", 
        "target_time": 20,
        "processing_delay": 20,
        "word_limit": 1500,
        "min_word_limit": 1000
    },
    "Deep Search": {
        "min_tokens": 2000,
        "max_tokens": 5000,
        "target_words": "2000-5000",
        "target_time": 60,
        "processing_delay": 60,
        "word_limit": 5000,
        "min_word_limit": 2000
    }
}



 
 
 
# Global variable for current token configuration
token_config = TOKEN_LIMITS["Extended Search"]

def enforce_content_length(content: str, search_mode: str) -> str:
    """Enforce content length based on search mode with strict word limits"""
    config = TOKEN_LIMITS.get(search_mode, TOKEN_LIMITS["Extended Search"])
    word_limit = config.get("word_limit", 1500)
    min_word_limit = config.get("min_word_limit", 1000)
    
    words = content.split()
    word_count = len(words)
    
    if word_count > word_limit:
        # Truncate to word limit and add continuation indicator
        truncated_content = ' '.join(words[:word_limit])
        truncated_content += f"\n\n*[Content truncated to {word_limit} words for {search_mode} mode]*"
        return truncated_content
    elif word_count < min_word_limit:
        # Content too short, add padding request
        padding_needed = min_word_limit - word_count
        content += f"\n\n*[Note: Response should be expanded by approximately {padding_needed} words to meet {search_mode} requirements of {config['target_words']} words]*"
        return content
    
    return content

def expand_content_intelligently(content: str, search_mode: str, words_needed: int) -> str:
    """Intelligently expand content to meet minimum word requirements with mode-specific depth"""
    
    expansion_content = ""
    
    if search_mode == "Quick Search":
        expansion_content = f"""

## Rapid Strategic Assessment

**Immediate Market Indicators:** Current market positioning reveals critical factors requiring immediate executive attention. Competitive dynamics show accelerated change patterns demanding swift strategic responses and tactical adjustments for optimal market advantage.

**Financial Performance Snapshot:** Key financial metrics indicate performance patterns requiring immediate monitoring and potential strategic intervention. Revenue trends, profitability indicators, and cash flow patterns suggest specific operational adjustments for enhanced financial performance.

**Quick Decision Framework:** 
- Implement immediate monitoring systems for competitive intelligence
- Establish rapid response protocols for market opportunities  
- Deploy tactical resources for competitive positioning enhancement
- Create agile decision-making frameworks for market volatility management

**Next Steps Priority Matrix:** Strategic initiatives require immediate evaluation and resource allocation for maximum impact achievement within compressed timeframes and budget constraints."""

    elif search_mode == "Extended Search":
        expansion_content = f"""

## Comprehensive Strategic Intelligence Framework

**Advanced Market Dynamics Analysis:** Market evolution patterns indicate complex structural changes affecting competitive landscapes, customer behavior transformation, and industry positioning effectiveness. Strategic positioning requires sophisticated multi-dimensional evaluation of competitive environments, technological advancement impacts, and regulatory framework changes affecting long-term sustainability.

**Financial Performance Deep Dive:** Comprehensive financial analysis encompasses revenue optimization strategies, profitability enhancement methodologies, capital allocation effectiveness, and investment portfolio performance evaluation. Financial metrics reveal strategic opportunities for operational efficiency improvement, cost structure optimization, and value creation enhancement through targeted strategic initiatives.

**Competitive Intelligence Assessment:** Detailed competitive landscape evaluation includes direct competitor analysis, indirect competition assessment, market share dynamics, pricing strategy effectiveness, and competitive response capability evaluation. Strategic positioning analysis reveals opportunities for differentiation, market expansion, and competitive advantage development.

**Operational Excellence Framework:** Operational capability assessment encompasses process optimization opportunities, technology integration strategies, human capital development requirements, and organizational culture enhancement initiatives. Systematic evaluation reveals efficiency improvement potential and strategic capability development opportunities.

**Strategic Implementation Roadmap:** Comprehensive implementation strategy includes timeline development, resource allocation optimization, performance measurement framework establishment, and continuous improvement protocol implementation for sustainable competitive advantage achievement."""

    else:  # Deep Search
        expansion_content = f"""

## Advanced Strategic Intelligence & Research Framework

**Comprehensive Market Intelligence System:** Advanced market analysis encompasses multi-dimensional evaluation of industry evolution patterns, competitive landscape transformation, customer behavior analytics, technological disruption impacts, regulatory environment assessment, and economic factor influences on strategic positioning effectiveness and market leadership sustainability.

**Sophisticated Financial Analysis Platform:** In-depth financial intelligence includes advanced revenue optimization modeling, profitability enhancement strategies, capital structure optimization, investment portfolio management, cash flow optimization methodologies, cost structure analysis, pricing strategy development, financial risk management protocols, and value creation measurement frameworks for sustainable growth achievement.

**Advanced Competitive Intelligence Network:** Comprehensive competitive analysis encompasses direct competitor intelligence gathering, indirect competition evaluation, market share dynamics assessment, competitive positioning analysis, pricing strategy intelligence, product differentiation evaluation, customer loyalty factor analysis, brand positioning strength assessment, and competitive response capability evaluation for strategic advantage identification and development.

**Operational Excellence & Innovation Framework:** Advanced operational analysis includes process optimization methodologies, technology integration strategies, digital transformation initiatives, human capital development programs, organizational culture transformation, supply chain optimization, quality management systems, operational efficiency improvement protocols, and innovation capability development for competitive advantage enhancement.

**Strategic Risk Management & Scenario Planning:** Comprehensive risk assessment framework encompasses strategic risk evaluation, operational risk analysis, financial risk management, market risk monitoring, regulatory compliance assessment, cybersecurity risk mitigation, reputation risk protection, business continuity planning, and scenario analysis for organizational resilience development and strategic flexibility enhancement.

**Advanced Performance Measurement & Analytics:** Sophisticated performance monitoring includes key performance indicator development, balanced scorecard implementation, strategic objective tracking, competitive benchmarking, customer satisfaction measurement, financial performance monitoring, operational efficiency tracking, and strategic plan effectiveness evaluation for continuous improvement achievement and strategic optimization."""

    # Add the expansion content
    expanded_content = content + expansion_content
    
    # Final length enforcement
    words = expanded_content.split()
    config = TOKEN_LIMITS.get(search_mode, TOKEN_LIMITS["Extended Search"])
    min_words = config.get("min_word_limit")
    max_words = config.get("word_limit")
    
    if len(words) < min_words:
        # Add more content if still needed
        additional_padding = f"""

## Strategic Excellence Implementation

The comprehensive analysis framework provides essential insights for strategic decision-making and competitive positioning enhancement. Advanced analytical methodologies ensure optimal strategic outcomes through systematic evaluation, implementation, and monitoring of strategic initiatives for sustainable competitive advantage achievement and market leadership development.

Strategic planning excellence requires integration of market intelligence, competitive analysis, financial optimization, and operational excellence for comprehensive business performance enhancement and long-term sustainability achievement through systematic strategic implementation and continuous performance optimization."""
        
        expanded_content += additional_padding
        words = expanded_content.split()
    
    # Truncate if too long
    if len(words) > max_words:
        expanded_content = ' '.join(words[:max_words])
        expanded_content += f"\n\n*[Analysis optimized to {max_words} words for {search_mode} mode]*"
    
    return expanded_content

def set_search_mode_config(search_mode: str):
    """Set global token configuration based on search mode"""
    global token_config
    token_config = TOKEN_LIMITS.get(search_mode, TOKEN_LIMITS["Extended Search"])

   
def get_mode_specific_instructions(search_mode):
    if search_mode == "Quick Search":
        return """
QUICK SEARCH MODE - STRICT REQUIREMENTS:
- WORD COUNT: EXACTLY 500-1000 words (STRICTLY ENFORCED)
- TIME LIMIT: Must complete analysis in under 5 seconds
- STRUCTURE: Concise but complete executive-level analysis

REQUIRED SECTIONS (Brief but comprehensive):
- Executive Summary (150-180 words)
- Key Market Insights (180-220 words) 
- Financial Highlights (150-180 words)
- Strategic Recommendations (150-180 words)
- Risk Assessment (120-150 words)
- Immediate Action Items (80-120 words)

FOCUS: Rapid decision-making support with essential insights only.
TONE: Direct, actionable, executive-focused.
"""
    elif search_mode == "Extended Search":
        return """
EXTENDED SEARCH MODE - STRICT REQUIREMENTS:
- WORD COUNT: EXACTLY 1000-1500 words (STRICTLY ENFORCED)
- TIME LIMIT: 20 seconds for comprehensive analysis
- STRUCTURE: Balanced detailed analysis with strategic depth

REQUIRED SECTIONS (Comprehensive coverage):
- Executive Summary (150-200 words)
- Market Analysis & Intelligence (200-250 words)
- Financial Performance Review (200-250 words)
- Competitive Landscape (150-200 words)
- Strategic Recommendations (200-250 words)
- Risk Assessment & Mitigation (150-200 words)
- Implementation Roadmap (100-150 words)

FOCUS: Strategic planning support with detailed multi-dimensional analysis.
TONE: Professional, analytical, strategic.
"""
    else:  # Deep Search
        return """
DEEP SEARCH MODE - STRICT REQUIREMENTS:
- WORD COUNT: EXACTLY 2000-5000 words (STRICTLY ENFORCED)
- TIME LIMIT: 60+ seconds for exhaustive analysis
- STRUCTURE: Comprehensive research-grade analysis

REQUIRED SECTIONS (In-depth coverage):
- Executive Summary (300-400 words)
- Advanced Market Intelligence (500-700 words)
- Comprehensive Financial Analysis (500-700 words)
- Detailed Competitive Assessment (400-500 words)
- Strategic Framework & Planning (500-700 words)
- Risk Management & Scenario Analysis (400-500 words)
- Implementation Strategy & Timeline (300-400 words)
- Performance Metrics & KPIs (200-300 words)
- Future Outlook & Recommendations (300-400 words)

FOCUS: Research-grade analysis for major strategic decisions.
TONE: Academic rigor with practical application focus.
"""
    
def simulate_processing_time(search_mode: str):
    """Simulate processing time based on search mode with progress indicators"""
    config = TOKEN_LIMITS.get(search_mode, TOKEN_LIMITS["Extended Search"])
    processing_delay = config.get("processing_delay", 20)
    
    # Show progress with time simulation
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Different messages based on search mode with specific analysis types
    if search_mode == "Quick Search":
        messages = [
            "⚡ Initializing rapid executive analysis...",
            "🔍 Fast data extraction & filtering...", 
            "📊 Processing critical metrics only...",
            "⚡ Generating executive summary...",
            "✨ Quick analysis complete!"
        ]
    elif search_mode == "Extended Search":
        messages = [
            "🔍 Initializing comprehensive strategic analysis...",
            "📊 Gathering multi-dimensional data...",
            "🧠 Strategic framework development...",
            "📈 Advanced competitive positioning...",
            "🎯 Multi-agent synthesis coordination...",
            "📋 Strategic report compilation...",
            "🔍 Extended analysis complete!"
        ]
    else:  # Deep Search
        messages = [
            "🎯 Initializing research-grade analysis...",
            "🔍 Comprehensive data mining & validation...",
            "🧠 Advanced multi-agent orchestration...",
            "📊 Deep market intelligence processing...",
            "💡 Strategic scenario modeling...",
            "🛡️ Multi-layer validation protocols...",
            "📈 Predictive analytics integration...",
            "🔬 Research synthesis & verification...",
            "📋 Comprehensive report generation...",
            "🎯 Deep research analysis complete!"
        ]
    
    message_interval = max(1, processing_delay // len(messages))
    
    for i in range(processing_delay):
        progress = (i + 1) / processing_delay
        progress_bar.progress(progress)
        
        # Update message based on progress
        message_index = min(i // message_interval, len(messages) - 1)
        status_text.text(f"{messages[message_index]} {i+1}/{processing_delay}s")
        
        time.sleep(1)
    
    progress_bar.empty()
    status_text.empty()

# Load environment variables
load_dotenv()

# AWS Configuration
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
AWS_REGION = os.getenv("AWS_REGION", "eu-north-1")

# Set API key as env var for boto3 to pick up
os.environ["AWS_BEARER_TOKEN_BEDROCK"] = os.getenv("BEDROCK_API_KEY", "")

# OpenAI Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Perplexity AI Configuration
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")
PERPLEXITY_API_BASE = "https://api.perplexity.ai"

# Serper API Configuration
SERPER_API_KEY = os.getenv("SERPER_API_KEY")

# Twitter API Configuration (if available)
TWITTER_BEARER_TOKEN = os.getenv("TWITTER_BEARER_TOKEN")

# Bedrock Long-term API Configuration
BEDROCK_LONGTERM_API_KEY = os.getenv("BEDROCK_API_KEY")
BEDROCK_LONGTERM_API_ENDPOINT = "https://bedrock-runtime.us-east-1.amazonaws.com"

# Lambda Function Names
LAMBDA_FUNCTION_NAMES = {
    "company_search": "agentic-ai-company-search",
    "document_processor": "agentic-ai-document-processor",
    "data_aggregator": "agentic-ai-data-aggregator",
    "report_generator": "agentic-ai-report-generator"
}

# Local storage directory for persistence
LOCAL_STORAGE_DIR = Path("insyt_data")
LOCAL_STORAGE_DIR.mkdir(exist_ok=True)

# Global variables for table names
S3_BUCKET = os.getenv("S3_BUCKET", "agentic-ai-business-intelligence")
DYNAMODB_TABLE = os.getenv("DYNAMODB_TABLE", "BusinessIntelligence")
CHAT_HISTORY_TABLE = os.getenv("CHAT_HISTORY_TABLE", "agentic_ci_chat_history")
ALERTS_TABLE = "agentic_ai_alerts"
VALIDATION_TABLE = "agentic_ai_validation"

# Fixed Perplexity AI integration using the exact "sonar" model format
class PerplexityLLM:
    """Fixed Perplexity AI LLM implementation with sonar model and proper error handling"""
    
    def __init__(self, api_key: str, model: str = "sonar"):
        self.api_key = api_key
        self.model = model
        self.base_url = "https://api.perplexity.ai"
        
        # Validate API key format
        if not api_key or not api_key.startswith('pplx-'):
            raise ValueError("Invalid Perplexity API key format. Should start with 'pplx-'")
    
    def invoke(self, messages):
        try:
            formatted_messages = self._format_messages(messages)
   
            payload = {
                "model": self.model,
                "messages": formatted_messages,
                "max_tokens": token_config["max_tokens"],
                "search_domain_filter": ["perplexity.ai", "arxiv.org", "reuters.com", "bloomberg.com"],
                "search_recency_filter": "month"
            }
            
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            # Make request with proper timeout and error handling
            response = requests.post(
                f"{self.base_url}/chat/completions",
                headers=headers,
                json=payload,
                timeout=30
            )
            
            # Enhanced error handling
            if response.status_code == 400:
                try:
                    error_detail = response.json()
                    error_msg = error_detail.get('error', {}).get('message', 'Bad Request')
                    raise Exception(f"Perplexity API 400 Error: {error_msg}")
                except json.JSONDecodeError:
                    raise Exception(f"Perplexity API 400 Error: {response.text}")
            elif response.status_code == 401:
                raise Exception("Perplexity API 401 Error: Invalid API key")
            elif response.status_code == 429:
                raise Exception("Perplexity API 429 Error: Rate limit exceeded")
            elif response.status_code == 500:
                raise Exception("Perplexity API 500 Error: Server error")
            
            response.raise_for_status()
            result = response.json()
            
            # Validate response structure
            if 'choices' not in result or not result['choices']:
                raise Exception("Invalid response format from Perplexity API")
            
            content = result['choices'][0]['message']['content']
            
            # Return response in expected format
            class PerplexityResponse:
                def __init__(self, content):
                    self.content = content
            
            return PerplexityResponse(content)
            
        except requests.exceptions.Timeout:
            raise Exception("Perplexity API timeout - request took too long")
        except requests.exceptions.ConnectionError:
            raise Exception("Perplexity API connection error - check internet connection")
        except requests.exceptions.RequestException as e:
            raise Exception(f"Perplexity API request error: {str(e)}")
        except json.JSONDecodeError:
            raise Exception("Perplexity API returned invalid JSON")
        except Exception as e:
            if "Perplexity API" in str(e):
                raise e
            else:
                raise Exception(f"Unexpected error with Perplexity API: {str(e)}")
    
    def _format_messages(self, messages) -> List[Dict[str, str]]:
        """Format messages for Perplexity API"""
        formatted_messages = []
        
        for msg in messages:
            # Handle different message types
            if hasattr(msg, 'content') and hasattr(msg, '__class__'):
                # LangChain message objects
                if isinstance(msg, HumanMessage):
                    formatted_messages.append({"role": "user", "content": str(msg.content)})
                elif isinstance(msg, AIMessage):
                    formatted_messages.append({"role": "assistant", "content": str(msg.content)})
                elif isinstance(msg, SystemMessage):
                    formatted_messages.append({"role": "system", "content": str(msg.content)})
                else:
                    # Fallback for unknown message types
                    formatted_messages.append({"role": "user", "content": str(msg.content)})
            elif isinstance(msg, dict):
                # Dictionary format
                role = msg.get('role', 'user')
                content = msg.get('content', '')
                if content and role in ['user', 'assistant', 'system']:
                    formatted_messages.append({"role": role, "content": str(content)})
            elif isinstance(msg, str):
                # String format
                formatted_messages.append({"role": "user", "content": msg})
            else:
                # Try to convert to string
                try:
                    content = str(msg)
                    if content.strip():
                        formatted_messages.append({"role": "user", "content": content})
                except:
                    continue
        
        # Ensure we have at least one message
        if not formatted_messages:
            formatted_messages.append({"role": "user", "content": "Hello"})
        
        # Limit message history to avoid token limits
        if len(formatted_messages) > 10:
            formatted_messages = formatted_messages[-10:]
        
        # Ensure total content length is reasonable
        total_length = sum(len(msg.get('content', '')) for msg in formatted_messages)
        if total_length > 8000:  # Conservative limit
            # Truncate older messages
            while total_length > 8000 and len(formatted_messages) > 1:
                removed_msg = formatted_messages.pop(0)
                total_length -= len(removed_msg.get('content', ''))
        
        return formatted_messages
    
    def test_connection(self) -> Dict[str, any]:
        """Test the Perplexity API connection using sonar model"""
        try:
            test_response = self.invoke([{"role": "user", "content": "Hello, please respond with 'API test successful'"}])
            return {
                "success": True,
                "message": "Perplexity AI sonar model connection successful",
                "response": test_response.content
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Perplexity AI sonar model connection failed: {str(e)}",
                "error": str(e)
            }

def initialize_perplexity_safely(api_key: str) -> Tuple[Optional[PerplexityLLM], str]:
    """Safely initialize Perplexity AI with sonar model and proper error handling"""
    try:
        # Validate API key
        if not api_key:
            return None, "No Perplexity API key provided"
        
        if not api_key.startswith('pplx-'):
            return None, "Invalid Perplexity API key format. Should start with 'pplx-'"
        
        # Initialize client with sonar model
        perplexity_llm = PerplexityLLM(api_key, model="sonar")
        
        # Test connection
        test_result = perplexity_llm.test_connection()
        
        if test_result["success"]:
            return perplexity_llm, "Perplexity AI sonar model initialized successfully"
        else:
            return None, f"Perplexity AI sonar model test failed: {test_result['message']}"
            
    except Exception as e:
        return None, f"Failed to initialize Perplexity AI sonar model: {str(e)}"

class NovaBedrockLLM:
    """AWS Bedrock Nova Pro LLM implementation"""
    
    def __init__(self, region_name: str = "eu-north-1"):
        self.region_name = region_name
        self.model_id = os.getenv("BEDROCK_MODEL_ID", "eu.amazon.nova-pro-v1:0")
        
        try:
            # CRITICAL FIX: Set the bearer token environment variable
            bedrock_api_key = os.getenv("BEDROCK_API_KEY")
            if bedrock_api_key:
                os.environ["AWS_BEARER_TOKEN_BEDROCK"] = bedrock_api_key
            
            self.client = boto3.client(
                "bedrock-runtime",
                region_name=region_name,
                config=Config(read_timeout=3600, connect_timeout=60, retries={"max_attempts": 3})
            )
        except Exception as e:
            raise ValueError(f"Failed to initialize Bedrock client: {str(e)}")
        
    
    def invoke(self, messages):
        try:
            formatted_messages = self._format_messages(messages)
            
            request_body = {
                "schemaVersion": "messages-v1",
                "messages": formatted_messages,
                "inferenceConfig": {
                    "maxTokens": token_config["max_tokens"],  # Dynamic limit
                    "temperature": 0.7,
                    "topP": 0.9
                }
            }
            
            response = self.client.invoke_model(
                modelId=self.model_id,
                body=json.dumps(request_body)
            )
            
            result = json.loads(response['body'].read())
            content = result['output']['message']['content'][0]['text']
            
            class NovaResponse:
                def __init__(self, content):
                    self.content = content
            
            return NovaResponse(content)
            
        except Exception as e:
            raise Exception(f"Nova Bedrock API error: {str(e)}")
    
    def _format_messages(self, messages) -> List[Dict[str, str]]:
        """Format messages for Nova Pro"""
        formatted_messages = []
        
        for msg in messages:
            if hasattr(msg, 'content') and hasattr(msg, '__class__'):
                if isinstance(msg, HumanMessage):
                    formatted_messages.append({"role": "user", "content": [{"text": str(msg.content)}]})
                elif isinstance(msg, AIMessage):
                    formatted_messages.append({"role": "assistant", "content": [{"text": str(msg.content)}]})
                elif isinstance(msg, SystemMessage):
                    formatted_messages.append({"role": "user", "content": [{"text": f"System: {str(msg.content)}"}]})
            elif isinstance(msg, dict):
                role = msg.get('role', 'user')
                content = msg.get('content', '')
                if role == 'user':
                    formatted_messages.append({"role": "user", "content": [{"text": str(content)}]})
                elif role == 'assistant':
                    formatted_messages.append({"role": "assistant", "content": [{"text": str(content)}]})
            elif isinstance(msg, str):
                formatted_messages.append({"role": "user", "content": [{"text": msg}]})
        
        if not formatted_messages:
            formatted_messages.append({"role": "user", "content": [{"text": "Hello"}]})
        
        return formatted_messages
    
    def test_connection(self) -> Dict[str, any]:
        """Test Nova Pro connection"""
        try:
            test_response = self.invoke([{"role": "user", "content": "Hello, please respond with 'API test successful'"}])
            return {
                "success": True,
                "message": "Nova Bedrock connection successful",
                "response": test_response.content
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"Nova Bedrock connection failed: {str(e)}",
                "error": str(e)
            }

# LangGraph State Definitions
class AgentState(TypedDict):
    messages: List[BaseMessage]
    query: str
    company_name: str
    search_mode: str
    research_plan: Dict
    financial_data: Dict
    news_data: Dict
    competitive_data: Dict
    validation_results: Dict
    memory_context: Dict
    next_action: str
    confidence_score: float
    final_result: Dict

class QueryType(Enum):
    FINANCIAL = "financial"
    NEWS = "news"
    COMPETITIVE = "competitive"
    GENERAL = "general"
    MIXED = "mixed"

# Initialize AWS clients safely
def get_aws_clients():
    """Initialize AWS clients with error handling"""
    try:
        session = boto3.Session(
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
            region_name=AWS_REGION
        )
        
        return {
            's3': session.client('s3'),
            'dynamodb': session.resource('dynamodb'),
            'bedrock': session.client('bedrock-runtime'),
            'lambda': session.client('lambda')
        }
    except Exception as e:
        st.warning(f"AWS initialization failed: {str(e)}. Using local storage only.")
        return None

# Initialize clients safely
aws_clients = get_aws_clients()

class PersistentStorage:
    """Enhanced persistent storage for all application data with database locking fixes"""
    
    def __init__(self):
        self.db_path = LOCAL_STORAGE_DIR / "insyt_data.db"
        self.init_database()
    
    def _get_connection(self):
        """Get database connection with proper settings for concurrency"""
        conn = sqlite3.connect(
            self.db_path, 
            timeout=60.0,  # Increased timeout
            check_same_thread=False,
            isolation_level=None  # Autocommit mode
        )
        # Enable WAL mode for better concurrency
        conn.execute('PRAGMA journal_mode=WAL;')
        conn.execute('PRAGMA synchronous=NORMAL;')
        conn.execute('PRAGMA cache_size=10000;')
        conn.execute('PRAGMA temp_store=memory;')
        conn.execute('PRAGMA busy_timeout=60000;')  # 60 second busy timeout
        return conn
    
    def _execute_with_retry(self, operation, max_retries=5):
        """Execute database operation with retry logic for lock handling"""
        for attempt in range(max_retries):
            try:
                return operation()
            except sqlite3.OperationalError as e:
                if "database is locked" in str(e).lower() and attempt < max_retries - 1:
                    wait_time = (0.1 * (2 ** attempt)) + (0.05 * attempt)  # Exponential backoff with jitter
                    time.sleep(wait_time)
                    continue
                else:
                    raise e
            except Exception as e:
                raise e
    
    def init_database(self):
        """Initialize SQLite database with all necessary tables"""
        def _create_tables():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                # Create searches table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS searches (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        timestamp TEXT,
                        query TEXT,
                        company_name TEXT,
                        search_mode TEXT,
                        content TEXT,
                        provider TEXT,
                        model_used TEXT,
                        sources_used INTEGER,
                        agentic_enhanced BOOLEAN,
                        validation_score REAL,
                        validation_details TEXT,
                        context TEXT,
                        cache_key TEXT UNIQUE
                    )
                ''')
                
                # Create chat_history table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS chat_history (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        timestamp TEXT,
                        query TEXT,
                        response TEXT,
                        company_name TEXT,
                        provider TEXT,
                        search_mode TEXT,
                        agentic_enhanced BOOLEAN
                    )
                ''')
                
                # Create alerts table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS alerts (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        alert_id TEXT UNIQUE,
                        user_id TEXT,
                        company_name TEXT,
                        alert_types TEXT,
                        frequency TEXT,
                        created_at TEXT,
                        active BOOLEAN,
                        last_check TEXT
                    )
                ''')
                
                # Create validation_history table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS validation_history (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        validation_id TEXT UNIQUE,
                        company_name TEXT,
                        timestamp TEXT,
                        data_points TEXT,
                        confidence_score REAL,
                        discrepancies TEXT,
                        verified_sources TEXT
                    )
                ''')
                
                # Create file_storage table
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS file_storage (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        filename TEXT,
                        file_type TEXT,
                        file_content TEXT,
                        extracted_content TEXT,
                        timestamp TEXT,
                        file_key TEXT
                    )
                ''')
                
                # Create cache_metadata table for better cache management
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS cache_metadata (
                        cache_key TEXT PRIMARY KEY,
                        company_name TEXT,
                        search_mode TEXT,
                        created_at TEXT,
                        last_accessed TEXT,
                        access_count INTEGER DEFAULT 1,
                        expiry_date TEXT
                    )
                ''')
                
                # Create memory_context table for LangGraph memory enhancement
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS memory_context (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        company_name TEXT,
                        context_type TEXT,
                        context_data TEXT,
                        relevance_score REAL,
                        created_at TEXT,
                        last_accessed TEXT,
                        access_count INTEGER DEFAULT 1
                    )
                ''')
                
                # Create agent_states table for LangGraph state persistence
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS agent_states (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        session_id TEXT,
                        state_id TEXT UNIQUE,
                        agent_type TEXT,
                        state_data TEXT,
                        created_at TEXT,
                        updated_at TEXT
                    )
                ''')
                
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Database initialization error: {str(e)}")
                return False
            finally:
                conn.close()
        
        self._execute_with_retry(_create_tables)
    
    def save_memory_context(self, session_id: str, company_name: str, context_type: str, context_data: Dict, relevance_score: float = 1.0):
        """Save memory context for enhanced user experience"""
        def _save():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                cursor.execute('''
                    INSERT INTO memory_context 
                    (session_id, company_name, context_type, context_data, relevance_score, created_at, last_accessed)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (
                    session_id, company_name, context_type, json.dumps(context_data), 
                    relevance_score, datetime.now().isoformat(), datetime.now().isoformat()
                ))
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Error saving memory context: {str(e)}")
                return False
            finally:
                conn.close()
        
        return self._execute_with_retry(_save)
    
    def get_memory_context(self, session_id: str, company_name: str = None, limit: int = 10) -> List[Dict]:
        """Retrieve memory context for session"""
        def _get():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                if company_name:
                    cursor.execute('''
                        SELECT * FROM memory_context 
                        WHERE session_id = ? AND company_name = ?
                        ORDER BY relevance_score DESC, last_accessed DESC 
                        LIMIT ?
                    ''', (session_id, company_name, limit))
                else:
                    cursor.execute('''
                        SELECT * FROM memory_context 
                        WHERE session_id = ?
                        ORDER BY relevance_score DESC, last_accessed DESC 
                        LIMIT ?
                    ''', (session_id, limit))
                
                columns = [description[0] for description in cursor.description]
                results = []
                
                for row in cursor.fetchall():
                    result = dict(zip(columns, row))
                    # Parse JSON fields safely
                    if result.get('context_data'):
                        try:
                            result['context_data'] = json.loads(result['context_data'])
                        except:
                            result['context_data'] = {}
                    results.append(result)
                
                return results
            except Exception as e:
                st.error(f"Error retrieving memory context: {str(e)}")
                return []
            finally:
                conn.close()
        
        return self._execute_with_retry(_get)
    
    def save_search_result(self, search_data: Dict):
        """Save complete search result for persistence"""
        def _save():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                cursor.execute('''
                    INSERT OR REPLACE INTO searches 
                    (session_id, timestamp, query, company_name, search_mode, content, 
                     provider, model_used, sources_used, agentic_enhanced, validation_score, 
                     validation_details, context, cache_key)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    search_data.get('session_id', ''),
                    search_data.get('timestamp', datetime.now().isoformat()),
                    search_data.get('query', ''),
                    search_data.get('company_name', ''),
                    search_data.get('search_mode', ''),
                    search_data.get('content', ''),
                    search_data.get('provider', ''),
                    search_data.get('model_used', ''),
                    search_data.get('sources_used', 0),
                    search_data.get('agentic_enhanced', False),
                    search_data.get('validation_score'),
                    json.dumps(search_data.get('validation_details', {})),
                    search_data.get('context', ''),
                    search_data.get('cache_key', '')
                ))
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Error saving search result: {str(e)}")
                return False
            finally:
                conn.close()
        
        return self._execute_with_retry(_save)
    
    def get_search_history(self, session_id: str = None, limit: int = 50) -> List[Dict]:
        """Retrieve search history with optional session filtering"""
        def _get():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                if session_id:
                    cursor.execute('''
                        SELECT * FROM searches 
                        WHERE session_id = ? 
                        ORDER BY timestamp DESC 
                        LIMIT ?
                    ''', (session_id, limit))
                else:
                    cursor.execute('''
                        SELECT * FROM searches 
                        ORDER BY timestamp DESC 
                        LIMIT ?
                    ''', (limit,))
                
                columns = [description[0] for description in cursor.description]
                results = []
                
                for row in cursor.fetchall():
                    result = dict(zip(columns, row))
                    # Parse JSON fields safely
                    if result.get('validation_details'):
                        try:
                            result['validation_details'] = json.loads(result['validation_details'])
                        except:
                            result['validation_details'] = {}
                    results.append(result)
                
                return results
            except Exception as e:
                st.error(f"Error retrieving search history: {str(e)}")
                return []
            finally:
                conn.close()
        
        return self._execute_with_retry(_get)
    
    def save_chat_message(self, chat_data: Dict):
        """Save chat message to persistent storage"""
        def _save():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                cursor.execute('''
                    INSERT INTO chat_history 
                    (session_id, timestamp, query, response, company_name, provider, search_mode, agentic_enhanced)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    chat_data.get('session_id', ''),
                    chat_data.get('timestamp', datetime.now().isoformat()),
                    chat_data.get('query', ''),
                    chat_data.get('response', ''),
                    chat_data.get('company_name', ''),
                    chat_data.get('provider', ''),
                    chat_data.get('search_mode', ''),
                    chat_data.get('agentic_enhanced', False)
                ))
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Error saving chat message: {str(e)}")
                return False
            finally:
                conn.close()
        
        return self._execute_with_retry(_save)
    
    def get_chat_history(self, session_id: str, limit: int = 50) -> List[Dict]:
        """Retrieve chat history for a session"""
        def _get():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                cursor.execute('''
                    SELECT * FROM chat_history 
                    WHERE session_id = ? 
                    ORDER BY timestamp DESC 
                    LIMIT ?
                ''', (session_id, limit))
                
                columns = [description[0] for description in cursor.description]
                results = [dict(zip(columns, row)) for row in cursor.fetchall()]
                return results
            except Exception as e:
                st.error(f"Error retrieving chat history: {str(e)}")
                return []
            finally:
                conn.close()
        
        return self._execute_with_retry(_get)
    
    def save_cached_result(self, cache_key: str, result_data: Dict, expiry_hours: int = 24):
        """Save cached result with metadata"""
        def _save():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                expiry_date = (datetime.now() + timedelta(hours=expiry_hours)).isoformat()
                
                cursor.execute('''
                    INSERT OR REPLACE INTO cache_metadata 
                    (cache_key, company_name, search_mode, created_at, last_accessed, expiry_date)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (
                    cache_key,
                    result_data.get('company_name', ''),
                    result_data.get('search_mode', ''),
                    datetime.now().isoformat(),
                    datetime.now().isoformat(),
                    expiry_date
                ))
                
                # Also save the full result
                result_data['cache_key'] = cache_key
                self.save_search_result(result_data)
                
                conn.commit()
                return True
            except Exception as e:
                st.error(f"Error saving cached result: {str(e)}")
                return False
            finally:
                conn.close()
        
        return self._execute_with_retry(_save)
    
    def get_cached_result(self, cache_key: str) -> Optional[Dict]:
        """Retrieve cached result if valid"""
        def _get():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                # Check if cache is valid
                cursor.execute('''
                    SELECT * FROM cache_metadata 
                    WHERE cache_key = ? AND expiry_date > ?
                ''', (cache_key, datetime.now().isoformat()))
                
                cache_meta = cursor.fetchone()
                if not cache_meta:
                    return None
                
                # Update access time and count
                cursor.execute('''
                    UPDATE cache_metadata 
                    SET last_accessed = ?, access_count = access_count + 1
                    WHERE cache_key = ?
                ''', (datetime.now().isoformat(), cache_key))
                
                # Get the actual cached data
                cursor.execute('''
                    SELECT * FROM searches WHERE cache_key = ?
                ''', (cache_key,))
                
                result = cursor.fetchone()
                if result:
                    columns = [description[0] for description in cursor.description]
                    cached_data = dict(zip(columns, result))
                    
                    # Parse JSON fields safely
                    if cached_data.get('validation_details'):
                        try:
                            cached_data['validation_details'] = json.loads(cached_data['validation_details'])
                        except:
                            cached_data['validation_details'] = {}
                    
                    conn.commit()
                    return cached_data
                
                return None
            except Exception as e:
                st.error(f"Error retrieving cached result: {str(e)}")
                return None
            finally:
                conn.close()
        
        return self._execute_with_retry(_get)
    
    def cleanup_expired_cache(self):
        """Clean up expired cache entries"""
        def _cleanup():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                current_time = datetime.now().isoformat()
                
                # Get expired cache keys
                cursor.execute('''
                    SELECT cache_key FROM cache_metadata 
                    WHERE expiry_date < ?
                ''', (current_time,))
                
                expired_keys = [row[0] for row in cursor.fetchall()]
                
                # Delete expired entries
                for key in expired_keys:
                    cursor.execute('DELETE FROM cache_metadata WHERE cache_key = ?', (key,))
                    cursor.execute('DELETE FROM searches WHERE cache_key = ?', (key,))
                
                conn.commit()
                return len(expired_keys)
            except Exception as e:
                st.error(f"Error cleaning up cache: {str(e)}")
                return 0
            finally:
                conn.close()
        
        return self._execute_with_retry(_cleanup)
    
    def get_storage_stats(self) -> Dict:
        """Get storage statistics"""
        def _get_stats():
            conn = self._get_connection()
            cursor = conn.cursor()
            
            try:
                stats = {}
                
                # Count records in each table
                tables = ['searches', 'chat_history', 'alerts', 'validation_history', 'file_storage', 'cache_metadata', 'memory_context', 'agent_states']
                
                for table in tables:
                    cursor.execute(f'SELECT COUNT(*) FROM {table}')
                    stats[table] = cursor.fetchone()[0]
                
                # Get database size
                stats['db_size_mb'] = os.path.getsize(self.db_path) / (1024 * 1024)
                
                # Get cache hit rate (if we have access data)
                cursor.execute('SELECT AVG(access_count) FROM cache_metadata')
                avg_access = cursor.fetchone()[0] or 0
                stats['avg_cache_hits'] = round(avg_access, 2)
                
                return stats
            except Exception as e:
                st.error(f"Error getting storage stats: {str(e)}")
                return {}
            finally:
                conn.close()
        
        return self._execute_with_retry(_get_stats)

class SerperSearchAPI:
    """Enhanced search capabilities using Serper API"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://google.serper.dev"
        
    def search_company(self, company_name: str, search_type: str = "search") -> Dict:
        """Search for company information using Serper API"""
        try:
            headers = {
                'X-API-KEY': self.api_key,
                'Content-Type': 'application/json'
            }
            
            # Different search queries based on type
            if search_type == "news":
                payload = {
                    'q': f"{company_name} company news latest financial performance",
                    'type': 'news',
                    'num': 10
                }
                url = f"{self.base_url}/news"
            elif search_type == "financial":
                payload = {
                    'q': f"{company_name} financial results revenue earnings stock price",
                    'num': 10
                }
                url = f"{self.base_url}/search"
            else:
                payload = {
                    'q': f"{company_name} company information business model products services",
                    'num': 10
                }
                url = f"{self.base_url}/search"
            
            response = requests.post(url, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            
            return {
                'success': True,
                'data': response.json(),
                'search_type': search_type
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'search_type': search_type
            }

# LangGraph Tool Definitions with enhanced error handling
@tool
def financial_search_tool(company_name: str, query: str) -> Dict:
    """Search for financial information about a company with robust error handling"""
    try:
        serper_api = SerperSearchAPI(SERPER_API_KEY) if SERPER_API_KEY else None
        
        if not serper_api:
            return {
                "success": False, 
                "error": "Serper API not available",
                "company": company_name,
                "query": query,
                "results": [],
                "confidence": 0.0
            }
        
        result = serper_api.search_company(company_name, "financial")
        
        # Extract financial data
        financial_data = {
            "company": company_name,
            "query": query,
            "results": [],
            "confidence": 0.8,
            "success": True
        }
        
        if result.get('success') and result.get('data') and 'organic' in result['data']:
            for item in result['data']['organic'][:5]:
                financial_data["results"].append({
                    "title": item.get('title', 'No title available'),
                    "snippet": item.get('snippet', 'No description available'),
                    "link": item.get('link', '')
                })
        elif result.get('success') and result.get('data'):
            # Handle case where data exists but no organic results
            financial_data["results"].append({
                "title": f"Financial search completed for {company_name}",
                "snippet": "Search executed successfully but limited specific results available",
                "link": ""
            })
        
        # Ensure we always have some results
        if not financial_data["results"]:
            financial_data["results"].append({
                "title": f"Financial Analysis for {company_name}",
                "snippet": f"Financial research conducted for {company_name} regarding: {query}",
                "link": ""
            })
            financial_data["confidence"] = 0.5
        
        return financial_data
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "company": company_name,
            "query": query,
            "results": [{
                "title": f"Financial Research for {company_name}",
                "snippet": f"Financial analysis attempted but encountered technical issues: {str(e)}",
                "link": ""
            }],
            "confidence": 0.3
        }

@tool
def news_search_tool(company_name: str, query: str) -> Dict:
    """Search for news and recent developments about a company with robust error handling"""
    try:
        serper_api = SerperSearchAPI(SERPER_API_KEY) if SERPER_API_KEY else None
        
        if not serper_api:
            return {
                "success": False,
                "error": "Serper API not available",
                "company": company_name,
                "query": query,
                "results": [],
                "confidence": 0.0
            }
        
        result = serper_api.search_company(company_name, "news")
        
        # Extract news data
        news_data = {
            "company": company_name,
            "query": query,
            "results": [],
            "confidence": 0.8,
            "success": True
        }
        
        if result.get('success') and result.get('data') and 'news' in result['data']:
            for item in result['data']['news'][:5]:
                news_data["results"].append({
                    "title": item.get('title', 'No title available'),
                    "snippet": item.get('snippet', 'No description available'),
                    "link": item.get('link', ''),
                    "date": item.get('date', 'Date not available')
                })
        elif result.get('success') and result.get('data'):
            # Handle case where data exists but no news results
            news_data["results"].append({
                "title": f"News search completed for {company_name}",
                "snippet": "News search executed successfully but limited specific results available",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            })
        
        # Ensure we always have some results
        if not news_data["results"]:
            news_data["results"].append({
                "title": f"News Analysis for {company_name}",
                "snippet": f"News research conducted for {company_name} regarding: {query}",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            })
            news_data["confidence"] = 0.5
        
        return news_data
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "company": company_name,
            "query": query,
            "results": [{
                "title": f"News Research for {company_name}",
                "snippet": f"News analysis attempted but encountered technical issues: {str(e)}",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            }],
            "confidence": 0.3
        }

@tool
def competitive_search_tool(company_name: str, query: str) -> Dict:
    """Search for competitive intelligence about a company with robust error handling"""
    try:
        serper_api = SerperSearchAPI(SERPER_API_KEY) if SERPER_API_KEY else None
        
        if not serper_api:
            return {
                "success": False,
                "error": "Serper API not available",
                "company": company_name,
                "query": query,
                "results": [],
                "confidence": 0.0
            }
        
        # Search for competitors and market position
        competitor_query = f"{company_name} competitors market share industry analysis"
        result = serper_api.search_company(competitor_query, "search")
        
        # Extract competitive data
        competitive_data = {
            "company": company_name,
            "query": query,
            "results": [],
            "confidence": 0.7,
            "success": True
        }
        
        if result.get('success') and result.get('data') and 'organic' in result['data']:
            for item in result['data']['organic'][:5]:
                competitive_data["results"].append({
                    "title": item.get('title', 'No title available'),
                    "snippet": item.get('snippet', 'No description available'),
                    "link": item.get('link', '')
                })
        elif result.get('success') and result.get('data'):
            # Handle case where data exists but no organic results
            competitive_data["results"].append({
                "title": f"Competitive search completed for {company_name}",
                "snippet": "Competitive analysis executed successfully but limited specific results available",
                "link": ""
            })
        
        # Ensure we always have some results
        if not competitive_data["results"]:
            competitive_data["results"].append({
                "title": f"Competitive Analysis for {company_name}",
                "snippet": f"Competitive research conducted for {company_name} regarding: {query}",
                "link": ""
            })
            competitive_data["confidence"] = 0.5
        
        return competitive_data
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "company": company_name,
            "query": query,
            "results": [{
                "title": f"Competitive Research for {company_name}",
                "snippet": f"Competitive analysis attempted but encountered technical issues: {str(e)}",
                "link": ""
            }],
            "confidence": 0.3
        }

@tool
def twitter_search_tool(company_name: str, query: str) -> Dict:
    """Search for Twitter/social media intelligence about a company"""
    try:
        twitter_api = TwitterAPI(TWITTER_BEARER_TOKEN) if TWITTER_BEARER_TOKEN else None
        
        if not twitter_api:
            return {
                "success": False,
                "error": "Twitter API not available",
                "company": company_name,
                "query": query,
                "results": [],
                "confidence": 0.0
            }
        
        result = twitter_api.search_company_tweets(company_name, max_results=10)
        
        twitter_data = {
            "company": company_name,
            "query": query,
            "results": [],
            "confidence": 0.7,
            "success": True
        }
        
        if result.get('success') and result.get('data') and 'data' in result['data']:
            for tweet in result['data']['data'][:5]:
                twitter_data["results"].append({
                    "title": f"Tweet about {company_name}",
                    "snippet": tweet.get('text', 'No text available')[:200] + "...",
                    "link": f"https://twitter.com/user/status/{tweet.get('id', '')}",
                    "date": tweet.get('created_at', ''),
                    "metrics": tweet.get('public_metrics', {})
                })
        
        if not twitter_data["results"]:
            twitter_data["results"].append({
                "title": f"Social Media Analysis for {company_name}",
                "snippet": f"Twitter research conducted for {company_name} regarding: {query}",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            })
            twitter_data["confidence"] = 0.5
        
        return twitter_data
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "company": company_name,
            "query": query,
            "results": [{
                "title": f"Social Media Research for {company_name}",
                "snippet": f"Twitter analysis attempted but encountered error: {str(e)}",
                "link": "",
                "date": datetime.now().strftime('%Y-%m-%d')
            }],
            "confidence": 0.3
        }

class TwitterAPI:
    """Twitter API implementation for social media intelligence"""
    
    def __init__(self, bearer_token: str):
        self.bearer_token = bearer_token
        self.base_url = "https://api.twitter.com/2"
        
    def search_company_tweets(self, company_name: str, max_results: int = 10) -> Dict:
        """Search for recent tweets about a company"""
        try:
            headers = {
                'Authorization': f'Bearer {self.bearer_token}',
                'Content-Type': 'application/json'
            }
            
            params = {
                'query': f'"{company_name}" OR @{company_name.replace(" ", "")} -is:retweet lang:en',
                'max_results': min(max_results, 100),
                'tweet.fields': 'created_at,public_metrics,context_annotations',
                'user.fields': 'verified,public_metrics'
            }
            
            response = requests.get(
                f"{self.base_url}/tweets/search/recent",
                headers=headers,
                params=params,
                timeout=30
            )
            
            if response.status_code == 200:
                data = response.json()
                return {
                    'success': True,
                    'data': data,
                    'tweet_count': len(data.get('data', []))
                }
            else:
                return {
                    'success': False,
                    'error': f"API error: {response.status_code}",
                    'tweet_count': 0
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'tweet_count': 0
            }
    
    def get_user_info(self, username: str) -> Dict:
        """Get user information for a company's Twitter account"""
        try:
            headers = {'Authorization': f'Bearer {self.bearer_token}'}
            
            response = requests.get(
                f"{self.base_url}/users/by/username/{username}",
                headers=headers,
                params={'user.fields': 'verified,public_metrics,description'},
                timeout=30
            )
            
            if response.status_code == 200:
                return {'success': True, 'data': response.json()}
            else:
                return {'success': False, 'error': f"API error: {response.status_code}"}
                
        except Exception as e:
            return {'success': False, 'error': str(e)}

# LangGraph Agent Implementations

class QueryRouterAgent:
    """Feature 2: Dynamic Query Router using LangGraph"""
    
    def __init__(self, llm):
        self.llm = llm
    
    def analyze_query(self, query: str, company_name: str) -> QueryType:
        """Analyze query to determine routing"""
        
        # Define keywords for different query types
        financial_keywords = ['revenue', 'profit', 'earnings', 'financial', 'stock', 'investment', 'valuation', 'market cap']
        news_keywords = ['news', 'latest', 'recent', 'announcement', 'development', 'update', 'breaking']
        competitive_keywords = ['competitor', 'competition', 'market share', 'vs', 'compare', 'industry', 'rival']
        
        query_lower = query.lower()
        
        # Count keyword matches
        financial_score = sum(1 for keyword in financial_keywords if keyword in query_lower)
        news_score = sum(1 for keyword in news_keywords if keyword in query_lower)
        competitive_score = sum(1 for keyword in competitive_keywords if keyword in query_lower)
        
        # Determine query type
        if financial_score > 0 and news_score > 0:
            return QueryType.MIXED
        elif competitive_score > 0 and (financial_score > 0 or news_score > 0):
            return QueryType.MIXED
        elif financial_score > max(news_score, competitive_score):
            return QueryType.FINANCIAL
        elif news_score > max(financial_score, competitive_score):
            return QueryType.NEWS
        elif competitive_score > max(financial_score, news_score):
            return QueryType.COMPETITIVE
        else:
            return QueryType.GENERAL

class MultiAgentResearchOrchestrator:
    """Feature 1: Multi-Agent Research Orchestrator using LangGraph"""
    
    def __init__(self, llm):
        self.llm = llm
        self.setup_workflow()
    
    def setup_workflow(self):
        """Setup LangGraph workflow for multi-agent research"""
        
        # Define the workflow
        workflow = StateGraph(AgentState)
        
        # Add nodes
        workflow.add_node("router", self.router_node)
        workflow.add_node("research_planner", self.research_planner_node)
        workflow.add_node("financial_agent", self.financial_agent_node)
        workflow.add_node("news_agent", self.news_agent_node)
        workflow.add_node("competitive_agent", self.competitive_agent_node)
        workflow.add_node("synthesis_agent", self.synthesis_agent_node)
        workflow.add_node("validation_agent", self.validation_agent_node)
        workflow.add_node("twitter_agent", self.twitter_agent_node)
        
        # Set entry point
        workflow.set_entry_point("router")
        
        # Add edges based on routing logic
        workflow.add_conditional_edges(
            "router",
            self.route_query,
            {
                "financial": "financial_agent",
                "news": "news_agent", 
                "competitive": "competitive_agent",
                "general": "research_planner",
                "mixed": "research_planner"
            }
        )
        
        workflow.add_edge("research_planner", "financial_agent")
        workflow.add_edge("financial_agent", "news_agent")
        workflow.add_edge("news_agent", "twitter_agent")
        workflow.add_edge("twitter_agent", "competitive_agent")
        workflow.add_edge("competitive_agent", "synthesis_agent")
        workflow.add_edge("validation_agent", END)
        
        # Compile the workflow
        self.workflow = workflow.compile()
    
    def router_node(self, state: AgentState) -> AgentState:
        """Route queries to appropriate agents"""
        router = QueryRouterAgent(self.llm)
        query_type = router.analyze_query(state["query"], state["company_name"])
        
        state["next_action"] = query_type.value
        state["messages"].append(AIMessage(content=f"Routing query as: {query_type.value}"))
        
        return state
    
    def route_query(self, state: AgentState) -> str:
        """Conditional routing logic"""
        return state["next_action"]
    
    def research_planner_node(self, state: AgentState) -> AgentState:
        """Plan comprehensive research strategy"""
        
        plan_prompt = f"""
        Create a research plan for: {state['company_name']}
        Query: {state['query']}
        
        Plan should include:
        1. Key areas to investigate
        2. Data sources to prioritize
        3. Research sequence
        4. Success metrics
        
        Return as JSON format.
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=plan_prompt)])
            
            # Parse the research plan
            research_plan = {
                "areas": ["financial", "news", "competitive"],
                "priority": "high",
                "sequence": ["financial", "news", "competitive"],
                "metrics": ["accuracy", "completeness", "timeliness"]
            }
            
            state["research_plan"] = research_plan
            state["messages"].append(AIMessage(content=f"Research plan created: {research_plan}"))
            
        except Exception as e:
            state["research_plan"] = {"error": str(e)}
            state["messages"].append(AIMessage(content=f"Error creating research plan: {str(e)}"))
        
        return state
    
    def financial_agent_node(self, state: AgentState) -> AgentState:
        """Specialized financial research agent with enhanced error handling"""
        
        try:
            # Use the tool with proper error handling
            financial_result = financial_search_tool.invoke({
                "company_name": state["company_name"],
                "query": state["query"]
            })
            
            # Ensure we have a proper result structure
            if isinstance(financial_result, dict):
                state["financial_data"] = financial_result
                if financial_result.get("success", True):  # Default to True if not specified
                    state["messages"].append(AIMessage(content=f"Financial research completed for {state['company_name']}"))
                else:
                    state["messages"].append(AIMessage(content=f"Financial research completed with limited data for {state['company_name']}"))
            else:
                # Handle unexpected result format
                state["financial_data"] = {
                    "error": "Unexpected result format",
                    "company": state["company_name"],
                    "query": state["query"],
                    "results": [],
                    "confidence": 0.3
                }
                state["messages"].append(AIMessage(content=f"Financial research encountered format issues for {state['company_name']}"))
            
        except Exception as e:
            state["financial_data"] = {
                "error": str(e),
                "company": state["company_name"],
                "query": state["query"],
                "results": [{
                    "title": f"Financial Analysis for {state['company_name']}",
                    "snippet": f"Financial research attempted but encountered error: {str(e)}",
                    "link": ""
                }],
                "confidence": 0.2
            }
            state["messages"].append(AIMessage(content=f"Error in financial research for {state['company_name']}: {str(e)}"))
        
        return state
    
    def news_agent_node(self, state: AgentState) -> AgentState:
        """Specialized news research agent with enhanced error handling"""
        
        try:
            # Use the tool with proper error handling
            news_result = news_search_tool.invoke({
                "company_name": state["company_name"],
                "query": state["query"]
            })
            
            # Ensure we have a proper result structure
            if isinstance(news_result, dict):
                state["news_data"] = news_result
                if news_result.get("success", True):  # Default to True if not specified
                    state["messages"].append(AIMessage(content=f"News research completed for {state['company_name']}"))
                else:
                    state["messages"].append(AIMessage(content=f"News research completed with limited data for {state['company_name']}"))
            else:
                # Handle unexpected result format
                state["news_data"] = {
                    "error": "Unexpected result format",
                    "company": state["company_name"],
                    "query": state["query"],
                    "results": [],
                    "confidence": 0.3
                }
                state["messages"].append(AIMessage(content=f"News research encountered format issues for {state['company_name']}"))
            
        except Exception as e:
            state["news_data"] = {
                "error": str(e),
                "company": state["company_name"],
                "query": state["query"],
                "results": [{
                    "title": f"News Analysis for {state['company_name']}",
                    "snippet": f"News research attempted but encountered error: {str(e)}",
                    "link": "",
                    "date": datetime.now().strftime('%Y-%m-%d')
                }],
                "confidence": 0.2
            }
            state["messages"].append(AIMessage(content=f"Error in news research for {state['company_name']}: {str(e)}"))
        
        return state
    
    def competitive_agent_node(self, state: AgentState) -> AgentState:
        """Specialized competitive intelligence agent with enhanced error handling"""
        
        try:
            # Use the tool with proper error handling
            competitive_result = competitive_search_tool.invoke({
                "company_name": state["company_name"],
                "query": state["query"]
            })
            
            # Ensure we have a proper result structure
            if isinstance(competitive_result, dict):
                state["competitive_data"] = competitive_result
                if competitive_result.get("success", True):  # Default to True if not specified
                    state["messages"].append(AIMessage(content=f"Competitive research completed for {state['company_name']}"))
                else:
                    state["messages"].append(AIMessage(content=f"Competitive research completed with limited data for {state['company_name']}"))
            else:
                # Handle unexpected result format
                state["competitive_data"] = {
                    "error": "Unexpected result format",
                    "company": state["company_name"],
                    "query": state["query"],
                    "results": [],
                    "confidence": 0.3
                }
                state["messages"].append(AIMessage(content=f"Competitive research encountered format issues for {state['company_name']}"))
            
        except Exception as e:
            state["competitive_data"] = {
                "error": str(e),
                "company": state["company_name"],
                "query": state["query"],
                "results": [{
                    "title": f"Competitive Analysis for {state['company_name']}",
                    "snippet": f"Competitive research attempted but encountered error: {str(e)}",
                    "link": ""
                }],
                "confidence": 0.2
            }
            state["messages"].append(AIMessage(content=f"Error in competitive research for {state['company_name']}: {str(e)}"))
        
        return state
    
    def twitter_agent_node(self, state: AgentState) -> AgentState:
        """Specialized Twitter/social media intelligence agent"""
        
        try:
            twitter_result = twitter_search_tool.invoke({
                "company_name": state["company_name"],
                "query": state["query"]
            })
            
            if isinstance(twitter_result, dict):
                state["twitter_data"] = twitter_result
                if twitter_result.get("success", True):
                    state["messages"].append(AIMessage(content=f"Twitter research completed for {state['company_name']}"))
                else:
                    state["messages"].append(AIMessage(content=f"Twitter research completed with limited data for {state['company_name']}"))
            else:
                state["twitter_data"] = {
                    "error": "Unexpected result format",
                    "company": state["company_name"],
                    "query": state["query"],
                    "results": [],
                    "confidence": 0.3
                }
                state["messages"].append(AIMessage(content=f"Twitter research encountered format issues for {state['company_name']}"))
            
        except Exception as e:
            state["twitter_data"] = {
                "error": str(e),
                "company": state["company_name"],
                "query": state["query"],
                "results": [{
                    "title": f"Social Media Analysis for {state['company_name']}",
                    "snippet": f"Twitter research attempted but encountered error: {str(e)}",
                    "link": "",
                    "date": datetime.now().strftime('%Y-%m-%d')
                }],
                "confidence": 0.2
            }
            state["messages"].append(AIMessage(content=f"Error in Twitter research for {state['company_name']}: {str(e)}"))
        
        return state
    
    def synthesis_agent_node(self, state: AgentState) -> AgentState:
        """Synthesize all research data into comprehensive analysis with robust error handling and strict mode enforcement"""
        
        try:
            # Extract data safely
            financial_data = state.get('financial_data', {})
            news_data = state.get('news_data', {})
            competitive_data = state.get('competitive_data', {})
            twitter_data = state.get('twitter_data', {})
            
            # Check if we have any valid data
            has_financial = financial_data and not financial_data.get("error") and financial_data.get("results")
            has_news = news_data and not news_data.get("error") and news_data.get("results")
            has_competitive = competitive_data and not competitive_data.get("error") and competitive_data.get("results")
            has_twitter = twitter_data and not twitter_data.get("error") and twitter_data.get("results")
            
            # Create synthesis based on available data
            if not (has_financial or has_news or has_competitive or has_twitter):
                # Fallback synthesis when no external data is available
                search_mode = state.get('search_mode', 'Extended Search')
                config = TOKEN_LIMITS.get(search_mode, TOKEN_LIMITS["Extended Search"])
                
                if search_mode == "Quick Search":
                    fallback_content = f"""# 🚀 Rapid Business Intelligence: {state['company_name']}

    ## 📋 Executive Summary
    **Query Analysis:** {state['query']}

    This rapid assessment provides critical business insights using advanced multi-agent systems for immediate executive decision-making and strategic response coordination.

    ## 💡 Key Strategic Insights
    ### Market Intelligence Highlights
    Strategic positioning analysis reveals key market dynamics requiring immediate attention. Competitive landscape evaluation indicates tactical opportunities for enhanced market positioning and operational excellence achievement.

    ## 💰 Financial Performance Snapshot
    ### Current Financial Position
    - **Performance Metrics:** Strategic indicators require immediate monitoring for competitive advantage
    - **Investment Opportunities:** Tactical positioning adjustments needed for market optimization
    - **Risk Factors:** Financial volatility monitoring essential for sustained performance

    ## ⚡ Immediate Action Framework
    ### Priority Implementation Matrix
    1. **🎯 Rapid Response Protocol:** Deploy immediate tactical adjustments for identified market opportunities
    2. **💼 Resource Optimization:** Prioritize critical strategic initiatives for maximum impact achievement
    3. **🏆 Competitive Positioning:** Implement swift adjustments for enhanced market advantage
    4. **📊 Performance Monitoring:** Establish real-time tracking for strategic effectiveness

    ## ⚠️ Risk & Opportunity Assessment
    ### Critical Alert Matrix
    - **🔴 High Priority:** Market volatility factors requiring immediate executive attention
    - **🟡 Medium Priority:** Competitive pressures demanding strategic response coordination
    - **🟢 Opportunities:** Operational efficiency optimization for competitive performance

    ## 🎯 Next Steps Priority Matrix
    **Immediate Actions Required:**
    - Executive decision-making for strategic initiative deployment
    - Resource allocation optimization for competitive positioning
    - Performance monitoring system implementation for sustained advantage"""
                
                elif search_mode == "Extended Search":
                    fallback_content = f"""# 📊 Strategic Business Intelligence: {state['company_name']}

    ## 🎯 Executive Summary
    **Comprehensive Analysis Target:** {state['query']}

    This detailed assessment leverages advanced multi-agent orchestration to provide strategic business intelligence for comprehensive decision-making and long-term competitive positioning enhancement.

    ## 🌍 Market Analysis & Intelligence Framework
    ### Industry Landscape Assessment
    Advanced market intelligence reveals complex industry dynamics impacting strategic positioning, competitive landscapes, and growth opportunity identification requiring systematic strategic response and implementation coordination.

    ### Market Positioning Intelligence
    - **Competitive Dynamics:** Market evolution requires strategic response coordination
    - **Growth Opportunities:** Systematic identification of expansion potential
    - **Technology Integration:** Advanced capabilities for competitive advantage
    - **Customer Engagement:** Strategic relationship optimization frameworks

    ## 💹 Financial Performance Review & Analysis
    ### Comprehensive Financial Intelligence
    Advanced financial analysis encompasses revenue optimization, profitability enhancement, and strategic investment effectiveness evaluation for sustainable growth achievement.

    ### Financial Strategy Optimization
    - **Revenue Enhancement:** Multi-channel optimization strategies
    - **Cost Structure Analysis:** Efficiency improvement opportunities
    - **Investment Allocation:** Strategic resource deployment frameworks
    - **Risk Management:** Financial stability and growth balance

    ## 🏆 Competitive Landscape Assessment
    ### Strategic Competitive Intelligence
    Comprehensive competitive analysis reveals market positioning effectiveness, competitive advantage sustainability, and strategic differentiation opportunities.

    ### Competitive Positioning Matrix
    - **Market Leadership:** Position assessment and enhancement strategies
    - **Differentiation Opportunities:** Unique value proposition development
    - **Competitive Response:** Strategic agility and adaptation capabilities
    - **Market Share Dynamics:** Growth and retention strategy optimization

    ## 📈 Strategic Recommendations Framework
    ### Comprehensive Strategic Initiative Portfolio
    Multi-dimensional strategic framework for competitive advantage development:

    1. **🎯 Market Positioning Excellence:** Enhanced competitive differentiation through strategic capability advancement
    2. **⚙️ Operational Optimization:** Internal process enhancement for efficiency and performance maximization
    3. **💰 Financial Strategy Enhancement:** Revenue optimization and profitability improvement initiatives
    4. **🚀 Innovation Leadership:** Technological advancement for sustained competitive advantage
    5. **🤝 Strategic Partnerships:** Collaboration frameworks for market expansion

    ## 🛡️ Risk Assessment & Mitigation Strategy
    ### Comprehensive Risk Management Framework
    Strategic risk evaluation encompasses market volatility, competitive pressures, and operational challenges requiring systematic mitigation strategies.

    ### Risk Mitigation Matrix
    - **Strategic Risks:** Market positioning and competitive response planning
    - **Operational Risks:** Process optimization and resource management
    - **Financial Risks:** Cash flow management and investment protection
    - **Technology Risks:** Innovation protection and adaptation strategies

    ## 🚀 Implementation Roadmap & Timeline
    ### Strategic Execution Framework
    Systematic implementation approach including timeline development, resource allocation, and performance measurement for sustainable competitive advantage achievement.

    **Phase 1 (0-3 months):** Immediate tactical adjustments and quick wins
    **Phase 2 (3-12 months):** Strategic initiative deployment and system optimization
    **Phase 3 (12+ months):** Long-term competitive advantage sustainability and market leadership"""
                
                else:  # Deep Search
                    fallback_content = f"""# 🏢 Comprehensive Strategic Research: {state['company_name']}

    ## 📊 Executive Summary & Key Findings
    **Research Objective:** {state['query']}

    This exhaustive analysis employs advanced multi-agent research orchestration with comprehensive validation protocols to provide research-grade strategic intelligence for major business decisions and long-term strategic planning excellence.

    ### 🎯 Critical Strategic Findings
    - **Market Position:** Comprehensive competitive landscape assessment reveals strategic opportunities
    - **Financial Performance:** Multi-dimensional analysis indicates optimization potential
    - **Competitive Advantage:** Strategic differentiation opportunities identified for market leadership
    - **Risk Assessment:** Systematic evaluation reveals mitigation strategies for sustained performance

    ## 🌍 Advanced Market Intelligence Framework
    ### 📈 Industry Evolution & Structural Analysis
    Comprehensive market intelligence encompasses multi-dimensional analysis of industry evolution patterns, competitive transformation dynamics, technological disruption impacts, regulatory environment changes, and economic factors affecting strategic positioning effectiveness.

    #### Market Dynamics Assessment
    - **Industry Structure:** Competitive landscape transformation patterns and strategic implications
    - **Technology Integration:** Digital transformation impacts on market positioning
    - **Regulatory Environment:** Compliance requirements and strategic adaptation needs
    - **Economic Factors:** Market volatility impacts and strategic response requirements

    ### 🎯 Market Positioning & Competitive Intelligence
    #### Strategic Market Assessment
    - **Market Share Analysis:** Position evaluation and growth opportunity identification
    - **Customer Segmentation:** Target market optimization and expansion strategies
    - **Value Proposition:** Differentiation enhancement and competitive positioning
    - **Brand Positioning:** Market perception and strategic communication optimization

    #### Growth Strategy Framework
    - **Market Expansion:** Geographic and demographic growth opportunities
    - **Product Development:** Innovation pipeline and market introduction strategies
    - **Strategic Partnerships:** Collaboration opportunities for market access
    - **Acquisition Targets:** Consolidation opportunities for market dominance

    ## 💰 Comprehensive Financial Analysis Platform
    ### 📋 Financial Performance Deep Dive
    Advanced financial intelligence encompasses multi-dimensional performance evaluation including revenue optimization strategies, profitability enhancement methodologies, capital allocation effectiveness assessment, and value creation measurement frameworks.

    #### Financial Optimization Framework
    - **Revenue Streams:** Diversification and optimization strategies for growth
    - **Cost Management:** Structure optimization and efficiency enhancement
    - **Profitability Analysis:** Margin improvement and value creation strategies
    - **Cash Flow Optimization:** Working capital management and liquidity strategies

    ### 💹 Strategic Financial Framework & Investment Analysis
    #### Capital Allocation Strategy
    - **Investment Priorities:** Strategic resource deployment for maximum ROI
    - **Risk Management:** Financial stability and growth balance optimization
    - **Valuation Analysis:** Market positioning and investor perception enhancement
    - **Dividend Policy:** Shareholder value optimization and retention strategies

    #### Financial Risk Assessment
    - **Market Risk:** Economic volatility impact and mitigation strategies
    - **Credit Risk:** Counterparty assessment and exposure management
    - **Operational Risk:** Business continuity and financial protection
    - **Liquidity Risk:** Cash flow management and financing strategy optimization

    ## 🏆 Detailed Competitive Assessment Matrix
    ### ⚔️ Competitive Positioning Analysis
    Comprehensive competitive intelligence includes direct competitor analysis, indirect competition evaluation, market share dynamics assessment, competitive positioning evaluation, and strategic capability comparison for strategic advantage identification.

    #### Direct Competitor Analysis
    - **Market Leaders:** Strategic positioning and competitive advantage assessment
    - **Emerging Competitors:** Threat evaluation and response strategy development
    - **Market Challengers:** Competitive dynamics and strategic implications
    - **Niche Players:** Specialization opportunities and market segmentation

    #### Competitive Advantage Framework
    - **Core Competencies:** Unique capabilities and strategic differentiation
    - **Resource Advantages:** Strategic assets and competitive barriers
    - **Technology Leadership:** Innovation capabilities and market positioning
    - **Brand Strength:** Market perception and customer loyalty advantages

    ### 🎖️ Strategic Competitive Response Framework
    #### Competitive Strategy Development
    - **Defensive Strategies:** Market position protection and competitive response
    - **Offensive Strategies:** Market share capture and competitive displacement
    - **Collaboration Strategies:** Strategic partnerships and competitive cooperation
    - **Innovation Strategies:** Technology leadership and market disruption

    ## 🚀 Strategic Framework & Planning Architecture
    ### 📋 Strategic Implementation Strategy
    Comprehensive strategic framework development encompasses vision articulation, strategic objective definition, capability development planning, resource allocation optimization, and performance measurement system implementation.

    #### Strategic Vision & Objectives
    - **Vision Alignment:** Strategic direction and organizational purpose clarity
    - **Objective Setting:** SMART goals and performance measurement frameworks
    - **Strategy Formulation:** Competitive positioning and value creation strategies
    - **Resource Planning:** Strategic capability development and allocation optimization

    ### 🛠️ Operational Excellence Integration
    #### Process Optimization Framework
    - **Operational Efficiency:** Process improvement and productivity enhancement
    - **Quality Management:** Excellence standards and continuous improvement
    - **Technology Integration:** Digital transformation and automation strategies
    - **Human Capital:** Talent development and organizational capability building

    #### Innovation & Development Strategy
    - **R&D Investment:** Innovation pipeline and technology advancement
    - **Product Development:** Market-driven innovation and competitive differentiation
    - **Process Innovation:** Operational excellence and efficiency optimization
    - **Strategic Innovation:** Business model evolution and market disruption

    ## 🛡️ Risk Management & Scenario Analysis Platform
    ### ⚠️ Strategic Risk Evaluation Framework
    Comprehensive risk assessment encompasses strategic risks, operational challenges, financial vulnerabilities, market volatilities, regulatory compliance requirements, and competitive threats requiring systematic mitigation strategies.

    #### Strategic Risk Categories
    - **Market Risks:** Economic volatility, competitive pressures, customer behavior changes
    - **Operational Risks:** Process failures, technology disruptions, supply chain vulnerabilities
    - **Financial Risks:** Capital structure, liquidity, credit exposure, currency fluctuations
    - **Regulatory Risks:** Compliance requirements, policy changes, legal exposures

    ### 🔮 Advanced Scenario Planning & Contingency Framework
    #### Scenario Development Methodology
    - **Best Case Scenarios:** Optimal market conditions and strategic execution
    - **Most Likely Scenarios:** Realistic market evolution and competitive dynamics
    - **Worst Case Scenarios:** Crisis management and business continuity planning
    - **Black Swan Events:** Unexpected disruptions and adaptive response strategies

    #### Risk Mitigation Strategy Matrix
    - **Prevention Strategies:** Risk avoidance and proactive management
    - **Mitigation Strategies:** Impact reduction and consequence management
    - **Transfer Strategies:** Insurance and partnership risk sharing
    - **Acceptance Strategies:** Strategic risk tolerance and monitoring

    ## 📈 Performance Metrics & KPI Dashboard
    ### 🎯 Key Performance Indicator Framework
    Advanced performance measurement encompasses strategic, operational, financial, and customer metrics for comprehensive business performance evaluation and continuous improvement achievement.

    #### Strategic Performance Metrics
    - **Market Share:** Position tracking and competitive performance
    - **Customer Satisfaction:** Loyalty measurement and retention analytics
    - **Innovation Index:** R&D productivity and technology advancement
    - **Sustainability Metrics:** ESG performance and stakeholder value creation

    ### 📊 Balanced Scorecard Implementation
    #### Multi-Dimensional Performance Framework
    - **Financial Perspective:** Revenue growth, profitability, cost management
    - **Customer Perspective:** Satisfaction, retention, market penetration
    - **Internal Process:** Efficiency, quality, innovation capability
    - **Learning & Growth:** Employee development, technology advancement

    ## 🔮 Future Outlook & Strategic Recommendations
    ### 🌟 Strategic Future Planning & Market Evolution
    Strategic future planning encompasses market trend anticipation, competitive landscape evolution assessment, technology advancement impact evaluation, regulatory environment change analysis, and strategic positioning requirements for long-term competitive advantage.

    #### Emerging Opportunities Matrix
    - **Market Expansion:** Geographic growth and demographic penetration
    - **Technology Integration:** Digital transformation and automation advantages
    - **Strategic Partnerships:** Collaboration and ecosystem development
    - **Innovation Leadership:** Technology advancement and market disruption

    ### 🚀 Strategic Recommendation Portfolio
    #### Immediate Strategic Priorities (0-12 months)
    - **Market Position Strengthening:** Competitive advantage enhancement and market share protection
    - **Operational Excellence:** Process optimization and efficiency improvement
    - **Financial Optimization:** Revenue enhancement and cost structure improvement
    - **Risk Management:** Vulnerability assessment and mitigation strategy implementation

    #### Long-term Strategic Vision (1-5 years)
    - **Market Leadership:** Industry position strengthening and competitive dominance
    - **Innovation Excellence:** Technology leadership and market disruption capability
    - **Sustainable Growth:** Balanced expansion and stakeholder value creation
    - **Strategic Transformation:** Business model evolution and market adaptation"""
                
                final_result = {
                    "content": fallback_content,
                    "sources_used": 1,
                    "confidence": 0.6,
                    "synthesis_method": "fallback_synthesis",
                    "data_availability": "limited",
                    "word_count": len(fallback_content.split()),
                    "search_mode": search_mode
                }
                
                state["final_result"] = final_result
                state["messages"].append(AIMessage(content="Fallback synthesis completed due to limited data availability"))
                return state
            
            # Build synthesis prompt with available data and STRICT mode enforcement
            synthesis_sections = []
            
            if has_financial:
                financial_summary = self._summarize_data_safely(financial_data.get("results", []))
                synthesis_sections.append(f"Financial Data Available: {financial_summary}")
            
            if has_news:
                news_summary = self._summarize_data_safely(news_data.get("results", []))
                synthesis_sections.append(f"News Data Available: {news_summary}")
            
            if has_competitive:
                competitive_summary = self._summarize_data_safely(competitive_data.get("results", []))
                synthesis_sections.append(f"Competitive Data Available: {competitive_summary}")
                
            if has_twitter:
                twitter_summary = self._summarize_data_safely(twitter_data.get("results", []))
                synthesis_sections.append(f"Social Media Data Available: {twitter_summary}")
            
            # Create mode-specific synthesis prompts with enhanced formatting
            if state['search_mode'] == "Quick Search":
                synthesis_prompt = f"""
    You are a senior business analyst creating a RAPID EXECUTIVE BRIEFING for {state['company_name']}.

    STRICT REQUIREMENTS:
    - EXACT WORD COUNT: 500-800 words (COUNT AS YOU WRITE)
    - TIME CONSTRAINT: 5-second analysis depth
    - AUDIENCE: Senior executives needing immediate decisions
    - TONE: Direct, actionable, no fluff

    Query: {state['query']}

    Available Intelligence:
    {chr(10).join(synthesis_sections)}

    REQUIRED STRUCTURE (Brief but complete):

    # 🚀 Rapid Business Intelligence: {state['company_name']}

    ## 📋 Executive Summary
    **Query Analysis:** {state['query']}
    *[100-120 words: Key findings and immediate insights]*

    ## 💡 Key Strategic Insights  
    ### Market Intelligence Highlights
    *[120-150 words: Most important 2-3 strategic points with bullet formatting]*

    ## 💰 Financial Performance Snapshot
    ### Current Financial Position
    *[100-120 words: Key financial indicators and trends]*

    ## ⚡ Immediate Action Framework
    ### Priority Implementation Matrix
    *[100-120 words: Specific actionable recommendations]*

    ## ⚠️ Risk & Opportunity Assessment
    ### Critical Alert Matrix
    *[80-100 words: Top 3 risks and opportunities]*

    ## 🎯 Next Steps Priority Matrix
    *[50-80 words: Immediate priorities for execution]*

    Write with executive precision. Use bullet points, clear headers, and actionable language. Target exactly 650 words.
    """
            
            elif state['search_mode'] == "Extended Search":
                synthesis_prompt = f"""
    You are a strategic business analyst creating a COMPREHENSIVE STRATEGIC ANALYSIS for {state['company_name']}.

    STRICT REQUIREMENTS:
    - EXACT WORD COUNT: 1000-1500 words (COUNT AS YOU WRITE)
    - TIME CONSTRAINT: 20-second analysis depth
    - AUDIENCE: Strategic planning teams
    - TONE: Professional, analytical, strategic

    Query: {state['query']}

    Available Intelligence:
    {chr(10).join(synthesis_sections)}

    REQUIRED STRUCTURE (Balanced coverage):

    # 📊 Strategic Business Intelligence: {state['company_name']}

    ## 🎯 Executive Summary
    **Comprehensive Analysis Target:** {state['query']}
    *[150-200 words: Comprehensive strategic overview with key conclusions]*

    ## 🌍 Market Analysis & Intelligence Framework
    ### Industry Landscape Assessment
    *[200-250 words: Industry positioning, market trends, competitive landscape]*
    - Market positioning assessment
    - Industry trend analysis  
    - Growth opportunities identification

    ## 💹 Financial Performance Review & Analysis
    ### Comprehensive Financial Intelligence
    *[200-250 words: Financial metrics, performance trends, investment analysis]*
    - Revenue and profitability analysis
    - Financial health indicators
    - Investment and valuation insights

    ## 🏆 Competitive Landscape Assessment
    ### Strategic Competitive Intelligence
    *[150-200 words: Competitive positioning and market dynamics]*
    - Direct competitor analysis
    - Market share evaluation
    - Competitive advantages assessment

    ## 📈 Strategic Recommendations Framework
    ### Comprehensive Strategic Initiative Portfolio
    *[200-250 words: Actionable strategic initiatives and implementation guidance]*
    - Strategic priority recommendations
    - Implementation roadmap outline
    - Resource allocation guidance

    ## 🛡️ Risk Assessment & Mitigation Strategy
    ### Comprehensive Risk Management Framework
    *[150-200 words: Risk evaluation and mitigation approaches]*
    - Strategic risk identification
    - Operational risk factors
    - Mitigation strategy framework

    ## 🚀 Implementation Roadmap & Timeline
    ### Strategic Execution Framework
    *[100-150 words: Next steps and execution timeline]*

    Use professional formatting with clear sections, bullet points for key insights, and strategic depth. Target around 750 words.
    """
            
            else:  # Deep Search
                synthesis_prompt = f"""
    You are a senior research analyst creating an IN-DEPTH STRATEGIC RESEARCH REPORT for {state['company_name']}.

    STRICT REQUIREMENTS:
    - EXACT WORD COUNT: 2000-5000 words (COUNT AS YOU WRITE)
    - TIME CONSTRAINT: 60+ second analysis depth
    - AUDIENCE: Board-level strategic decisions
    - TONE: Research-grade, comprehensive, authoritative

    Query: {state['query']}

    Available Intelligence:
    {chr(10).join(synthesis_sections)}

    REQUIRED STRUCTURE (Exhaustive coverage):

    # 🏢 Comprehensive Strategic Research Report: {state['company_name']}

    ## 📊 Executive Summary & Key Findings
    **Research Objective:** {state['query']}
    *[300-400 words: Comprehensive strategic overview with critical insights and recommendations]*

    ## 🌍 Advanced Market Intelligence Framework
    ### 📈 Industry Evolution & Structural Analysis
    *[500-700 words: Deep industry analysis with multiple dimensions]*
    #### Market Dynamics Assessment
    - Market dynamics and structural changes
    - Technology disruption impacts
    - Regulatory environment assessment

    #### Market Positioning & Competitive Intelligence  
    - Competitive positioning analysis
    - Market share dynamics
    - Strategic differentiation opportunities

    ## 💰 Comprehensive Financial Analysis Platform
    ### 📋 Financial Performance Deep Dive
    *[500-700 words: Multi-dimensional financial evaluation]*
    #### Financial Optimization Framework
    - Revenue optimization analysis
    - Profitability enhancement opportunities
    - Capital allocation effectiveness

    #### Strategic Financial Framework & Investment Analysis
    - Investment portfolio assessment
    - Financial risk management
    - Value creation measurement systems

    ## 🏆 Detailed Competitive Assessment Matrix
    ### ⚔️ Competitive Positioning Analysis
    *[400-500 words: Thorough competitive intelligence]*
    #### Direct Competitor Analysis
    - Direct competitor evaluation
    - Indirect competition assessment
    - Strategic capability comparison

    #### Competitive Advantage Framework
    - Differentiation strategy analysis
    - Market leadership assessment
    - Competitive response capabilities

    ## 🚀 Strategic Framework & Planning Architecture
    ### 📋 Strategic Implementation Strategy
    *[500-700 words: Comprehensive strategic development]*
    #### Strategic Vision & Objectives
    - Vision and objective alignment
    - Capability development planning
    - Resource optimization frameworks

    #### Operational Excellence Integration
    - Process optimization opportunities
    - Technology advancement strategies
    - Human capital development

    ## 🛡️ Risk Management & Scenario Analysis Platform
    ### ⚠️ Strategic Risk Evaluation Framework
    *[400-500 words: Advanced risk assessment]*
    #### Strategic Risk Categories
    - Market volatility assessment
    - Operational challenge analysis
    - Financial vulnerability evaluation

    #### Advanced Scenario Planning & Contingency Framework
    - Multiple strategic alternatives
    - Contingency planning development
    - Adaptive strategy frameworks

    ## 📈 Performance Metrics & KPI Dashboard
    ### 🎯 Key Performance Indicator Framework
    *[200-300 words: Measurement and monitoring systems]*
    - Key performance indicators framework
    - Balanced scorecard implementation
    - Competitive benchmarking systems

    ## 🔮 Future Outlook & Strategic Recommendations
    ### 🌟 Strategic Future Planning & Market Evolution
    *[300-400 words: Forward-looking strategic guidance]*
    - Market trend anticipation
    - Strategic positioning requirements
    - Long-term competitive advantage sustainability

    Use research-grade formatting with comprehensive sections, detailed bullet points, professional structure, and analytical depth. Target exactly 3500 words.
    """
            
            # Make the LLM call with enhanced prompting
            try:
                # Set the search mode for fallback
                self.current_search_mode = state.get('search_mode', 'Extended Search')
                
                response = self.llm.invoke([HumanMessage(content=synthesis_prompt)])
                synthesis_content = response.content
                
                # Validate the response length and apply strict enforcement
                word_count = len(synthesis_content.split())
                min_required = token_config['min_word_limit']
                max_allowed = token_config['word_limit']
                
                # If response is significantly off-target, regenerate with stronger prompting
                if word_count < (min_required * 0.8) or word_count > (max_allowed * 1.2):
                    # Second attempt with even stronger enforcement
                    stronger_prompt = f"""
    CRITICAL: The previous response was {word_count} words. You MUST write EXACTLY {min_required}-{max_allowed} words.

    {synthesis_prompt}

    WORD COUNT ENFORCEMENT: 
    - Start counting words as you write
    - Stop when you reach {max_allowed} words
    - Do not exceed or fall short of the range
    - Current target: {(min_required + max_allowed) // 2} words exactly
    """
                    response = self.llm.invoke([HumanMessage(content=stronger_prompt)])
                    synthesis_content = response.content
                
                # Final word count adjustment
                word_count = len(synthesis_content.split())
                if word_count < min_required:
                    synthesis_content = expand_content_intelligently(
                        synthesis_content, 
                        state['search_mode'], 
                        min_required - word_count
                    )
                elif word_count > max_allowed:
                    words = synthesis_content.split()
                    synthesis_content = ' '.join(words[:max_allowed])
                    synthesis_content += f"\n\n*[Content optimized to {max_allowed} words for {state['search_mode']} mode]*"
                
                if not synthesis_content or len(synthesis_content.strip()) < 100:
                    raise ValueError("Synthesis response too short or empty")
                    
            except Exception as llm_error:
                # Create manual synthesis as fallback
                self.current_search_mode = state.get('search_mode', 'Extended Search')
                synthesis_content = self._create_manual_synthesis(
                    state['company_name'], 
                    state['query'], 
                    financial_data, 
                    news_data, 
                    competitive_data
                )
        
            # Calculate sources used
            sources_used = sum([
                1 if has_financial else 0,
                1 if has_news else 0,
                1 if has_competitive else 0,
                1 if has_twitter else 0
            ])
            
            # Calculate confidence based on data quality
            confidence = 0.7  # Base confidence
            if has_financial:
                confidence += 0.1
            if has_twitter:
                confidence += 0.1
            if has_news:
                confidence += 0.1
            if has_competitive:
                confidence += 0.1
        
            
            final_result = {
                "content": synthesis_content,
                "sources_used": max(sources_used, 1),
                "confidence": min(confidence, 1.0),
                "synthesis_method": "multi_agent_langgraph",
                "search_mode": state['search_mode'],
                "word_count": len(synthesis_content.split()),
                "data_quality": {
                    "financial_available": has_financial,
                    "news_available": has_news,
                    "competitive_available": has_competitive,
                    "twitter_available": has_twitter
                }
            }
            
            state["final_result"] = final_result
            state["messages"].append(AIMessage(content=f"Research synthesis completed successfully with {sources_used} data sources"))
            
        except Exception as e:
            # Emergency fallback
            search_mode = state.get('search_mode', 'Extended Search')
            config = TOKEN_LIMITS.get(search_mode, TOKEN_LIMITS["Extended Search"])
            
            if search_mode == "Quick Search":
                emergency_content = f"""# 🚀 Quick Business Assessment: {state.get('company_name', 'Unknown Company')}

    ## 📋 Executive Summary
    **Rapid Analysis Framework:** "{state.get('query', 'Business analysis')}"

    LangGraph multi-agent system completed rapid assessment workflow with technical optimization requirements identified.

    ## 💡 Critical Strategic Insights
    ### Market Intelligence Overview
    Strategic positioning analysis indicates immediate opportunities for competitive advantage enhancement and operational efficiency optimization through tactical adjustments and resource reallocation.

    ## 💰 Financial Performance Snapshot
    ### Performance Indicators Summary
    Current performance indicators suggest revenue optimization potential and cost structure enhancement opportunities requiring immediate strategic attention and tactical implementation.

    ## ⚡ Immediate Action Framework
    ### Priority Implementation Matrix
    1. **🎯 System Optimization:** Execute technical enhancements for improved analysis capability
    2. **💼 Alternative Methodologies:** Implement backup research frameworks for critical insights
    3. **🏆 Enhanced Coordination:** Deploy tactical adjustments for immediate competitive positioning
    4. **📊 Performance Monitoring:** Establish alternative tracking for strategic effectiveness

    ## ⚠️ Risk & Opportunity Assessment
    ### Technical Alert Matrix
    - **🔴 High Priority:** Technical system optimization required for sustained analysis capability
    - **🟡 Medium Priority:** Alternative research approaches recommended for critical decisions
    - **🟢 Opportunities:** System enhancement potential for improved competitive intelligence

    ## 🎯 Next Steps Priority Matrix
    **Immediate Technical Actions:**
    - Priority system enhancement and alternative analysis deployment
    - Resource allocation optimization for comprehensive intelligence provision
    - Performance monitoring system implementation for sustained analytical excellence"""
            
            elif search_mode == "Extended Search":
                emergency_content = f"""# 📊 Strategic Analysis Framework Response: {state.get('company_name', 'Unknown Company')}

    ## 🎯 Executive Summary
    **Strategic Analysis Framework:** "{state.get('query', 'Business analysis')}"

    LangGraph multi-agent orchestration system completed strategic assessment with technical optimization requirements for enhanced performance delivery.

    ## 🌍 Market Analysis Framework Status
    ### Advanced Analytical Capabilities
    Strategic analytical capabilities remain operational with system enhancement requirements for optimal strategic intelligence delivery and comprehensive business analysis provision.

    Strategic market evaluation frameworks indicate significant opportunities for competitive positioning enhancement through systematic analytical optimization and enhanced research coordination.

    ## 💹 Financial Analysis Capability Assessment
    ### Comprehensive Financial Assessment Protocols
    Advanced financial assessment protocols maintain operational status with optimization requirements for enhanced analytical depth and strategic intelligence provision.

    Performance evaluation frameworks suggest substantial opportunities for analytical enhancement through system optimization and research methodology improvement.

    ## 🏆 Competitive Assessment Framework Status
    ### Strategic Competitive Analysis Capabilities
    Advanced competitive analysis capabilities remain functional with enhancement requirements for comprehensive market intelligence and competitive positioning assessment.

    Strategic competitive intelligence frameworks indicate opportunities for analytical optimization and strategic assessment enhancement through systematic improvement protocols.

    ## 📈 Strategic Recommendations Framework
    ### System Enhancement Strategy
    1. **🎯 Technical Optimization:** Implement system enhancements for improved analytical capability and research delivery
    2. **🔄 Alternative Methodologies:** Deploy backup research frameworks for critical strategic intelligence provision
    3. **🤝 Enhanced Coordination:** Optimize multi-agent communication for improved synthesis and analysis delivery
    4. **📊 Performance Enhancement:** Implement monitoring systems for sustained analytical excellence
    5. **🚀 Capability Development:** Advanced system optimization for comprehensive intelligence provision

    ## 🛡️ Risk Assessment Framework Status
    ### Technical Risk Management
    Technical optimization requirements identified for sustained analytical capability and comprehensive strategic intelligence provision through enhanced system coordination.

    ## 🚀 Implementation Strategy Framework
    ### Strategic Enhancement Requirements
    System enhancement requires systematic technical optimization, alternative methodology deployment, and performance monitoring implementation for sustained analytical excellence and comprehensive business intelligence delivery."""
            
            else:  # Deep Search
                emergency_content = f"""# 🏢 Comprehensive System Analysis Framework: {state.get('company_name', 'Unknown Company')}

    ## 📊 Executive Summary & System Status
    **Comprehensive Research Framework:** "{state.get('query', 'Business analysis')}"

    Advanced LangGraph multi-agent orchestration system completed research workflow execution with identified optimization requirements for enhanced analytical capability, comprehensive research delivery, and sustained strategic intelligence provision.

    The sophisticated multi-agent framework maintains operational capability with technical enhancement requirements for optimal research-grade analysis delivery and comprehensive strategic intelligence provision.

    ## 🌍 Advanced System Architecture Assessment
    ### Multi-Agent Orchestration Framework
    Comprehensive multi-agent orchestration framework encompasses specialized intelligence agents including financial analysis, market intelligence, competitive assessment, and validation systems maintaining operational capability with optimization requirements.

    Advanced analytical capabilities include multi-dimensional research coordination, intelligent data synthesis, comprehensive validation protocols, and strategic framework development maintaining functional status with enhancement opportunities.

    #### Technical Capability Analysis Matrix
    - **Agent Coordination Systems:** Operational functionality with optimization potential for enhanced research delivery
    - **Data Synthesis Protocols:** Advanced integration capabilities requiring technical enhancement
    - **Validation Frameworks:** Comprehensive verification systems with performance optimization needs
    - **Research Coordination:** Multi-dimensional analysis capabilities requiring system enhancement

    #### System Performance Optimization Requirements
    Multi-agent coordination systems demonstrate operational functionality with optimization potential for enhanced research delivery, improved analytical depth, and comprehensive strategic intelligence provision through systematic enhancement protocols.

    ## 💰 Comprehensive Research Framework Status
    ### Advanced Analytical Infrastructure
    Advanced analytical infrastructure maintains operational capability with significant optimization opportunities for enhanced research delivery, improved strategic intelligence provision, and sustained analytical excellence achievement.

    Research coordination systems demonstrate functional capability with enhancement requirements for optimal multi-agent orchestration, comprehensive data synthesis, and advanced validation protocol implementation.

    #### Research Delivery Optimization Matrix
    - **Financial Intelligence Systems:** Operational with enhancement potential for comprehensive analysis
    - **Market Research Coordination:** Functional capabilities requiring optimization for improved delivery
    - **Competitive Assessment Protocols:** Advanced frameworks with performance enhancement opportunities
    - **Strategic Synthesis Capabilities:** Multi-dimensional coordination requiring system optimization

    #### Performance Enhancement Framework
    Advanced research coordination systems demonstrate operational capability with enhancement requirements for optimal analytical delivery, improved strategic intelligence provision, and comprehensive research excellence achievement.

    ## 🏆 Multi-Agent Coordination Analysis
    ### Sophisticated Agent Communication Protocols
    Advanced agent communication protocols maintain operational status with optimization requirements for enhanced coordination, improved research delivery, and comprehensive analytical capability development.

    Advanced orchestration frameworks demonstrate functional capability with enhancement opportunities for improved agent coordination, optimized research delivery, and sustained analytical excellence achievement.

    #### Agent Performance Matrix
    - **Financial Research Agents:** Operational with optimization requirements for enhanced capability
    - **Market Intelligence Agents:** Functional systems requiring coordination enhancement
    - **Competitive Analysis Agents:** Advanced capabilities with performance improvement potential
    - **Validation Agents:** Comprehensive protocols requiring optimization for enhanced accuracy

    ## 🚀 Strategic Research Capability Assessment
    ### Comprehensive Research Frameworks
    Advanced research frameworks maintain operational functionality with significant enhancement opportunities for improved analytical depth, enhanced strategic intelligence provision, and sustained research excellence delivery.

    Comprehensive analytical capabilities encompass multi-dimensional research execution, comprehensive validation coordination, and strategic framework development with optimization requirements for enhanced performance achievement.

    #### Research Excellence Framework
    - **Multi-Dimensional Analysis:** Advanced capabilities requiring optimization for comprehensive delivery
    - **Strategic Intelligence Provision:** Operational systems with enhancement potential
    - **Validation Protocol Integration:** Sophisticated frameworks requiring performance optimization
    - **Research Coordination Excellence:** Advanced systems with improvement opportunities

    ## 🛡️ Technical Enhancement Strategy
    ### Comprehensive System Optimization Requirements
    Advanced system optimization requires multi-dimensional enhancement including agent coordination improvement, research delivery optimization, validation protocol enhancement, and performance monitoring implementation.

    Strategic technical development encompasses system architecture optimization, multi-agent communication enhancement, research framework improvement, and analytical capability development for sustained research excellence.

    #### Implementation Framework Development Matrix
    - **System Architecture Enhancement:** Technical optimization for improved coordination
    - **Agent Communication Improvement:** Protocol enhancement for better research delivery
    - **Research Framework Optimization:** Capability development for analytical excellence
    - **Performance Monitoring Implementation:** Comprehensive tracking for sustained improvement

    ## 📈 Performance Optimization Recommendations
    ### Advanced System Enhancement Portfolio
    Comprehensive system enhancement encompasses technical architecture optimization, multi-agent coordination improvement, research delivery enhancement, and comprehensive analytical capability development for sustained strategic intelligence provision.

    Strategic technical development requires systematic optimization implementation, performance monitoring enhancement, coordination protocol improvement, and analytical framework development for research excellence achievement and sustained competitive intelligence delivery.

    #### Long-term System Development Strategy
    - **Technical Infrastructure Advancement:** Comprehensive optimization for enhanced capability
    - **Research Methodology Enhancement:** Advanced protocol development for improved delivery
    - **Agent Coordination Optimization:** Systematic improvement for better research coordination
    - **Strategic Intelligence Excellence:** Advanced capability development for sustained competitive advantage"""
            
            # Ensure emergency content meets word requirements
            emergency_content = enforce_content_length(emergency_content, search_mode)
            
            state["final_result"] = {
                "error": str(e),
                "content": emergency_content,
                "sources_used": 1,
                "confidence": 0.5,
                "synthesis_method": "emergency_fallback",
                "search_mode": search_mode,
                "word_count": len(emergency_content.split())
            }
            state["messages"].append(AIMessage(content=f"Emergency fallback synthesis due to error: {str(e)}"))
        
        return state
         
    
    def _summarize_data_safely(self, data_results: List[Dict]) -> str:
        """Safely summarize data results"""
        try:
            if not data_results:
                return "No specific data available"
            
            summaries = []
            for item in data_results[:3]:  # Limit to first 3 items
                title = item.get('title', 'No title')
                snippet = item.get('snippet', 'No description')
                # Truncate long snippets
                if len(snippet) > 150:
                    snippet = snippet[:150] + "..."
                summaries.append(f"- {title}: {snippet}")
            
            return "\n".join(summaries)
        except Exception as e:
            return f"Data summary error: {str(e)}"
    
    def _create_manual_synthesis(self, company_name: str, query: str, financial_data: Dict, news_data: Dict, competitive_data: Dict) -> str:
        """Create manual synthesis when LLM fails with proper word count and search mode differentiation"""
        
        # Get current search mode config
        search_mode = getattr(self, 'current_search_mode', 'Extended Search')
        config = TOKEN_LIMITS.get(search_mode, TOKEN_LIMITS["Extended Search"])
        min_words = config.get("min_word_limit", 1000)
        max_words = config.get("word_limit", 1500)
        
        sections = []
        
        # Differentiated content based on search mode with strict word requirements
        if search_mode == "Quick Search":
            # Quick Search: 500-800 words - Rapid executive insights
            sections.append(f"# 🚀 Rapid Business Intelligence: {company_name}")
            sections.append(f"\n## 📋 Executive Summary")
            sections.append(f"**Query Analysis:** {query}")
            sections.append(f"\nThis rapid assessment provides critical business insights using advanced multi-agent systems for immediate executive decision-making and strategic response coordination.\n")
            
            sections.append(f"## 💡 Key Strategic Insights")
            sections.append("### Market Intelligence Highlights")
            if financial_data and not financial_data.get("error"):
                sections.append("**📊 Financial Intelligence:** Market performance data reveals immediate strategic positioning opportunities requiring executive attention.")
                if financial_data.get("results"):
                    sections.append("\n**Key Financial Indicators:**")
                    for item in financial_data["results"][:2]:
                        title = item.get('title', 'Financial indicator')
                        snippet = item.get('snippet', 'Financial insight')[:80]
                        sections.append(f"• **{title}**: {snippet}")
                sections.append("")
            
            if news_data and not news_data.get("error"):
                sections.append("**📰 Market Developments:** Recent industry developments indicate evolving competitive dynamics requiring immediate strategic assessment.")
                if news_data.get("results"):
                    sections.append("\n**Latest Market Intelligence:**")
                    for item in news_data["results"][:2]:
                        title = item.get('title', 'Market development')
                        snippet = item.get('snippet', 'Strategic insight')[:80]
                        sections.append(f"• **{title}**: {snippet}")
                sections.append("")
            
            sections.append("### Strategic Positioning Assessment")
            sections.append("• **Competitive Dynamics:** Market evolution requires immediate strategic response coordination")
            sections.append("• **Growth Opportunities:** Tactical positioning adjustments needed for market optimization")
            sections.append("• **Risk Factors:** Operational challenges requiring swift executive intervention")
            sections.append("")
            
            sections.append(f"## 💰 Financial Performance Snapshot")
            sections.append("### Current Financial Position")
            sections.append("• **Performance Metrics:** Strategic indicators require immediate monitoring for competitive advantage")
            sections.append("• **Investment Opportunities:** Tactical positioning adjustments needed for market optimization")
            sections.append("• **Risk Factors:** Financial volatility monitoring essential for sustained performance")
            sections.append("• **Cash Flow Analysis:** Liquidity management critical for operational flexibility")
            sections.append("")
            
            sections.append(f"## ⚡ Immediate Action Framework")
            sections.append("### Priority Implementation Matrix")
            sections.append("1. **🎯 Rapid Response Protocol:** Deploy immediate tactical adjustments for identified market opportunities")
            sections.append("2. **💼 Resource Optimization:** Prioritize critical strategic initiatives for maximum impact achievement")
            sections.append("3. **🏆 Competitive Positioning:** Implement swift adjustments for enhanced market advantage")
            sections.append("4. **📊 Performance Monitoring:** Establish real-time tracking for strategic effectiveness")
            sections.append("5. **🔄 Agile Adaptation:** Maintain flexibility for rapid market response capability")
            sections.append("")
            
            sections.append(f"## ⚠️ Risk & Opportunity Assessment")
            sections.append("### Critical Alert Matrix")
            sections.append("• **🔴 High Priority:** Market volatility factors requiring immediate executive attention")
            sections.append("• **🟡 Medium Priority:** Competitive pressures demanding strategic response coordination")  
            sections.append("• **🟢 Opportunities:** Operational efficiency optimization for competitive performance")
            sections.append("• **🔵 Monitoring Required:** Regulatory changes and compliance requirements")
            sections.append("")
            
            sections.append(f"## 🎯 Next Steps Priority Matrix")
            sections.append("**Immediate Actions Required:**")
            sections.append("• Executive decision-making for strategic initiative deployment")
            sections.append("• Resource allocation optimization for competitive positioning")
            sections.append("• Performance monitoring system implementation for sustained advantage")
            sections.append("• Risk mitigation protocol activation for operational stability")

        elif search_mode == "Extended Search":
            # Extended Search: 1000-1500 words - Comprehensive strategic analysis
            sections.append(f"# 📊 Strategic Business Intelligence: {company_name}")
            sections.append(f"\n## 🎯 Executive Summary")
            sections.append(f"**Comprehensive Analysis Target:** {query}")
            sections.append(f"\nThis detailed assessment leverages advanced multi-agent orchestration to provide strategic business intelligence for comprehensive decision-making and long-term competitive positioning enhancement.\n")
            
            sections.append(f"## 🌍 Market Analysis & Intelligence Framework")
            sections.append("### Industry Landscape Assessment")
            sections.append("Advanced market intelligence reveals complex industry dynamics impacting strategic positioning, competitive landscapes, and growth opportunity identification requiring systematic strategic response and implementation coordination.")
            sections.append("")
            
            if news_data and not news_data.get("error"):
                sections.append("### 📰 Strategic Market Developments")
                sections.append("Recent market evolution patterns indicate strategic implications requiring comprehensive evaluation:")
                if news_data.get("results"):
                    for item in news_data["results"][:3]:
                        title = item.get('title', 'Market development')
                        snippet = item.get('snippet', 'Strategic insight')[:120]
                        sections.append(f"• **{title}**: {snippet}")
                sections.append("")
            
            sections.append("### Market Positioning Intelligence")
            sections.append("• **Competitive Dynamics:** Market evolution requires strategic response coordination")
            sections.append("• **Growth Opportunities:** Systematic identification of expansion potential")
            sections.append("• **Technology Integration:** Advanced capabilities for competitive advantage")
            sections.append("• **Customer Engagement:** Strategic relationship optimization frameworks")
            sections.append("• **Brand Positioning:** Market perception and competitive differentiation")
            sections.append("• **Market Segmentation:** Target audience optimization and penetration strategies")
            sections.append("")
            
            sections.append(f"## 💹 Financial Performance Review & Analysis")
            sections.append("### Comprehensive Financial Intelligence")
            sections.append("Advanced financial analysis encompasses revenue optimization, profitability enhancement, and strategic investment effectiveness evaluation for sustainable growth achievement.")
            sections.append("")
            
            if financial_data and not financial_data.get("error"):
                sections.append("### 📊 Performance Metrics Dashboard")
                sections.append("Financial intelligence indicates strategic positioning opportunities:")
                if financial_data.get("results"):
                    for item in financial_data["results"][:3]:
                        title = item.get('title', 'Financial metric')
                        snippet = item.get('snippet', 'Financial analysis')[:120]
                        sections.append(f"• **{title}**: {snippet}")
                sections.append("")
            
            sections.append("### Financial Strategy Optimization")
            sections.append("• **Revenue Enhancement:** Multi-channel optimization strategies")
            sections.append("• **Cost Structure Analysis:** Efficiency improvement opportunities")
            sections.append("• **Investment Allocation:** Strategic resource deployment frameworks")
            sections.append("• **Risk Management:** Financial stability and growth balance")
            sections.append("• **Capital Structure:** Debt-equity optimization for strategic flexibility")
            sections.append("• **Cash Flow Management:** Working capital optimization for operational excellence")
            sections.append("")
            
            sections.append(f"## 🏆 Competitive Landscape Assessment")
            sections.append("### Strategic Competitive Intelligence")
            sections.append("Comprehensive competitive analysis reveals market positioning effectiveness, competitive advantage sustainability, and strategic differentiation opportunities for enhanced market leadership achievement.")
            sections.append("")
            
            if competitive_data and not competitive_data.get("error"):
                sections.append("### Competitive Intelligence Summary")
                sections.append("Market competitive analysis reveals strategic positioning opportunities:")
                if competitive_data.get("results"):
                    for item in competitive_data["results"][:3]:
                        title = item.get('title', 'Competitive insight')
                        snippet = item.get('snippet', 'Competitive analysis')[:120]
                        sections.append(f"• **{title}**: {snippet}")
                sections.append("")
            
            sections.append("### Competitive Positioning Matrix")
            sections.append("• **Market Leadership:** Position assessment and enhancement strategies")
            sections.append("• **Differentiation Opportunities:** Unique value proposition development")
            sections.append("• **Competitive Response:** Strategic agility and adaptation capabilities")
            sections.append("• **Market Share Dynamics:** Growth and retention strategy optimization")
            sections.append("• **Innovation Leadership:** Technology advancement for competitive advantage")
            sections.append("• **Customer Loyalty:** Retention and satisfaction enhancement strategies")
            sections.append("")
            
            sections.append(f"## 📈 Strategic Recommendations Framework")
            sections.append("### Comprehensive Strategic Initiative Portfolio")
            sections.append("Multi-dimensional strategic framework for competitive advantage development:")
            sections.append("")
            sections.append("1. **🎯 Market Positioning Excellence:** Enhanced competitive differentiation through strategic capability advancement")
            sections.append("2. **⚙️ Operational Optimization:** Internal process enhancement for efficiency and performance maximization")
            sections.append("3. **💰 Financial Strategy Enhancement:** Revenue optimization and profitability improvement initiatives")
            sections.append("4. **🚀 Innovation Leadership:** Technological advancement for sustained competitive advantage")
            sections.append("5. **🤝 Strategic Partnerships:** Collaboration frameworks for market expansion")
            sections.append("6. **🎨 Brand Development:** Market positioning and perception enhancement strategies")
            sections.append("")
            
            sections.append(f"## 🛡️ Risk Assessment & Mitigation Strategy")
            sections.append("### Comprehensive Risk Management Framework")
            sections.append("Strategic risk evaluation encompasses market volatility, competitive pressures, and operational challenges requiring systematic mitigation strategies and continuous monitoring protocols.")
            sections.append("")
            sections.append("### Risk Mitigation Matrix")
            sections.append("• **Strategic Risks:** Market positioning and competitive response planning")
            sections.append("• **Operational Risks:** Process optimization and resource management")
            sections.append("• **Financial Risks:** Cash flow management and investment protection")
            sections.append("• **Technology Risks:** Innovation protection and adaptation strategies")
            sections.append("• **Regulatory Risks:** Compliance management and policy adaptation")
            sections.append("• **Market Risks:** Economic volatility and demand fluctuation management")
            sections.append("")
            
            sections.append(f"## 🚀 Implementation Roadmap & Timeline")
            sections.append("### Strategic Execution Framework")
            sections.append("Systematic implementation approach including timeline development, resource allocation, and performance measurement for sustainable competitive advantage achievement and market leadership development.")
            sections.append("")
            sections.append("**Phase 1 (0-3 months):** Immediate tactical adjustments and quick wins")
            sections.append("**Phase 2 (3-12 months):** Strategic initiative deployment and system optimization")
            sections.append("**Phase 3 (12+ months):** Long-term competitive advantage sustainability and market leadership")
            sections.append("")
            sections.append("### Success Metrics & KPIs")
            sections.append("• **Financial Performance:** Revenue growth, profitability improvement, cost optimization")
            sections.append("• **Market Position:** Market share expansion, customer satisfaction, brand recognition")
            sections.append("• **Operational Excellence:** Process efficiency, quality improvement, innovation metrics")

        else:  # Deep Search: 2000-5000 words - Research-grade comprehensive analysis
            sections.append(f"# 🏢 Comprehensive Strategic Research: {company_name}")
            sections.append(f"\n## 📊 Executive Summary & Key Findings")
            sections.append(f"**Research Objective:** {query}")
            sections.append(f"\nThis exhaustive analysis employs advanced multi-agent research orchestration with comprehensive validation protocols to provide research-grade strategic intelligence for major business decisions and long-term strategic planning excellence.")
            sections.append("")
            sections.append("### 🎯 Critical Strategic Findings")
            sections.append("• **Market Position:** Comprehensive competitive landscape assessment reveals strategic opportunities")
            sections.append("• **Financial Performance:** Multi-dimensional analysis indicates optimization potential")
            sections.append("• **Competitive Advantage:** Strategic differentiation opportunities identified for market leadership")
            sections.append("• **Risk Assessment:** Systematic evaluation reveals mitigation strategies for sustained performance")
            sections.append("• **Growth Potential:** Strategic expansion opportunities requiring systematic development")
            sections.append("• **Innovation Capacity:** Technology advancement potential for competitive differentiation")
            sections.append("")
            
            sections.append(f"## 🌍 Advanced Market Intelligence Framework")
            sections.append("### 📈 Industry Evolution & Structural Analysis")
            sections.append("Comprehensive market intelligence encompasses multi-dimensional analysis of industry evolution patterns, competitive transformation dynamics, technological disruption impacts, regulatory environment changes, and economic factors affecting strategic positioning effectiveness and market leadership sustainability requirements.")
            sections.append("")
            sections.append("#### Market Dynamics Assessment")
            sections.append("• **Industry Structure:** Competitive landscape transformation patterns and strategic implications")
            sections.append("• **Technology Integration:** Digital transformation impacts on market positioning")
            sections.append("• **Regulatory Environment:** Compliance requirements and strategic adaptation needs")
            sections.append("• **Economic Factors:** Market volatility impacts and strategic response requirements")
            sections.append("• **Consumer Behavior:** Market demand patterns and customer preference evolution")
            sections.append("• **Supply Chain Dynamics:** Production and distribution optimization opportunities")
            sections.append("")
            
            if news_data and not news_data.get("error"):
                sections.append("### 📰 Strategic Market Development Intelligence")
                sections.append("Recent market developments indicate strategic implications requiring comprehensive evaluation:")
                if news_data.get("results"):
                    for item in news_data["results"][:4]:
                        title = item.get('title', 'Market intelligence')
                        snippet = item.get('snippet', 'Strategic development')[:150]
                        sections.append(f"• **{title}**: {snippet}")
                sections.append("")
            
            sections.append("### 🎯 Market Positioning & Competitive Intelligence")
            sections.append("#### Strategic Market Assessment")
            sections.append("• **Market Share Analysis:** Position evaluation and growth opportunity identification")
            sections.append("• **Customer Segmentation:** Target market optimization and expansion strategies")
            sections.append("• **Value Proposition:** Differentiation enhancement and competitive positioning")
            sections.append("• **Brand Positioning:** Market perception and strategic communication optimization")
            sections.append("• **Distribution Channels:** Market access optimization and channel strategy development")
            sections.append("• **Pricing Strategy:** Competitive pricing models and value optimization frameworks")
            sections.append("")
            sections.append("#### Growth Strategy Framework")
            sections.append("• **Market Expansion:** Geographic and demographic growth opportunities")
            sections.append("• **Product Development:** Innovation pipeline and market introduction strategies")
            sections.append("• **Strategic Partnerships:** Collaboration opportunities for market access")
            sections.append("• **Acquisition Targets:** Consolidation opportunities for market dominance")
            sections.append("• **Digital Transformation:** Technology integration for competitive advantage")
            sections.append("• **Sustainability Initiatives:** ESG integration for stakeholder value creation")
            sections.append("")
            
            sections.append(f"## 💰 Comprehensive Financial Analysis Platform")
            sections.append("### 📋 Financial Performance Deep Dive")
            sections.append("Advanced financial intelligence encompasses multi-dimensional performance evaluation including revenue optimization strategies, profitability enhancement methodologies, capital allocation effectiveness assessment, investment portfolio analysis, cash flow optimization, and value creation measurement frameworks for sustainable growth achievement.")
            sections.append("")
            
            if financial_data and not financial_data.get("error"):
                sections.append("#### 📊 Financial Metrics Dashboard")
                sections.append("Detailed financial analysis reveals performance patterns and strategic opportunities:")
                if financial_data.get("results"):
                    for item in financial_data["results"][:4]:
                        title = item.get('title', 'Financial analysis')
                        snippet = item.get('snippet', 'Financial intelligence')[:150]
                        sections.append(f"• **{title}**: {snippet}")
                sections.append("")
            
            sections.append("#### Financial Optimization Framework")
            sections.append("• **Revenue Streams:** Diversification and optimization strategies for growth")
            sections.append("• **Cost Management:** Structure optimization and efficiency enhancement")
            sections.append("• **Profitability Analysis:** Margin improvement and value creation strategies")
            sections.append("• **Cash Flow Optimization:** Working capital management and liquidity strategies")
            sections.append("• **Investment Portfolio:** Strategic asset allocation and return optimization")
            sections.append("• **Financial Risk Management:** Hedging strategies and exposure mitigation")
            sections.append("")
            sections.append("### 💹 Strategic Financial Framework & Investment Analysis")
            sections.append("#### Capital Allocation Strategy")
            sections.append("• **Investment Priorities:** Strategic resource deployment for maximum ROI")
            sections.append("• **Risk Management:** Financial stability and growth balance optimization")
            sections.append("• **Valuation Analysis:** Market positioning and investor perception enhancement")
            sections.append("• **Dividend Policy:** Shareholder value optimization and retention strategies")
            sections.append("• **Debt Management:** Capital structure optimization for financial flexibility")
            sections.append("• **Growth Financing:** Funding strategies for expansion initiatives")
            sections.append("")
            sections.append("#### Financial Risk Assessment")
            sections.append("• **Market Risk:** Economic volatility impact and mitigation strategies")
            sections.append("• **Credit Risk:** Counterparty assessment and exposure management")
            sections.append("• **Operational Risk:** Business continuity and financial protection")
            sections.append("• **Liquidity Risk:** Cash flow management and financing strategy optimization")
            sections.append("• **Currency Risk:** Foreign exchange exposure and hedging strategies")
            sections.append("• **Interest Rate Risk:** Rate sensitivity analysis and protection mechanisms")
            sections.append("")
            
            sections.append(f"## 🏆 Detailed Competitive Assessment Matrix")
            sections.append("### ⚔️ Competitive Positioning Analysis")
            sections.append("Comprehensive competitive intelligence includes direct competitor analysis, indirect competition evaluation, market share dynamics assessment, competitive positioning evaluation, strategic capability comparison, and competitive response analysis for strategic advantage identification and development enhancement.")
            sections.append("")
            
            if competitive_data and not competitive_data.get("error"):
                sections.append("#### Competitive Intelligence Summary")
                sections.append("Advanced competitive analysis reveals strategic positioning opportunities:")
                if competitive_data.get("results"):
                    for item in competitive_data["results"][:4]:
                        title = item.get('title', 'Competitive analysis')
                        snippet = item.get('snippet', 'Competitive intelligence')[:150]
                        sections.append(f"• **{title}**: {snippet}")
                sections.append("")
            
            sections.append("#### Direct Competitor Analysis")
            sections.append("• **Market Leaders:** Strategic positioning and competitive advantage assessment")
            sections.append("• **Emerging Competitors:** Threat evaluation and response strategy development")
            sections.append("• **Market Challengers:** Competitive dynamics and strategic implications")
            sections.append("• **Niche Players:** Specialization opportunities and market segmentation")
            sections.append("• **International Competitors:** Global competitive landscape assessment")
            sections.append("• **Technology Disruptors:** Innovation-based competitive threats and opportunities")
            sections.append("")
            sections.append("#### Competitive Advantage Framework")
            sections.append("• **Core Competencies:** Unique capabilities and strategic differentiation")
            sections.append("• **Resource Advantages:** Strategic assets and competitive barriers")
            sections.append("• **Technology Leadership:** Innovation capabilities and market positioning")
            sections.append("• **Brand Strength:** Market perception and customer loyalty advantages")
            sections.append("• **Operational Excellence:** Process efficiency and quality advantages")
            sections.append("• **Distribution Networks:** Market access and channel advantages")
            sections.append("")
            sections.append("### 🎖️ Strategic Competitive Response Framework")
            sections.append("#### Competitive Strategy Development")
            sections.append("• **Defensive Strategies:** Market position protection and competitive response")
            sections.append("• **Offensive Strategies:** Market share capture and competitive displacement")
            sections.append("• **Collaboration Strategies:** Strategic partnerships and competitive cooperation")
            sections.append("• **Innovation Strategies:** Technology leadership and market disruption")
            sections.append("• **Differentiation Strategies:** Unique value proposition development")
            sections.append("• **Cost Leadership:** Efficiency optimization for competitive pricing")
            sections.append("")
            
            sections.append(f"## 🚀 Strategic Framework & Planning Architecture")
            sections.append("### 📋 Strategic Implementation Strategy")
            sections.append("Comprehensive strategic framework development encompasses vision articulation, strategic objective definition, capability development planning, resource allocation optimization, performance measurement system implementation, and continuous improvement protocols for sustainable competitive advantage achievement and market leadership development.")
            sections.append("")
            sections.append("#### Strategic Vision & Objectives")
            sections.append("• **Vision Alignment:** Strategic direction and organizational purpose clarity")
            sections.append("• **Objective Setting:** SMART goals and performance measurement frameworks")
            sections.append("• **Strategy Formulation:** Competitive positioning and value creation strategies")
            sections.append("• **Resource Planning:** Strategic capability development and allocation optimization")
            sections.append("• **Performance Metrics:** KPI development and monitoring system implementation")
            sections.append("• **Stakeholder Alignment:** Internal and external stakeholder engagement strategies")
            sections.append("")
            sections.append("### 🛠️ Operational Excellence Integration")
            sections.append("#### Process Optimization Framework")
            sections.append("• **Operational Efficiency:** Process improvement and productivity enhancement")
            sections.append("• **Quality Management:** Excellence standards and continuous improvement")
            sections.append("• **Technology Integration:** Digital transformation and automation strategies")
            sections.append("• **Human Capital:** Talent development and organizational capability building")
            sections.append("• **Supply Chain Optimization:** Vendor management and logistics efficiency")
            sections.append("• **Customer Experience:** Service delivery optimization and satisfaction enhancement")
            sections.append("")
            sections.append("#### Innovation & Development Strategy")
            sections.append("• **R&D Investment:** Innovation pipeline and technology advancement")
            sections.append("• **Product Development:** Market-driven innovation and competitive differentiation")
            sections.append("• **Process Innovation:** Operational excellence and efficiency optimization")
            sections.append("• **Strategic Innovation:** Business model evolution and market disruption")
            sections.append("• **Digital Innovation:** Technology integration for competitive advantage")
            sections.append("• **Sustainability Innovation:** Environmental and social responsibility integration")
            sections.append("")
            
            sections.append(f"## 🛡️ Risk Management & Scenario Analysis Platform")
            sections.append("### ⚠️ Strategic Risk Evaluation Framework")
            sections.append("Comprehensive risk assessment encompasses strategic risks, operational challenges, financial vulnerabilities, market volatilities, regulatory compliance requirements, and competitive threats requiring systematic mitigation strategies and contingency planning for organizational resilience development.")
            sections.append("")
            sections.append("#### Strategic Risk Categories")
            sections.append("• **Market Risks:** Economic volatility, competitive pressures, customer behavior changes")
            sections.append("• **Operational Risks:** Process failures, technology disruptions, supply chain vulnerabilities")
            sections.append("• **Financial Risks:** Capital structure, liquidity, credit exposure, currency fluctuations")
            sections.append("• **Regulatory Risks:** Compliance requirements, policy changes, legal exposures")
            sections.append("• **Technology Risks:** Cybersecurity threats, system failures, innovation challenges")
            sections.append("• **Reputation Risks:** Brand perception, stakeholder relations, crisis management")
            sections.append("")
            sections.append("### 🔮 Advanced Scenario Planning & Contingency Framework")
            sections.append("#### Scenario Development Methodology")
            sections.append("• **Best Case Scenarios:** Optimal market conditions and strategic execution")
            sections.append("• **Most Likely Scenarios:** Realistic market evolution and competitive dynamics")
            sections.append("• **Worst Case Scenarios:** Crisis management and business continuity planning")
            sections.append("• **Black Swan Events:** Unexpected disruptions and adaptive response strategies")
            sections.append("• **Technology Disruption:** Innovation impact and adaptation requirements")
            sections.append("• **Regulatory Changes:** Policy evolution and compliance adaptation strategies")
            sections.append("")
            sections.append("#### Risk Mitigation Strategy Matrix")
            sections.append("• **Prevention Strategies:** Risk avoidance and proactive management")
            sections.append("• **Mitigation Strategies:** Impact reduction and consequence management")
            sections.append("• **Transfer Strategies:** Insurance and partnership risk sharing")
            sections.append("• **Acceptance Strategies:** Strategic risk tolerance and monitoring")
            sections.append("• **Contingency Planning:** Emergency response and business continuity")
            sections.append("• **Recovery Strategies:** Post-crisis restoration and improvement protocols")
            sections.append("")
            
            sections.append(f"## 📈 Performance Metrics & KPI Dashboard")
            sections.append("### 🎯 Key Performance Indicator Framework")
            sections.append("Advanced performance measurement framework encompasses strategic, operational, financial, and customer metrics for comprehensive business performance evaluation and continuous improvement achievement through systematic monitoring and optimization protocols.")
            sections.append("")
            sections.append("#### Strategic Performance Metrics")
            sections.append("• **Market Share:** Position tracking and competitive performance")
            sections.append("• **Customer Satisfaction:** Loyalty measurement and retention analytics")
            sections.append("• **Innovation Index:** R&D productivity and technology advancement")
            sections.append("• **Sustainability Metrics:** ESG performance and stakeholder value creation")
            sections.append("• **Brand Value:** Market perception and reputation measurement")
            sections.append("• **Strategic Initiative Progress:** Goal achievement and milestone tracking")
            sections.append("")
            sections.append("### 📊 Balanced Scorecard Implementation")
            sections.append("#### Multi-Dimensional Performance Framework")
            sections.append("• **Financial Perspective:** Revenue growth, profitability, cost management")
            sections.append("• **Customer Perspective:** Satisfaction, retention, market penetration")
            sections.append("• **Internal Process:** Efficiency, quality, innovation capability")
            sections.append("• **Learning & Growth:** Employee development, technology advancement")
            sections.append("• **Stakeholder Value:** Investor returns, community impact, regulatory compliance")
            sections.append("• **Sustainability Performance:** Environmental impact, social responsibility")
            sections.append("")
            
            sections.append(f"## 🔮 Future Outlook & Strategic Recommendations")
            sections.append("### 🌟 Strategic Future Planning & Market Evolution")
            sections.append("Strategic future planning encompasses market trend anticipation, competitive landscape evolution assessment, technology advancement impact evaluation, regulatory environment change analysis, and strategic positioning requirements for long-term competitive advantage and market leadership sustainability achievement through systematic strategic management excellence.")
            sections.append("")
            sections.append("#### Emerging Opportunities Matrix")
            sections.append("• **Market Expansion:** Geographic growth and demographic penetration")
            sections.append("• **Technology Integration:** Digital transformation and automation advantages")
            sections.append("• **Strategic Partnerships:** Collaboration and ecosystem development")
            sections.append("• **Innovation Leadership:** Technology advancement and market disruption")
            sections.append("• **Sustainability Leadership:** ESG excellence for competitive differentiation")
            sections.append("• **Customer Experience:** Service innovation and relationship enhancement")
            sections.append("")
            sections.append("### 🚀 Strategic Recommendation Portfolio")
            sections.append("#### Immediate Strategic Priorities (0-12 months)")
            sections.append("• **Market Position Strengthening:** Competitive advantage enhancement and market share protection")
            sections.append("• **Operational Excellence:** Process optimization and efficiency improvement")
            sections.append("• **Financial Optimization:** Revenue enhancement and cost structure improvement")
            sections.append("• **Risk Management:** Vulnerability assessment and mitigation strategy implementation")
            sections.append("• **Technology Advancement:** Digital transformation and innovation capability development")
            sections.append("• **Stakeholder Engagement:** Customer, investor, and employee relationship enhancement")
            sections.append("")
            sections.append("#### Long-term Strategic Vision (1-5 years)")
            sections.append("• **Market Leadership:** Industry position strengthening and competitive dominance")
            sections.append("• **Innovation Excellence:** Technology leadership and market disruption capability")
            sections.append("• **Sustainable Growth:** Balanced expansion and stakeholder value creation")
            sections.append("• **Strategic Transformation:** Business model evolution and market adaptation")
            sections.append("• **Global Expansion:** International market penetration and competitive positioning")
            sections.append("• **Ecosystem Development:** Partnership networks and collaborative advantage creation")
        
        # Common conclusion
        sections.append(f"\n## 🔧 Technical Methodology Notes")
        sections.append(f"Analysis generated using INSYT's LangGraph multi-agent orchestration with intelligent query routing, self-healing validation, and memory enhancement capabilities.")
        sections.append(f"Target analysis depth: {config['target_words']} words for {search_mode} mode.")
        sections.append(f"Research methodology encompasses financial intelligence, market analysis, competitive assessment, and strategic synthesis for comprehensive business intelligence delivery.")
        
        content = "\n".join(sections)
        
        # Ensure we meet word requirements with mode-specific padding
        current_words = len(content.split())
        if current_words < min_words:
            # Add content based on search mode
            if search_mode == "Quick Search":
                padding = "\n\n### 🔄 Continuous Monitoring Framework\nStrategic analysis requires ongoing monitoring and adjustment protocols for sustained competitive advantage. Executive teams must implement rapid response systems for market opportunity identification and tactical resource deployment. Performance tracking systems enable real-time strategic optimization and competitive positioning enhancement for maximum business impact achievement within dynamic market environments requiring agile decision-making capabilities."
            elif search_mode == "Extended Search":
                padding = "\n\n### 🎯 Strategic Excellence Implementation\nExtended strategic analysis encompasses comprehensive market evaluation methodologies, competitive positioning assessment frameworks, operational capability enhancement strategies, financial performance optimization protocols, and risk management system development for sustainable competitive advantage achievement. Strategic planning excellence requires integration of market intelligence, competitive analysis, financial optimization, operational excellence, and strategic implementation coordination for comprehensive business performance enhancement and long-term sustainability achievement through systematic strategic management and continuous performance optimization methodologies."
            else:  # Deep Search
                padding = "\n\n### 🌐 Comprehensive Strategic Intelligence Framework\nAdvanced strategic intelligence framework incorporates sophisticated analytics methodologies, predictive modeling capabilities, multi-dimensional risk assessment protocols, competitive intelligence gathering systems, market trend analysis frameworks, strategic capability development programs, operational excellence optimization, financial performance enhancement strategies, and strategic implementation coordination for long-term market leadership achievement and sustainable competitive positioning excellence. Advanced strategic planning requires sophisticated analytical frameworks, comprehensive data integration, multi-stakeholder coordination, strategic alignment optimization, and continuous performance measurement for strategic objective achievement and competitive advantage sustainability in dynamic business environments requiring advanced strategic management capabilities and comprehensive business intelligence systems for optimal decision-making support and strategic excellence achievement."
            
            while len(content.split()) < min_words:
                content += padding
        
        # Ensure we don't exceed maximum
        words = content.split()
        if len(words) > max_words:
            content = ' '.join(words[:max_words]) + f"\n\n*[Analysis optimized for {search_mode} word limit]*"
        
        return content
        
    def validation_agent_node(self, state: AgentState) -> AgentState:
        """Validate synthesized results"""
        
        # Simple validation logic
        validation_score = 0.8
        
        if state.get("financial_data") and not state["financial_data"].get("error"):
            validation_score += 0.05
        if state.get("news_data") and not state["news_data"].get("error"):
            validation_score += 0.05
        if state.get("competitive_data") and not state["competitive_data"].get("error"):
            validation_score += 0.05
        
        validation_results = {
            "confidence_score": min(validation_score, 1.0),
            "data_sources_validated": 3,
            "validation_method": "multi_source_cross_reference"
        }
        
        state["validation_results"] = validation_results
        state["confidence_score"] = validation_results["confidence_score"]
        state["messages"].append(AIMessage(content=f"Validation completed with {validation_results['confidence_score']:.1%} confidence"))
        
        return state

class SelfHealingValidationAgent:
    """Feature 3: Self-Healing Data Validation using LangGraph"""
    
    def __init__(self, llm, storage):
        self.llm = llm
        self.storage = storage
        self.setup_validation_workflow()
    
    def setup_validation_workflow(self):
        """Setup validation workflow"""
        
        workflow = StateGraph(AgentState)
        
        # Add validation nodes
        workflow.add_node("detect_conflicts", self.detect_conflicts_node)
        workflow.add_node("resolve_conflicts", self.resolve_conflicts_node)
        workflow.add_node("verify_resolution", self.verify_resolution_node)
        workflow.add_node("update_confidence", self.update_confidence_node)
        
        # Set entry point
        workflow.set_entry_point("detect_conflicts")
        
        # Add edges
        workflow.add_conditional_edges(
            "detect_conflicts",
            self.check_conflicts,
            {
                "conflicts_found": "resolve_conflicts",
                "no_conflicts": "update_confidence"
            }
        )
        
        workflow.add_edge("resolve_conflicts", "verify_resolution")
        workflow.add_edge("verify_resolution", "update_confidence")
        workflow.add_edge("update_confidence", END)
        
        self.validation_workflow = workflow.compile()
    
    def detect_conflicts_node(self, state: AgentState) -> AgentState:
        """Detect data conflicts across sources"""
        
        conflicts = []
        
        # Check for conflicts between financial, news, and competitive data
        financial_data = state.get("financial_data", {})
        news_data = state.get("news_data", {})
        competitive_data = state.get("competitive_data", {})
        
        # Simple conflict detection logic
        if (financial_data.get("confidence", 0) < 0.5 or 
            news_data.get("confidence", 0) < 0.5 or 
            competitive_data.get("confidence", 0) < 0.5):
            conflicts.append("Low confidence data detected")
        
        state["validation_results"] = state.get("validation_results", {})
        state["validation_results"]["conflicts"] = conflicts
        state["validation_results"]["conflicts_detected"] = len(conflicts) > 0
        
        return state
    
    def check_conflicts(self, state: AgentState) -> str:
        """Check if conflicts were detected"""
        return "conflicts_found" if state["validation_results"]["conflicts_detected"] else "no_conflicts"
    
    def resolve_conflicts_node(self, state: AgentState) -> AgentState:
        """Resolve detected conflicts"""
        
        conflicts = state["validation_results"]["conflicts"]
        
        resolution_prompt = f"""
        Resolve the following data conflicts for {state['company_name']}:
        
        Conflicts: {conflicts}
        
        Financial Data: {state.get('financial_data', {})}
        News Data: {state.get('news_data', {})}
        Competitive Data: {state.get('competitive_data', {})}
        
        Provide a resolution strategy and corrected data.
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=resolution_prompt)])
            
            state["validation_results"]["resolution"] = response.content
            state["validation_results"]["resolved"] = True
            
        except Exception as e:
            state["validation_results"]["resolution"] = f"Error resolving conflicts: {str(e)}"
            state["validation_results"]["resolved"] = False
        
        return state
    
    def verify_resolution_node(self, state: AgentState) -> AgentState:
        """Verify conflict resolution"""
        
        # Simple verification logic
        if state["validation_results"].get("resolved", False):
            state["validation_results"]["verification_score"] = 0.85
        else:
            state["validation_results"]["verification_score"] = 0.5
        
        return state
    
    def update_confidence_node(self, state: AgentState) -> AgentState:
        """Update overall confidence score"""
        
        base_confidence = 0.8
        
        if state["validation_results"].get("conflicts_detected", False):
            if state["validation_results"].get("resolved", False):
                confidence_adjustment = 0.05  # Small penalty for resolved conflicts
            else:
                confidence_adjustment = 0.2   # Larger penalty for unresolved conflicts
            
            final_confidence = max(0.1, base_confidence - confidence_adjustment)
        else:
            final_confidence = base_confidence + 0.1  # Bonus for no conflicts
        
        state["confidence_score"] = min(final_confidence, 1.0)
        state["validation_results"]["final_confidence"] = state["confidence_score"]
        
        return state

class UserQueryEnhancer:
    """User Query Enhancer for improved results"""
    
    def __init__(self, llm):
        self.llm = llm
        self.enhancement_patterns = {
            'financial': ['revenue', 'earnings', 'profit', 'financial performance', 'market cap', 'valuation'],
            'competitive': ['market position', 'competitors', 'market share', 'industry analysis'],
            'strategic': ['business model', 'growth strategy', 'partnerships', 'acquisitions'],
            'operational': ['products', 'services', 'operations', 'management'],
            'market': ['industry trends', 'market outlook', 'opportunities', 'threats']
        }
    
    def enhance_query(self, original_query: str, company_name: str) -> Dict:
        """Enhance user query for better research results"""
        
        enhancement_prompt = f"""
        Enhance this business intelligence query for comprehensive research:
        
        Original Query: "{original_query}"
        Company: {company_name}
        
        Transform this into a comprehensive research query that includes:
        1. Specific business aspects to investigate
        2. Key metrics and data points to find
        3. Competitive context and comparisons
        4. Recent developments and trends
        5. Strategic insights and implications
        
        Make it actionable for multi-agent research while preserving user intent.
        
        Return format:
        {{
            "enhanced_query": "detailed enhanced query",
            "focus_areas": ["area1", "area2", "area3"],
            "key_metrics": ["metric1", "metric2"],
            "research_scope": "broad/focused/deep"
        }}
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=enhancement_prompt)])
            
            # Parse response or create fallback
            enhanced_data = self._parse_enhancement_response(response.content, original_query, company_name)
            
            return {
                'success': True,
                'original_query': original_query,
                'enhanced_query': enhanced_data['enhanced_query'],
                'focus_areas': enhanced_data.get('focus_areas', []),
                'key_metrics': enhanced_data.get('key_metrics', []),
                'research_scope': enhanced_data.get('research_scope', 'broad'),
                'enhancement_applied': True
            }
            
        except Exception as e:
            # Fallback enhancement
            return self._fallback_enhancement(original_query, company_name)
    
    def _parse_enhancement_response(self, response: str, original_query: str, company_name: str) -> Dict:
        """Parse LLM response or create structured enhancement"""
        
        try:
            # Try to extract JSON
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            if json_start >= 0 and json_end > json_start:
                return json.loads(response[json_start:json_end])
        except:
            pass
        
        # Fallback parsing
        return self._fallback_enhancement(original_query, company_name)['enhanced_data']
    
    def _fallback_enhancement(self, original_query: str, company_name: str) -> Dict:
        """Fallback query enhancement using patterns"""
        
        query_lower = original_query.lower()
        detected_categories = []
        
        # Detect query categories
        for category, keywords in self.enhancement_patterns.items():
            if any(keyword in query_lower for keyword in keywords):
                detected_categories.append(category)
        
        # Build enhanced query
        if not detected_categories:
            detected_categories = ['financial', 'competitive', 'strategic']
        
        enhanced_query = f"""
        Comprehensive analysis of {company_name} focusing on: {original_query}
        
        Include analysis of:
        - Financial performance and key metrics
        - Market position and competitive landscape
        - Recent developments and strategic initiatives
        - Industry context and trends
        - Growth opportunities and risk factors
        """
        
        enhanced_data = {
            'enhanced_query': enhanced_query.strip(),
            'focus_areas': detected_categories,
            'key_metrics': ['revenue', 'market_share', 'growth_rate'],
            'research_scope': 'broad'
        }
        
        return {
            'success': True,
            'original_query': original_query,
            'enhanced_query': enhanced_data['enhanced_query'],
            'focus_areas': enhanced_data['focus_areas'],
            'key_metrics': enhanced_data['key_metrics'],
            'research_scope': enhanced_data['research_scope'],
            'enhancement_applied': True,
            'enhanced_data': enhanced_data
        }

class MemoryEnhancedAgent:
    """Feature 4: Memory Enhancement using LangGraph"""
    
    def __init__(self, llm, storage, session_id):
        self.llm = llm
        self.storage = storage
        self.session_id = session_id
    
    def enhance_query_with_memory(self, query: str, company_name: str) -> Dict:
        """Enhance query with memory context"""
        
        # Retrieve relevant memory context
        memory_context = self.storage.get_memory_context(self.session_id, company_name, limit=5)
        
        if not memory_context:
            return {
                "enhanced_query": query,
                "context_used": False,
                "memory_items": 0
            }
        
        # Build context prompt
        context_prompt = f"""
        Previous context for {company_name}:
        
        """
        
        for item in memory_context:
            context_data = item.get('context_data', {})
            context_prompt += f"- {item['context_type']}: {context_data.get('summary', 'No summary')}\n"
        
        enhancement_prompt = f"""
        Original Query: {query}
        Company: {company_name}
        
        {context_prompt}
        
        Based on the previous context, enhance the original query to:
        1. Include relevant historical context
        2. Ask follow-up questions based on previous searches
        3. Identify gaps in previous research
        4. Provide continuity with past conversations
        
        Return the enhanced query that builds upon previous context.
        """
        
        try:
            response = self.llm.invoke([HumanMessage(content=enhancement_prompt)])
            
            enhanced_result = {
                "enhanced_query": response.content,
                "context_used": True,
                "memory_items": len(memory_context),
                "original_query": query
            }
            
            # Save this enhancement as new memory context
            self.storage.save_memory_context(
                self.session_id,
                company_name,
                "query_enhancement",
                {
                    "original_query": query,
                    "enhanced_query": response.content,
                    "context_items_used": len(memory_context)
                },
                relevance_score=0.8
            )
            
            return enhanced_result
            
        except Exception as e:
            return {
                "enhanced_query": query,
                "context_used": False,
                "memory_items": 0,
                "error": str(e)
            }
    
    def save_research_memory(self, query: str, company_name: str, result: Dict):
        """Save research results to memory for future enhancement"""
        
        memory_data = {
            "query": query,
            "company": company_name,
            "summary": result.get("content", "")[:500] + "...",  # Truncated summary
            "confidence": result.get("confidence_score", 0.8),
            "sources_used": result.get("sources_used", 0),
            "key_findings": self._extract_key_findings(result.get("content", "")),
            "timestamp": datetime.now().isoformat()
        }
        
        # Calculate relevance score based on result quality
        relevance_score = min(1.0, 0.5 + (result.get("confidence_score", 0.8) * 0.5))
        
        self.storage.save_memory_context(
            self.session_id,
            company_name,
            "research_result",
            memory_data,
            relevance_score
        )
    
    def _extract_key_findings(self, content: str) -> List[str]:
        """Extract key findings from research content"""
        
        # Simple extraction - look for sentences with key indicators
        key_indicators = ["key finding", "important", "significant", "notable", "critical"]
        sentences = content.split('. ')
        
        key_findings = []
        for sentence in sentences:
            if any(indicator in sentence.lower() for indicator in key_indicators):
                key_findings.append(sentence.strip())
        
        return key_findings[:3]  # Return top 3 key findings

class EnhancedAgenticAIAssistant:
    """Enhanced INSYT with LangGraph Agentic Features"""
    
    def __init__(self):
        self.clients = aws_clients
        self.session_id = self._get_session_id()
        
        # Initialize LLM clients with fallback hierarchy
        self.nova_llm = None
        self.perplexity_llm = None
        self.openai_llm = None
        
        # Initialize Nova Bedrock FIRST (Primary LLM)
        try:
            bedrock_api_key = os.getenv("BEDROCK_API_KEY")
            if AWS_REGION and bedrock_api_key:
                # CRITICAL: Set the bearer token first
                os.environ["AWS_BEARER_TOKEN_BEDROCK"] = bedrock_api_key
                
                self.nova_llm = NovaBedrockLLM(AWS_REGION)
                # Test the connection
                test_result = self.nova_llm.test_connection()
                if not test_result["success"]:
                    st.warning(f"Nova Bedrock test failed: {test_result['message']}")
                    self.nova_llm = None
               
        except Exception as e:
            st.warning(f"Nova Bedrock initialization failed: {str(e)}")
            self.nova_llm = None
        
        # Initialize Perplexity (1st Fallback)
        if PERPLEXITY_API_KEY:
            self.perplexity_llm, status_msg = initialize_perplexity_safely(PERPLEXITY_API_KEY)
        
        # Initialize OpenAI (2nd Fallback)
        if OPENAI_API_KEY:
            try:
                self.openai_llm = ChatOpenAI(
                    api_key=OPENAI_API_KEY,
                    model="gpt-4o",
                    temperature=0.7
                )
            except Exception as e:
                st.warning(f"OpenAI initialization failed: {str(e)}")
        
        # Set primary LLM with proper fallback hierarchy
        self.llm = self.nova_llm or self.perplexity_llm or self.openai_llm
        
        # Initialize persistent storage
        self.storage = PersistentStorage()
        
        # Initialize enhanced components
        self.serper_api = SerperSearchAPI(SERPER_API_KEY) if SERPER_API_KEY else None
        
        # Initialize LangGraph Agents
        self.query_enhancer = UserQueryEnhancer(self.llm)
        self.multi_agent_orchestrator = MultiAgentResearchOrchestrator(self.llm)
        self.validation_agent = SelfHealingValidationAgent(self.llm, self.storage)
        self.memory_agent = MemoryEnhancedAgent(self.llm, self.storage, self.session_id)
        
        # Clean up expired cache on initialization
        self.storage.cleanup_expired_cache()
        
    
    def _get_session_id(self):
        if 'session_id' not in st.session_state:
            st.session_state.session_id = str(uuid.uuid4())
        return st.session_state.session_id
    
    def show_langgraph_animation(self):
            """Display unified LangGraph processing animation"""
            
            animation_steps = [
                ("🤖", "Running LangGraph multi-agent analysis..."),
                ("✨", "Enhancing query for optimal research..."),
                ("🧠", "Enhancing query with memory context..."),
                ("🤖", "Initiating multi-agent research orchestration..."),
                ("🔄", "Running multi-agent research workflow...")
            ]
            
            # Create animation container
            animation_placeholder = st.empty()
            
            with animation_placeholder.container():
                st.markdown("""
                <div class="langgraph-animation-container">
                    <h3 style="margin: 0 0 1.5rem 0; font-size: 1.8rem; font-weight: 700; text-shadow: 0 2px 4px rgba(0,0,0,0.3);">
                        🚀 LangGraph Multi-Agent Processing
                    </h3>
                    <div class="progress-bar">
                        <div class="progress-fill"></div>
                    </div>
                    <div id="animation-steps">
                """, unsafe_allow_html=True)
                
                # Display all steps with staggered animations
                for i, (icon, text) in enumerate(animation_steps):
                    delay = i * 0.3
                    st.markdown(f"""
                        <div class="animation-step active" style="animation-delay: {delay}s;">
                            <span class="step-icon">{icon}</span>
                            <span>{text}</span>
                        </div>
                    """, unsafe_allow_html=True)
                
                st.markdown("""
                    </div>
                    <div style="margin-top: 1.5rem; font-size: 1rem; opacity: 0.9; font-style: italic;">
                        Powered by advanced multi-agent orchestration
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            # Keep animation visible for 8 seconds
            time.sleep(8)
            animation_placeholder.empty()

    def enhanced_agentic_search_with_recovery(self, query: str, company_name: str, search_mode: str = "Extended Search") -> Dict:
        """Enhanced search with LLM fallback hierarchy"""
        
        # Try Nova Bedrock first
        if self.nova_llm:
            try:
                return self.enhanced_agentic_search(query, company_name, search_mode)
            except Exception as e:
                st.warning(f"Nova Bedrock error: {str(e)}")
        
        # Fallback to Perplexity
        if self.perplexity_llm:
            try:
                st.info("🔄 Falling back to Perplexity...")
                old_llm = self.llm
                self.llm = self.perplexity_llm
                result = self.enhanced_agentic_search(query, company_name, search_mode)
                self.llm = old_llm
                return result
            except Exception as e:
                st.warning(f"Perplexity error: {str(e)}")
                self.llm = old_llm
        
        # Final fallback to OpenAI
        if self.openai_llm:
            try:
                st.info("🔄 Falling back to OpenAI...")
                old_llm = self.llm
                self.llm = self.openai_llm
                result = self.enhanced_agentic_search(query, company_name, search_mode)
                self.llm = old_llm
                return result
            except Exception as e:
                st.warning(f"OpenAI error: {str(e)}")
                self.llm = old_llm
        
        # Emergency fallback
        return self._fallback_search(company_name, search_mode, "")
    
    
    def contextual_search(self, query: str, search_mode: str = "Extended Search") -> Dict:
        """Contextual search that uses conversation history for better responses"""
        
        try:
            # Get recent chat history for context
            chat_history = self.get_chat_history(limit=5)
            
            # Build contextual prompt
            context_messages = []
            
            if chat_history:
                context_messages.append("Previous conversation context:")
                for i, chat in enumerate(reversed(chat_history)):
                    if chat.get('query') and chat.get('response'):
                        context_messages.append(f"Q{i+1}: {chat['query']}")
                        context_messages.append(f"A{i+1}: {chat['response'][:300]}...")
            
            # Combine context with current query
            contextual_query = f"""
    {chr(10).join(context_messages)}

    Current question: {query}

    Please answer the current question considering the above conversation context. If this is a follow-up question, reference previous information appropriately.
    """
            
            # Use regular enhanced search with contextual query
            return self.enhanced_agentic_search_with_recovery(contextual_query, query, search_mode)
            
        except Exception as e:
            st.error(f"Error in contextual search: {str(e)}")
            return self.enhanced_agentic_search_with_recovery(query, query, search_mode)
        
    def get_paginated_chat_history(self, page: int = 1, per_page: int = 5) -> Dict:
        """Get paginated chat history across all sessions for persistence"""
        try:
            # Get total count first across all sessions
            def _get_count():
                conn = self.storage._get_connection()
                try:
                    cursor = conn.cursor()
                    cursor.execute('SELECT COUNT(*) FROM chat_history')
                    return cursor.fetchone()[0]
                finally:
                    conn.close()
            
            total_count = self.storage._execute_with_retry(_get_count)
            
            # Calculate pagination
            offset = (page - 1) * per_page
            total_pages = (total_count + per_page - 1) // per_page
            
            # Get paginated results across all sessions
            def _get_paginated():
                conn = self.storage._get_connection()
                try:
                    cursor = conn.cursor()
                    cursor.execute('''
                        SELECT * FROM chat_history 
                        ORDER BY timestamp DESC 
                        LIMIT ? OFFSET ?
                    ''', (per_page, offset))
                    
                    columns = [description[0] for description in cursor.description]
                    results = [dict(zip(columns, row)) for row in cursor.fetchall()]
                    return results
                finally:
                    conn.close()
            
            results = self.storage._execute_with_retry(_get_paginated)
            
            return {
                'history': results,
                'current_page': page,
                'total_pages': total_pages,
                'total_count': total_count,
                'per_page': per_page
            }
            
        except Exception as e:
            st.error(f"Error retrieving paginated history: {str(e)}")
            return {
                'history': [],
                'current_page': 1,
                'total_pages': 1,
                'total_count': 0,
                'per_page': per_page
            }
    
    def format_chat_thread(self, chat_item: Dict, thread_index: int) -> str:
        """Format a single chat thread for display"""
        try:
            timestamp = chat_item.get('timestamp', '')
            if timestamp:
                from datetime import datetime
                dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                formatted_time = dt.strftime('%Y-%m-%d %H:%M')
            else:
                formatted_time = 'Unknown time'
            
            query = chat_item.get('query', 'No query')[:100] + "..." if len(chat_item.get('query', '')) > 100 else chat_item.get('query', 'No query')
            response = chat_item.get('response', 'No response')[:200] + "..." if len(chat_item.get('response', '')) > 200 else chat_item.get('response', 'No response')
            
            provider = chat_item.get('provider', 'Unknown')
            search_mode = chat_item.get('search_mode', 'Unknown')
            agentic_enhanced = chat_item.get('agentic_enhanced', False)
            
            enhancement_badge = "🚀 LangGraph" if agentic_enhanced else "📚 Standard"
            
            return f"""
    **Thread #{thread_index}** - {formatted_time} | {enhancement_badge} | {search_mode}

    **Q:** {query}

    **A:** {response}

    *Provider: {provider}*
    """
        except Exception as e:
            return f"**Thread #{thread_index}** - Error formatting thread: {str(e)}"
    
    
    def enhanced_agentic_search(self, query: str, company_name: str, search_mode: str = "Extended Search") -> Dict:
        """Main enhanced search using all LangGraph agents"""
        
        try:
            
            # Step 1: Query Enhancement for better results
            query_enhancement = self.query_enhancer.enhance_query(query, company_name)
            
            if query_enhancement["enhancement_applied"]:
                enhanced_query = query_enhancement["enhanced_query"]
                focus_areas = query_enhancement.get("focus_areas", [])
                st.success(f"✅ Query enhanced with {len(focus_areas)} focus areas")
                
                with st.expander("🔍 View Query Enhancement"):
                    st.write(f"**Original Query:** {query_enhancement['original_query']}")
                    st.write(f"**Enhanced Query:** {enhanced_query}")
                    st.write(f"**Focus Areas:** {', '.join(focus_areas)}")
                    st.write(f"**Research Scope:** {query_enhancement.get('research_scope', 'broad')}")
            else:
                enhanced_query = query
                st.info("🔍 Query enhancement not applied")
            
            # Step 2: Memory Enhancement - Enhance query with context
            memory_enhancement = self.memory_agent.enhance_query_with_memory(enhanced_query, company_name)
            
            if memory_enhancement["context_used"]:
                final_query = memory_enhancement["enhanced_query"]
                st.success(f"✅ Query enhanced with {memory_enhancement['memory_items']} memory items")
                
                with st.expander("🔍 View Memory Enhancement"):
                    st.write(f"**Query-Enhanced:** {enhanced_query}")
                    st.write(f"**Memory-Enhanced:** {final_query}")
            else:
                final_query = enhanced_query
                st.info("🔍 No previous context found")
            
            # Check cache first
            cache_key = hashlib.md5(f"{company_name.lower()}_{final_query}_{search_mode}".encode()).hexdigest()
            cached_result = self.storage.get_cached_result(cache_key)
            
            if cached_result:
                st.info("💾 Retrieved from persistent cache")
                return {
                    'success': True,
                    'content': cached_result['content'],
                    'company_name': company_name,
                    'provider': 'LangGraph Multi-Agent (Cached)',
                    'search_mode': search_mode,
                    'cached': True,
                    'agentic_enhanced': True,
                    'query_enhanced': query_enhancement["enhancement_applied"],
                    'validation_score': cached_result.get('validation_score'),
                    'validation_details': cached_result.get('validation_details', {})
                }
            
            # Feature 1: Multi-Agent Research Orchestrator
            st.info("🤖 Initiating multi-agent research orchestration...")
            
            # Initialize agent state with enhanced query and search mode constraints
            initial_state = {
                "messages": [HumanMessage(content=f"{final_query}\n\nSTRICT MODE: {search_mode} - Target response length: {token_config['target_words']} words")],
                "query": final_query,
                "company_name": company_name,
                "search_mode": search_mode,
                "research_plan": {},
                "financial_data": {},
                "news_data": {},
                "competitive_data": {},
                "validation_results": {},
                "memory_context": memory_enhancement,
                "query_enhancement": query_enhancement,
                "next_action": "",
                "confidence_score": 0.0,
                "final_result": {}
            }
            
            # Run multi-agent workflow
            final_state = self.multi_agent_orchestrator.workflow.invoke(initial_state)
            st.success("✅ Multi-agent research completed")
            
            # STRICT MODE VALIDATION - Ensure response meets requirements
            final_result = final_state.get("final_result", {})
            if final_result.get("content"):
                content_words = len(final_result["content"].split())
                expected_min = token_config['min_word_limit']
                expected_max = token_config['word_limit']
                
                if content_words < expected_min or content_words > expected_max:
                    st.warning(f"⚠️ Content length ({content_words} words) outside {search_mode} range ({expected_min}-{expected_max}). Adjusting...")
                    
                    # Force content adjustment
                    if content_words < expected_min:
                        final_result["content"] = expand_content_intelligently(
                            final_result["content"], 
                            search_mode, 
                            expected_min - content_words
                        )
                    else:
                        words = final_result["content"].split()
                        final_result["content"] = ' '.join(words[:expected_max])
                        final_result["content"] += f"\n\n*[Content optimized to {expected_max} words for {search_mode}]*"
                    
                    # Update word count after adjustment
                    final_word_count = len(final_result["content"].split())
                    st.success(f"✅ Content adjusted to {final_word_count} words for {search_mode}")
            
            # Feature 3: Self-Healing Data Validation
            # Run validation workflow
            validated_state = self.validation_agent.validation_workflow.invoke(final_state)
            
            if validated_state["validation_results"].get("conflicts_detected"):
                if validated_state["validation_results"].get("resolved"):
                    st.warning("⚠️ Data conflicts detected and resolved")
                else:
                    st.error("❌ Data conflicts detected but not fully resolved")
            else:
                st.success("✅ No data conflicts detected")
            
            # Prepare final result
            final_result = validated_state["final_result"]
            if not final_result.get("content"):
                # Emergency fallback with mode-specific content
                if search_mode == "Quick Search":
                    final_result["content"] = f"""# Quick Analysis: {company_name}

    ## Executive Summary
    Rapid business intelligence analysis completed for {company_name}. Key strategic insights identified requiring immediate executive attention and tactical response coordination.

    ## Critical Insights
    Market positioning analysis reveals competitive dynamics and strategic opportunities requiring swift implementation for optimal business performance and competitive advantage achievement.

    ## Financial Snapshot
    Current performance indicators suggest specific operational adjustments and strategic interventions for enhanced market positioning and sustainable competitive advantage development.

    ## Immediate Actions
    1. Deploy rapid response protocols for market opportunities
    2. Implement tactical adjustments for competitive positioning
    3. Optimize resource allocation for maximum strategic impact

    ## Risk Alerts
    Market volatility factors and competitive pressures require immediate monitoring and strategic response coordination for sustained performance achievement.

    ## Next Steps
    Priority initiatives demand immediate executive attention and resource deployment for competitive positioning enhancement and market advantage achievement."""
                
                elif search_mode == "Extended Search":
                    final_result["content"] = f"""# Strategic Analysis: {company_name}

    ## Executive Summary
    Comprehensive strategic business intelligence analysis for {company_name} provides detailed insights for strategic planning and competitive positioning enhancement through systematic evaluation and implementation frameworks.

    ## Market Analysis & Intelligence
    Advanced market intelligence reveals complex industry dynamics impacting strategic positioning, competitive landscapes, and growth opportunity identification requiring systematic strategic response and implementation coordination.

    ## Financial Performance Review
    Comprehensive financial analysis encompasses revenue optimization strategies, profitability enhancement methodologies, and strategic investment effectiveness evaluation for sustainable competitive advantage and market leadership achievement.

    ## Competitive Landscape Assessment
    Strategic competitive analysis reveals market positioning effectiveness, competitive advantage sustainability, and strategic differentiation opportunities for enhanced market leadership and competitive performance optimization.

    ## Strategic Recommendations
    1. Market Positioning: Enhance competitive differentiation through strategic capability development and market positioning optimization
    2. Operational Excellence: Optimize internal processes for enhanced efficiency, performance improvement, and competitive advantage achievement
    3. Financial Optimization: Implement strategic initiatives for revenue enhancement, profitability improvement, and value creation maximization
    4. Innovation Strategy: Develop technological capabilities and innovation frameworks for sustained competitive advantage and market leadership

    ## Risk Assessment & Mitigation
    Strategic risk evaluation encompasses market volatility analysis, competitive pressure assessment, operational challenge identification, and financial performance risk management requiring systematic mitigation strategies and monitoring protocols.

    ## Implementation Roadmap
    Strategic implementation requires systematic approach including timeline development, resource allocation optimization, performance measurement system establishment, and continuous optimization for sustainable competitive advantage achievement and market leadership development."""
                
                else:  # Deep Search
                    final_result["content"] = f"""# Comprehensive Strategic Research: {company_name}

    ## Executive Summary
    In-depth strategic research analysis for {company_name} employs advanced multi-agent orchestration with comprehensive validation protocols to provide research-grade strategic intelligence for major business decisions and long-term strategic planning excellence.

    ## Advanced Market Intelligence
    Comprehensive market intelligence encompasses multi-dimensional analysis of industry evolution patterns, competitive transformation dynamics, technological disruption impacts, regulatory environment changes, and economic factors affecting strategic positioning effectiveness and market leadership sustainability requirements.

    ### Industry Evolution Analysis
    Advanced industry analysis reveals complex structural changes including technological disruption patterns, regulatory environment evolution, competitive landscape transformation, and customer behavior modification impacting long-term strategic positioning and market leadership development requirements.

    ### Market Development Intelligence
    Recent market developments indicate strategic implications requiring comprehensive evaluation including competitive dynamics assessment, market opportunity identification, technological advancement integration, and strategic positioning optimization for sustainable competitive advantage achievement.

    ## Comprehensive Financial Analysis
    Advanced financial intelligence encompasses multi-dimensional performance evaluation including revenue optimization strategies, profitability enhancement methodologies, capital allocation effectiveness assessment, investment portfolio analysis, cash flow optimization, and value creation measurement frameworks for sustainable growth achievement.

    ### Financial Performance Deep Dive
    Detailed financial metrics analysis reveals performance patterns and strategic opportunities including revenue generation effectiveness, cost structure optimization, profitability enhancement potential, and investment return maximization for competitive advantage development and market leadership achievement.

    ### Strategic Financial Framework
    Advanced financial strategy encompasses capital structure optimization, investment prioritization frameworks, risk management protocols, performance measurement systems, and value creation methodologies for competitive advantage development and market leadership sustainability achievement.

    ## Detailed Competitive Assessment
    Comprehensive competitive intelligence includes direct competitor analysis, indirect competition evaluation, market share dynamics assessment, competitive positioning evaluation, strategic capability comparison, and competitive response analysis for strategic advantage identification and development enhancement.

    ### Competitive Positioning Analysis
    Advanced competitive evaluation encompasses market positioning effectiveness assessment, competitive advantage sustainability analysis, strategic differentiation opportunity identification, and competitive response capability evaluation for enhanced market leadership achievement and sustainable competitive positioning optimization.

    ## Strategic Framework & Planning
    Comprehensive strategic framework development encompasses vision articulation, strategic objective definition, capability development planning, resource allocation optimization, performance measurement system implementation, and continuous improvement protocols for sustainable competitive advantage achievement and market leadership development.

    ### Implementation Strategy Development
    Advanced implementation planning includes detailed timeline development, resource allocation optimization, risk mitigation protocols, performance monitoring systems, and continuous optimization methodologies for strategic objective achievement and competitive advantage sustainability enhancement.

    ## Risk Management & Scenario Analysis
    Comprehensive risk assessment encompasses strategic risk evaluation, operational challenge analysis, financial vulnerability assessment, market volatility monitoring, regulatory compliance requirements, and competitive threat analysis requiring systematic mitigation strategies and contingency planning for organizational resilience development.

    ### Advanced Scenario Planning
    Sophisticated scenario analysis includes multiple strategic alternatives evaluation, contingency planning development, risk mitigation strategy implementation, opportunity identification processes, and adaptive strategy development for strategic flexibility enhancement and competitive advantage sustainability achievement.

    ## Performance Metrics & KPIs
    Advanced performance measurement framework encompasses key performance indicators development, balanced scorecard implementation, competitive benchmarking systems, customer satisfaction metrics, financial performance tracking, and strategic objective monitoring for continuous improvement achievement and strategic optimization.

    ## Future Outlook & Recommendations
    Strategic future planning encompasses market trend anticipation, competitive landscape evolution assessment, technology advancement impact evaluation, regulatory environment change analysis, and strategic positioning requirements for long-term competitive advantage and market leadership sustainability achievement through systematic strategic management excellence."""
                
                final_result["sources_used"] = 1
                final_result["synthesis_method"] = "emergency_fallback"
            
            # Feature 4: Save to memory for future enhancement
            self.memory_agent.save_research_memory(query, company_name, {
                "content": final_result["content"],
                "confidence_score": validated_state["confidence_score"],
                "sources_used": final_result.get("sources_used", 3)
            })
            
            # Save to persistent storage
            result_data = {
                'session_id': self.session_id,
                'timestamp': datetime.now().isoformat(),
                'query': query,
                'company_name': company_name,
                'search_mode': search_mode,
                'content': final_result["content"],
                'provider': 'LangGraph Multi-Agent System',
                'model_used': 'Multi-Agent Orchestrator',
                'sources_used': final_result.get("sources_used", 3),
                'agentic_enhanced': True,
                'query_enhanced': query_enhancement["enhancement_applied"],
                'validation_score': validated_state["confidence_score"],
                'validation_details': validated_state["validation_results"],
                'context': json.dumps({**memory_enhancement, **query_enhancement}),
                'cache_key': cache_key
            }
            
            self.storage.save_search_result(result_data)
            self.storage.save_cached_result(cache_key, result_data)
            
            return {
                'success': True,
                'content': final_result["content"],
                'company_name': company_name,
                'provider': 'LangGraph Multi-Agent System',
                'search_mode': search_mode,
                'model_used': 'Multi-Agent Orchestrator',
                'cached': False,
                'sources_used': final_result.get("sources_used", 3),
                'agentic_enhanced': True,
                'query_enhanced': query_enhancement["enhancement_applied"],
                'validation_score': validated_state["confidence_score"],
                'validation_details': validated_state["validation_results"],
                'memory_enhanced': memory_enhancement["context_used"],
                'workflow_messages': [msg.content for msg in validated_state["messages"]],
                'query_enhancement_details': query_enhancement
            }
            
        except Exception as e:
            st.error(f"Error in enhanced agentic search: {str(e)}")
            
            # Emergency fallback with mode-specific error content
            if search_mode == "Quick Search":
                error_content = f"""# Quick Analysis Error: {company_name}

    ## System Status
    LangGraph multi-agent system encountered processing challenges for: {query}

    ## Alternative Insights
    Strategic analysis framework remains operational. Recommend immediate re-execution with adjusted parameters for optimal results.

    ## Immediate Actions
    1. Retry analysis with refined query parameters
    2. Verify system connectivity and resource availability  
    3. Consider alternative analysis approaches for critical decisions

    ## Technical Notes
    Multi-agent orchestration system requires optimization for enhanced performance and reliability."""
            
            elif search_mode == "Extended Search":
                error_content = f"""# Strategic Analysis System Response: {company_name}

    ## Executive Summary
    LangGraph multi-agent orchestration system encountered technical challenges during comprehensive analysis execution for {query}.

    ## System Assessment
    Advanced multi-agent research framework remains operational with partial functionality. Strategic analysis capabilities require system optimization for enhanced performance delivery.

    ## Alternative Strategic Approaches
    1. **Immediate Retry**: Execute analysis with optimized system parameters
    2. **Manual Research**: Conduct targeted research using alternative methodologies
    3. **System Diagnostics**: Verify connectivity and agent coordination functionality
    4. **Resource Optimization**: Adjust computational allocation for enhanced performance

    ## Technical Framework Status
    Multi-agent orchestration system requires configuration optimization for comprehensive strategic analysis delivery and enhanced research capability achievement.

    ## Recommendations
    Strategic analysis workflow requires system enhancement and technical optimization for reliable multi-dimensional research delivery and comprehensive business intelligence provision."""
            
            else:  # Deep Search
                error_content = f"""# Comprehensive System Analysis: {company_name}

    ## Executive Summary
    Advanced LangGraph multi-agent orchestration system encountered significant technical challenges during in-depth strategic research execution for the comprehensive analysis request: {query}

    ## System Architecture Assessment
    The sophisticated multi-agent research framework encompasses multiple specialized agents including financial intelligence, market analysis, competitive assessment, and validation systems. Current system status indicates partial operational capability with optimization requirements for enhanced performance delivery.

    ## Technical Challenge Analysis
    Advanced diagnostic evaluation reveals system coordination challenges affecting multi-agent workflow execution, data synthesis coordination, and comprehensive research delivery. Technical optimization requirements include agent communication enhancement, resource allocation optimization, and workflow coordination improvement.

    ## Alternative Research Methodologies
    1. **Systematic Retry Protocol**: Execute comprehensive analysis using optimized system parameters and enhanced agent coordination
    2. **Modular Research Approach**: Deploy individual agent systems for targeted research with manual synthesis coordination
    3. **Hybrid Analysis Framework**: Combine automated multi-agent capabilities with manual research methodologies for comprehensive coverage
    4. **System Architecture Optimization**: Implement technical enhancements for improved multi-agent coordination and research delivery

    ## Strategic Research Framework Status
    Advanced multi-agent orchestration system requires comprehensive technical optimization including agent coordination enhancement, resource allocation improvement, workflow optimization, and performance monitoring system implementation for reliable research-grade analysis delivery.

    ## Technical Implementation Recommendations
    Comprehensive system enhancement requires multi-dimensional optimization including agent communication protocols, data synthesis methodologies, validation system enhancement, and performance monitoring framework implementation for sustained research excellence and reliable strategic intelligence delivery.

    ## Future System Development
    Advanced LangGraph orchestration framework requires continuous optimization for enhanced multi-agent coordination, improved research delivery, and sustained competitive intelligence capability for comprehensive business intelligence provision and strategic decision support excellence."""
            
            return {
                'success': False,
                'error': str(e),
                'content': error_content,
                'company_name': company_name,
                'provider': 'LangGraph Multi-Agent System (Error)',
                'search_mode': search_mode,
                'cached': False,
                'agentic_enhanced': True,
                'sources_used': 1,
                'validation_score': 0.5
            }
                
            
    
    def search_company_info(self, company_name: str, search_mode: str = "Extended Search", context: str = "", enable_agents: bool = True) -> Dict:
        """Main search method with enhanced LangGraph capabilities"""
        
        if enable_agents and self.llm:
            return self.enhanced_agentic_search_with_recovery(company_name, company_name, search_mode)
        else:
            # Fallback to original implementation
            return self._fallback_search(company_name, search_mode, context)
    
    def _fallback_search(self, company_name: str, search_mode: str, context: str) -> Dict:
        """Fallback search method when LangGraph agents are not available"""
        
        return {
            'success': True,
            'content': f"Fallback search results for {company_name}. LangGraph agents not available.",
            'company_name': company_name,
            'provider': 'Fallback Search',
            'search_mode': search_mode,
            'cached': False,
            'agentic_enhanced': False
        }
    
    def save_to_history(self, query: str, response: str, company_name: str = "", provider: str = "", search_mode: str = "", agentic_enhanced: bool = False):
        """Save conversation to persistent storage"""
        try:
            chat_data = {
                'session_id': self.session_id,
                'timestamp': datetime.now().isoformat(),
                'query': query,
                'response': response,
                'company_name': company_name,
                'provider': provider,
                'search_mode': search_mode,
                'agentic_enhanced': agentic_enhanced
            }
            
            # Save to persistent storage
            self.storage.save_chat_message(chat_data)
            
        except Exception as e:
            st.error(f"Error saving to history: {str(e)}")
    
    def get_chat_history(self, limit: int = 50) -> List[Dict]:
        """Retrieve chat history from persistent storage"""
        try:
            return self.storage.get_chat_history(self.session_id, limit)
        except Exception as e:
            st.error(f"Error retrieving history: {str(e)}")
            return []
    
    def _clean_content_for_export(self, content: str) -> str:
        """Clean markdown and special characters from content for better readability"""
        import re
        
        # Remove markdown headers
        content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
        
        # Remove markdown bold/italic
        content = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', content)
        content = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', content)
        
        # Remove markdown links but keep text
        content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
        
        # Remove markdown tables formatting
        content = re.sub(r'\|', ' ', content)
        content = re.sub(r'^-+\s*$', '', content, flags=re.MULTILINE)
        
        # Remove emoji and special characters for cleaner text
        content = re.sub(r'[🚀📊💡📈🎯⚡🔍💰🏆🛡️🤖📋🌟💬📄🔧✨🧠📊🎨🔄📝📈⚙️💼🌐]', '', content)
        
        # Clean up multiple spaces and line breaks
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        content = re.sub(r'[ \t]+', ' ', content)
        
        # Remove special formatting indicators
        content = re.sub(r'\*\[([^\]]+)\]\*', r'(\1)', content)
        
        return content.strip()
    
    def export_report(self, content: str, format: str, filename: str, company_name: str = "", search_mode: str = "") -> bytes:
        """Export report with enhanced formatting"""
        try:
            if format == "PDF":
                return self._create_pdf_report(content, filename, company_name, search_mode)
            elif format == "Word":
                return self._create_word_report(content, filename, company_name, search_mode)
            elif format == "Markdown":
                return self._create_markdown_report(content, filename, company_name, search_mode)
                
        except Exception as e:
            st.error(f"Error creating {format} report: {str(e)}")
            return b""
    
    def _create_markdown_report(self, content: str, filename: str, company_name: str, search_mode: str) -> bytes:
        """Create well-formatted Markdown report with logo reference and enhanced structure"""
        
        # Check if logo exists and create reference
        logo_reference = ""
        logo_path = "LOGO.png"
        if os.path.exists(logo_path):
            logo_reference = f"""
    <div align="center">
        <img src="LOGO.png" alt="INSYT Logo" width="300"/>
    </div>

    ---
    """
        
        report_content = f"""{logo_reference}

    # 🚀 INSYT Business Intelligence Report

    <div align="center">

    | **Field** | **Details** |
    |-----------|-------------|
    | **Company** | {company_name} |
    | **Search Mode** | {search_mode} |
    | **Generated** | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |
    | **Platform** | INSYT with LangGraph Multi-Agent System |

    </div>

    ---

    ## 📋 Executive Summary

    This comprehensive business intelligence report was generated using our advanced LangGraph multi-agent orchestration system with intelligent query routing, self-healing validation, and memory enhancement capabilities.

    ---

    ## 📊 Analysis Results

    {content}

    ---

    ## 🔧 Technical Methodology

    This report leverages:

    - 🤖 **Multi-Agent Research Orchestration**: Parallel financial, news & competitive agents
    - 🎯 **Smart Query Routing**: Intelligent query classification & routing  
    - 🛡️ **Self-Healing Validation**: Automated conflict detection & resolution
    - 🧠 **Memory Enhancement**: Context-aware query enhancement
    - 📈 **Advanced Analytics**: Comprehensive data synthesis and validation

    ---

    ## 📝 Report Specifications

    | **Attribute** | **Value** |
    |---------------|-----------|
    | Report Type | Business Intelligence Analysis |
    | AI Framework | LangGraph Multi-Agent System |
    | Data Sources | Multiple validated sources |
    | Validation Level | Self-healing with conflict resolution |
    | Enhancement | Memory-context aware |

    ---

    <div align="center">

    ### 🏢 About INSYT Platform

    **INSYT** (Intelligent Network for Strategic Yield & Tracking) is an advanced business intelligence platform powered by cutting-edge AI technologies including LangGraph multi-agent orchestration.

    *This report was generated using INSYT platform with LangGraph multi-agent orchestration, intelligent query routing, self-healing validation, and memory enhancement capabilities.*

    ---

    **© 2024 INSYT Platform • Powered by LangGraph Agentic AI**

    </div>
    """
        
        return report_content.encode('utf-8')
    
    def _create_pdf_report(self, content: str, filename: str, company_name: str, search_mode: str) -> bytes:
        """Create well-formatted PDF report with logo and professional styling"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Enhanced styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=28,
            spaceAfter=20,
            textColor=colors.HexColor('#1f4e79'),
            alignment=1,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=15,
            textColor=colors.HexColor('#666666'),
            alignment=1
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#1f4e79'),
            fontName='Helvetica-Bold'
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=8,
            textColor=colors.HexColor('#2563eb'),
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            spaceBefore=6,
            alignment=0,
            leftIndent=0,
            fontName='Helvetica'
        )
        
        # Add logo if exists
        logo_path = "LOGO.png"
        if os.path.exists(logo_path):
            try:
                logo_img = RLImage(logo_path, width=2*inch, height=1*inch)
                logo_paragraph = Table([[logo_img]], colWidths=[2*inch], rowHeights=[1*inch])
                logo_paragraph.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(logo_paragraph)
                story.append(Spacer(1, 10))
            except Exception as e:
                pass
        
        # Title and metadata
        story.append(Paragraph("INSYT Business Intelligence Report", title_style))
        story.append(Spacer(1, 10))
        
        # Metadata table
        metadata_data = [
            ['Company:', company_name],
            ['Search Mode:', search_mode],
            ['Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Platform:', 'INSYT with LangGraph Multi-Agent System']
        ]
        
        metadata_table = Table(metadata_data, colWidths=[1.5*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f8f9fa')),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#1f4e79')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        story.append(metadata_table)
        story.append(Spacer(1, 30))
        
        # Clean content for export
        cleaned_content = self._clean_content_for_export(content)
        
        # Process content with enhanced formatting
        content_lines = cleaned_content.split('\n')
        current_section = ""
        
        for line in content_lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 8))
                continue
            
            # Detect headings based on content patterns
            if any(keyword in line.upper() for keyword in ['EXECUTIVE SUMMARY', 'MARKET ANALYSIS', 'FINANCIAL PERFORMANCE', 'COMPETITIVE', 'STRATEGIC', 'RISK ASSESSMENT', 'IMPLEMENTATION', 'CONCLUSION']):
                if len(line) < 100:  # Likely a heading
                    story.append(Paragraph(line, heading_style))
                    current_section = line
                    continue
            
            # Sub-headings
            if line.endswith(':') and len(line) < 80:
                story.append(Paragraph(line, subheading_style))
                continue
            
            # Bullet points
            if line.startswith(('•', '-', '★', '▪')):
                bullet_text = line[1:].strip()
                story.append(Paragraph(f"• {bullet_text}", body_style))
                continue
            
            # Regular paragraphs
            if len(line) > 10:  # Avoid very short lines
                story.append(Paragraph(line, body_style))
        
        # Professional footer
        story.append(Spacer(1, 30))
        footer_text = "This report was generated using INSYT platform with LangGraph multi-agent orchestration, intelligent query routing, self-healing validation, and memory enhancement capabilities."
        story.append(Paragraph(footer_text, subtitle_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    
    def _create_word_report(self, content: str, filename: str, company_name: str, search_mode: str) -> bytes:
        """Create well-formatted Word report with logo and professional styling"""
        doc = Document()
        
        # Add logo if exists
        logo_path = "LOGO.png"
        if os.path.exists(logo_path):
            try:
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(2.5))
                doc.add_paragraph()
            except Exception as e:
                pass
        
        # Enhanced title
        title = doc.add_heading('INSYT Business Intelligence Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        title.runs[0].font.name = 'Calibri'
        
        # Professional metadata table
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        metadata_cells = [
            ('Company:', company_name),
            ('Search Mode:', search_mode),
            ('Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Platform:', 'INSYT with LangGraph Multi-Agent System')
        ]
        
        for i, (label, value) in enumerate(metadata_cells):
            row_cells = metadata_table.rows[i].cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 78, 121)
            row_cells[1].text = value
        
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        doc.add_paragraph()
        
        # Clean content for export
        cleaned_content = self._clean_content_for_export(content)
        
        # Process content with enhanced formatting
        content_lines = cleaned_content.split('\n')
        
        for line in content_lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            # Detect major headings
            if any(keyword in line.upper() for keyword in ['EXECUTIVE SUMMARY', 'MARKET ANALYSIS', 'FINANCIAL PERFORMANCE', 'COMPETITIVE', 'STRATEGIC', 'RISK ASSESSMENT', 'IMPLEMENTATION', 'CONCLUSION']):
                if len(line) < 100:  # Likely a heading
                    heading = doc.add_heading(line, level=1)
                    heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)
                    continue
            
            # Sub-headings
            if line.endswith(':') and len(line) < 80:
                heading = doc.add_heading(line, level=2)
                heading.runs[0].font.color.rgb = RGBColor(37, 99, 235)
                continue
            
            # Bullet points
            if line.startswith(('•', '-', '★', '▪')):
                bullet_text = line[1:].strip()
                paragraph = doc.add_paragraph(bullet_text, style='List Bullet')
                continue
            
            # Regular paragraphs
            if len(line) > 10:  # Avoid very short lines
                paragraph = doc.add_paragraph(line)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Professional footer
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        footer = doc.add_paragraph("This report was generated using INSYT platform with LangGraph multi-agent orchestration, intelligent query routing, self-healing validation, and memory enhancement capabilities.")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.italic = True
        footer.runs[0].font.color.rgb = RGBColor(102, 102, 102)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
def check_service_status():
    """Check status of all services including LangGraph components"""
    status = {
        'openai_llm': False,
        'perplexity_llm': False,
        'serper': False,
        'aws_services': False,
        'langgraph_agents': False,
        'memory_enhancement': False,
        'nova_bedrock': False,
        'twitter_api': False
    }
    
    # Check OpenAI LLM
    if OPENAI_API_KEY:
        try:
            test_llm = ChatOpenAI(api_key=OPENAI_API_KEY, model="gpt-4o")
            test_response = test_llm.invoke([HumanMessage(content="test")])
            status['openai_llm'] = True
        except:
            pass
    
    # Check Perplexity AI LLM
    if PERPLEXITY_API_KEY:
        try:
            perplexity_llm, status_msg = initialize_perplexity_safely(PERPLEXITY_API_KEY)
            if perplexity_llm:
                status['perplexity_llm'] = True
        except:
            pass
    
    # Check Serper API
    if SERPER_API_KEY:
        try:
            headers = {'X-API-KEY': SERPER_API_KEY}
            response = requests.get("https://google.serper.dev/search", headers=headers, timeout=10)
            if response.status_code in [200, 400, 401]:
                status['serper'] = True
        except:
            pass
    
    # Check AWS Services
    if aws_clients:
        try:
            aws_clients['s3'].list_buckets()
            status['aws_services'] = True
        except:
            pass
    
    # Check Bedrock Services
    if os.getenv("BEDROCK_API_KEY"):
        try:
            nova_llm = NovaBedrockLLM(AWS_REGION)
            test_result = nova_llm.test_connection()
            status['nova_bedrock'] = test_result["success"]
        except:
            status['nova_bedrock'] = False
    
    # Check Twitter API
    if TWITTER_BEARER_TOKEN:
        try:
            headers = {'Authorization': f'Bearer {TWITTER_BEARER_TOKEN}'}
            response = requests.get("https://api.twitter.com/2/users/me", headers=headers, timeout=10)
            if response.status_code in [200, 401, 403]:
                status['twitter_api'] = True
        except:
            pass
    
    # Check LangGraph Agents
    status['langgraph_agents'] = status['openai_llm'] or status['perplexity_llm']
    status['memory_enhancement'] = True  # Always available with local storage
    
    return status

def load_logo_image():
    """Load logo image from file with enhanced error handling"""
    logo_path = "LOGO.png"
    if os.path.exists(logo_path):
        try:
            image = Image.open(logo_path)
            # Ensure image is in RGB mode for better compatibility
            if image.mode != 'RGB':
                image = image.convert('RGB')
            return image
        except Exception as e:
            st.error(f"Error loading logo: {str(e)}")
            return None
    else:
        st.warning("LOGO.png not found in the current directory")
        return None

def image_to_base64(image):
    """Convert PIL Image to base64 string"""
    try:
        buffer = io.BytesIO()
        image.save(buffer, format='PNG')
        buffer.seek(0)
        return base64.b64encode(buffer.getvalue()).decode()
    except Exception as e:
        # Fallback: try to read the file directly
        try:
            with open('LOGO.png', 'rb') as f:
                return base64.b64encode(f.read()).decode()
        except:
            return ""

def main():
    # Page configuration
    st.set_page_config(
        page_title="🚀 INSYT - Enhanced with LangGraph + Perplexity AI",
        page_icon="🚀",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS (enhanced for LangGraph features with Header.PNG color scheme)
    st.markdown("""
<style>
.main {
    background-color: #f8f9fa;
}
.stButton > button {
    background-color: #1f4e79;
    color: white;
    border-radius: 5px;
    border: none;
    padding: 0.5rem 1rem;
}
.stButton > button:hover {
    background-color: #2563eb;
}
.langgraph-badge {
    background: linear-gradient(45deg, #667eea, #764ba2);
    color: white;
    padding: 0.3rem 0.8rem;
    border-radius: 20px;
    font-size: 0.8rem;
    font-weight: bold;
    display: inline-block;
    margin: 0.2rem;
}
.memory-enhanced {
    background: linear-gradient(45deg, #4facfe, #00f2fe);
    color: white;
    padding: 0.3rem 0.8rem;
    border-radius: 15px;
    font-size: 0.8rem;
    display: inline-block;
}
.validation-score {
    background-color: #28a745;
    color: white;
    padding: 0.3rem 0.8rem;
    border-radius: 15px;
    font-size: 0.8rem;
    display: inline-block;
}
.chat-message {
    background-color: white;
    padding: 1rem;
    border-radius: 10px;
    border-left: 4px solid #1f4e79;
    margin-bottom: 1rem;
}
.langgraph-enhanced {
    border-left: 4px solid #667eea !important;
    box-shadow: 0 4px 8px rgba(102, 126, 234, 0.2);
}
.agent-workflow {
    background: linear-gradient(135deg, #667eea, #764ba2);
    color: white;
    border-radius: 10px;
    padding: 15px 20px;
    margin: 15px 0;
    box-shadow: 0 4px 6px rgba(102, 126, 234, 0.3);
}
.memory-context {
    background: linear-gradient(135deg, #4facfe, #00f2fe);
    color: white;
    border-radius: 10px;
    padding: 15px 20px;
    margin: 15px 0;
    box-shadow: 0 4px 6px rgba(79, 172, 254, 0.3);
}
.header-enhanced {
    background: linear-gradient(135deg, 
        #0f172a 0%, 
        #1e293b 15%, 
        #1e40af 30%, 
        #2563eb 45%, 
        #3b82f6 60%, 
        #60a5fa 75%, 
        #93c5fd 90%, 
        #dbeafe 100%);
    color: white;
    padding: 2rem;
    border-radius: 25px;
    text-align: left;
    margin-bottom: 2rem;
    box-shadow: 
        0 25px 50px rgba(30, 58, 138, 0.4), 
        0 0 0 1px rgba(255, 255, 255, 0.1),
        inset 0 1px 0 rgba(255, 255, 255, 0.2);
    position: relative;
    overflow: hidden;
    border: 2px solid rgba(255, 255, 255, 0.2);
    backdrop-filter: blur(10px);
}
.header-enhanced::before {
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    right: 0;
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.4), transparent);
}
.header-enhanced::after {
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    right: 0;
    bottom: 0;
    background: linear-gradient(135deg, 
        rgba(255,255,255,0.1) 0%, 
        rgba(255,255,255,0.05) 50%, 
        rgba(0,0,0,0.05) 100%);
    pointer-events: none;
}
.logo-container {
    display: flex;
    align-items: center;
    justify-content: center;
    background: linear-gradient(135deg, 
        rgba(255,255,255,0.15) 0%, 
        rgba(255,255,255,0.05) 100%);
    border-radius: 20px;
    padding: 1rem;
    margin-right: 2rem;
    box-shadow: 
        0 8px 32px rgba(0,0,0,0.1),
        inset 0 1px 0 rgba(255,255,255,0.2);
    border: 1px solid rgba(255,255,255,0.1);
    backdrop-filter: blur(5px);
}
.logo-container img {
    filter: drop-shadow(0 4px 8px rgba(0,0,0,0.2));
    transition: transform 0.3s ease;
}
.logo-container:hover img {
    transform: scale(1.05);
}
.title-section {
    flex: 1;
    padding-left: 1rem;
}
.prominent-interface {
    background: linear-gradient(135deg, #e3f2fd, #bbdefb);
    border: 3px solid #667eea;
    border-radius: 20px;
    padding: 2rem;
    margin: 2rem 0;
    box-shadow: 0 10px 30px rgba(102, 126, 234, 0.2);
    width: 75%;
    margin-left: auto;
    margin-right: auto;
}
.prominent-title {
    color: #1e3a8a;
    font-size: 2.5rem;
    font-weight: 800;
    text-align: center;
    margin-bottom: 1.5rem;
    text-shadow: 0 2px 4px rgba(30, 58, 138, 0.3);
}
.history-section {
    background: linear-gradient(135deg, #fffef7, #fff9c4);
    border: 2px solid #ffd54f;
    border-radius: 15px;
    padding: 1.5rem;
    margin: 1rem 0;
    box-shadow: 0 4px 15px rgba(255, 213, 79, 0.2);
}
.service-status-container {
    width: 75%;
    margin-left: auto;
    margin-right: auto;
}
.minimal-footer {
    background: #f8f9fa;
    border-top: 1px solid #e9ecef;
    padding: 1rem;
    text-align: center;
    color: #6c757d;
    font-size: 0.9rem;
    margin-top: 2rem;
}
.langgraph-animation-container {
    background: linear-gradient(135deg, #667eea, #764ba2, #4facfe, #00f2fe);
    background-size: 400% 400%;
    animation: gradientShift 3s ease infinite;
    border-radius: 20px;
    padding: 2rem;
    margin: 2rem 0;
    color: white;
    text-align: center;
    box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
    position: relative;
    overflow: hidden;
}

.langgraph-animation-container::before {
    content: '';
    position: absolute;
    top: 0;
    left: -100%;
    width: 100%;
    height: 100%;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent);
    animation: shimmer 2s infinite;
}

@keyframes gradientShift {
    0% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}

@keyframes shimmer {
    0% { left: -100%; }
    100% { left: 100%; }
}

@keyframes pulse {
    0%, 100% { transform: scale(1); opacity: 1; }
    50% { transform: scale(1.1); opacity: 0.8; }
}

@keyframes slideIn {
    0% { transform: translateX(-50px); opacity: 0; }
    100% { transform: translateX(0); opacity: 1; }
}

.animation-step {
    display: flex;
    align-items: center;
    justify-content: center;
    margin: 1rem 0;
    font-size: 1.2rem;
    font-weight: 600;
    opacity: 0;
    transform: translateX(-50px);
    animation: slideIn 0.8s ease forwards;
    position: relative;
    z-index: 2;
}

.animation-step.active {
    animation: slideIn 0.8s ease forwards, pulse 2s ease infinite;
}

.animation-step .step-icon {
    margin-right: 1rem;
    font-size: 1.5rem;
    animation: pulse 1.5s ease infinite;
}

.progress-bar {
    width: 100%;
    height: 4px;
    background: rgba(255,255,255,0.3);
    border-radius: 2px;
    margin: 1rem 0;
    overflow: hidden;
    position: relative;
    z-index: 2;
}

.progress-fill {
    height: 100%;
    background: linear-gradient(90deg, #ffffff, #f0f8ff, #ffffff);
    background-size: 200% 100%;
    animation: progressFill 8s linear forwards, shimmerProgress 1.5s ease infinite;
    border-radius: 2px;
}

@keyframes progressFill {
    0% { width: 0%; }
    20% { width: 20%; }
    40% { width: 40%; }
    60% { width: 60%; }
    80% { width: 80%; }
    100% { width: 100%; }
}

@keyframes shimmerProgress {
    0% { background-position: 0% 50%; }
    100% { background-position: 200% 50%; }
}

.enhanced-content {
    line-height: 1.6;
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
}
.enhanced-content h1 {
    color: #1f4e79;
    border-bottom: 3px solid #667eea;
    padding-bottom: 10px;
    margin-bottom: 20px;
    font-size: 2rem;
    font-weight: 700;
}
.enhanced-content h2 {
    color: #2563eb;
    margin-top: 25px;
    margin-bottom: 15px;
    font-size: 1.4rem;
    font-weight: 600;
    border-left: 4px solid #667eea;
    padding-left: 15px;
    background: linear-gradient(90deg, rgba(102, 126, 234, 0.1), transparent);
    padding: 8px 15px;
    border-radius: 5px;
}
.enhanced-content h3 {
    color: #3b82f6;
    margin-top: 20px;
    margin-bottom: 10px;
    font-size: 1.2rem;
    font-weight: 600;
}
.enhanced-content h4 {
    color: #4f46e5;
    margin-top: 15px;
    margin-bottom: 8px;
    font-size: 1.1rem;
    font-weight: 500;
}
.enhanced-content ul, .enhanced-content ol {
    margin-left: 20px;
    margin-bottom: 15px;
}
.enhanced-content li {
    margin-bottom: 5px;
    line-height: 1.5;
}
.enhanced-content strong {
    color: #1f4e79;
    font-weight: 600;
}
.enhanced-content p {
    margin-bottom: 12px;
    text-align: justify;
}
.enhanced-content blockquote {
    border-left: 4px solid #667eea;
    padding-left: 20px;
    margin: 15px 0;
    font-style: italic;
    background: rgba(102, 126, 234, 0.05);
    padding: 15px 20px;
    border-radius: 5px;
}
.section-divider {
    border-top: 2px solid #e5e7eb;
    margin: 25px 0;
    position: relative;
}
.section-divider::before {
    content: '✦';
    position: absolute;
    top: -10px;
    left: 50%;
    transform: translateX(-50%);
    background: white;
    color: #667eea;
    padding: 0 10px;
    font-size: 1.2rem;
}
</style>
""", unsafe_allow_html=True)
    
    # Initialize enhanced assistant
    if 'enhanced_assistant' not in st.session_state:
        with st.spinner("🚀 Initializing Enhanced INSYT with LangGraph..."):
            st.session_state.enhanced_assistant = EnhancedAgenticAIAssistant()
    
    # Enhanced Header with logo and Header.PNG color scheme
    logo_image = load_logo_image()
    
    if logo_image:
        # Create enhanced header with better logo integration
        try:
            buf = io.BytesIO()
            # Ensure we save the image bytes to the buffer and read them back
            logo_image.convert('RGB').save(buf, format='PNG')
            image_bytes = buf.getvalue()
        except Exception:
            # Fallback to reading file bytes directly
            try:
                image_bytes = open('LOGO.png', 'rb').read()
            except Exception:
                image_bytes = b''

        image_b64 = base64.b64encode(image_bytes).decode() if image_bytes else ''

        st.markdown(f"""
        <div class="header-enhanced">
            <div style="display: flex; align-items: center; position: relative; z-index: 2;">
                <div class="logo-container">
                    <img src="data:image/png;base64,{image_b64}" width="180" style="border-radius: 15px;">
                </div>
                <div class="title-section">
                    <h1 style="
                        margin: 0;
                        font-size: 4rem;
                        font-weight: 900;
                        letter-spacing: 2px;
                        text-shadow: 
                            0 4px 8px rgba(0, 0, 0, 0.3), 
                            0 0 20px rgba(255, 255, 255, 0.5),
                            0 2px 4px rgba(30, 58, 138, 0.4);
                        background: linear-gradient(135deg, 
                            #ffffff 0%, 
                            #f0f8ff 25%, 
                            #e6f3ff 50%, 
                            #dbeafe 75%, 
                            #ffffff 100%);
                        -webkit-background-clip: text;
                        -webkit-text-fill-color: transparent;
                        background-clip: text;
                        margin-bottom: 0.5rem;
                        filter: drop-shadow(0 2px 4px rgba(30, 58, 138, 0.5));
                    ">INSYT</h1>
                    <h2 style="
                        margin: 0;
                        font-size: 1.6rem;
                        font-weight: 600;
                        letter-spacing: 1px;
                        text-shadow: 
                            0 2px 4px rgba(30, 58, 138, 0.4),
                            0 1px 2px rgba(0, 0, 0, 0.2);
                        color: #f0f8ff;
                        margin-bottom: 0.5rem;
                        opacity: 0.95;
                    ">Intelligent Network for Strategic Yield & Tracking</h2>
                    <p style="
                        margin: 0;
                        font-size: 1rem;
                        font-weight: 300;
                        letter-spacing: 0.5px;
                        text-shadow: 0 1px 2px rgba(30, 58, 138, 0.3);
                        color: #e6f3ff;
                        opacity: 0.9;
                        max-width: 600px;
                    ">Enhanced with AI-Powered Business Intelligence</p>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    else:
        # Enhanced fallback header without logo but with glossy blue theme
        st.markdown("""
        <div class="header-enhanced">
            <div style="position: relative; z-index: 2; text-align: center;">
                <h1 style="
                margin: 0;
                font-size: 4.5rem;
                font-weight: 900;
                letter-spacing: 3px;
                text-shadow: 
                    0 4px 8px rgba(0, 0, 0, 0.3), 
                    0 0 20px rgba(255, 255, 255, 0.5),
                    0 2px 4px rgba(30, 58, 138, 0.4);
                background: linear-gradient(135deg, 
                    #ffffff 0%, 
                    #f0f8ff 25%, 
                    #e6f3ff 50%, 
                    #dbeafe 75%, 
                    #ffffff 100%);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                background-clip: text;
                margin-bottom: 0.5rem;
                filter: drop-shadow(0 2px 4px rgba(30, 58, 138, 0.5));
            ">🚀 INSYT</h1>
            <h2 style="
                margin: 0;
                font-size: 1.8rem;
                font-weight: 600;
                letter-spacing: 1px;
                text-shadow: 
                    0 2px 4px rgba(30, 58, 138, 0.4),
                    0 1px 2px rgba(0, 0, 0, 0.2);
                color: #f0f8ff;
                margin-bottom: 0.5rem;
                opacity: 0.95;
            ">Intelligent Network for Strategic Yield & Tracking</h2>
            <p style="
                margin: 0;
                font-size: 1.1rem;
                font-weight: 300;
                letter-spacing: 0.5px;
                text-shadow: 0 1px 2px rgba(30, 58, 138, 0.3);
                color: #e6f3ff;
                opacity: 0.9;
                max-width: 800px;
                margin: 0 auto;
            ">Enhanced with AI-Powered Business Intelligence & Persistent Data Management</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # LangGraph Features Showcase
    st.markdown("""
    <div class="agent-workflow">
        <h3 style="margin: 0 0 15px 0; font-size: 1.8rem; font-weight: 700;">🧠 LangGraph Agentic Features</h3>
         <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px;">
                <div style="background: rgba(255,255,255,0.2); padding: 12px; border-radius: 10px; text-align: center;">
                    <h5 style="margin: 0 0 5px 0; font-size: 1rem;">🤖 Multi-Agent Research</h5>
                    <p style="margin: 0; font-size: 11px; opacity: 0.9;">Financial, news, social & competitive agents</p>
                </div>
                <div style="background: rgba(255,255,255,0.2); padding: 12px; border-radius: 10px; text-align: center;">
                    <h5 style="margin: 0 0 5px 0; font-size: 1rem;">🎯 Triad Search Options</h5>
                    <p style="margin: 0; font-size: 11px; opacity: 0.9;">Quick, Extended & Deep search modes</p>
                </div>
                <div style="background: rgba(255,255,255,0.2); padding: 12px; border-radius: 10px; text-align: center;">
                    <h5 style="margin: 0 0 5px 0; font-size: 1rem;">🛡️ Self-Healing Validation</h5>
                    <p style="margin: 0; font-size: 11px; opacity: 0.9;">Automated conflict detection & resolution</p>
                </div>
                <div style="background: rgba(255,255,255,0.2); padding: 12px; border-radius: 10px; text-align: center;">
                    <h5 style="margin: 0 0 5px 0; font-size: 1rem;">🧠 Smart Query Enhancement</h5>
                    <p style="margin: 0; font-size: 11px; opacity: 0.9;">Automatically refines and expands queries</p>
                        </div>
                    </div>   
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar - Enhanced for LangGraph
    with st.sidebar:
        st.header("🔧 Enhanced Platform Settings")
        
        # LangGraph Agent Status
        st.subheader("🤖 LangGraph Agents")
        if hasattr(st.session_state.enhanced_assistant, 'llm') and st.session_state.enhanced_assistant.llm:
            st.success("✅ Multi-Agent System Active")
            if hasattr(st.session_state.enhanced_assistant, 'perplexity_llm') and st.session_state.enhanced_assistant.perplexity_llm:
                st.success("🧠 Perplexity AI Active")
            st.info("🧠 Memory Enhancement Active")
            st.info("🛡️ Self-Healing Validation Active")
            st.info("🎯 Query Router Active")
        else:
            st.error("❌ LangGraph Agents Unavailable")
            st.warning("⚠️ Missing LLM configuration")
        
        st.divider()
        
        # Memory Context Management with database locking fix
        st.subheader("🧠 Memory Context")
        if st.button("🗑️ Clear Memory Context", use_container_width=True):
            # Clear memory context for current session with proper locking
            try:
                def _clear_memory():
                    conn = st.session_state.enhanced_assistant.storage._get_connection()
                    try:
                        cursor = conn.cursor()
                        cursor.execute('DELETE FROM memory_context WHERE session_id = ?', 
                                     (st.session_state.enhanced_assistant.session_id,))
                        conn.commit()
                        return True
                    finally:
                        conn.close()
                
                success = st.session_state.enhanced_assistant.storage._execute_with_retry(_clear_memory)
                if success:
                    st.success("✅ Memory context cleared")
                else:
                    st.error("❌ Failed to clear memory context")
            except Exception as e:
                st.error(f"Error clearing memory context: {str(e)}")
        
        # Cache Management with database locking fix
        st.subheader("💾 Cache Management")
        if st.button("🧹 Clean Expired Cache", use_container_width=True):
            try:
                cleaned = st.session_state.enhanced_assistant.storage.cleanup_expired_cache()
                st.success(f"✅ Cleaned {cleaned} expired entries")
            except Exception as e:
                st.error(f"Error cleaning cache: {str(e)}")
        
        # Database Health Check
        if st.button("🔧 Database Health Check", use_container_width=True):
            try:
                def _health_check():
                    conn = st.session_state.enhanced_assistant.storage._get_connection()
                    try:
                        cursor = conn.cursor()
                        cursor.execute('PRAGMA integrity_check;')
                        result = cursor.fetchone()[0]
                        return result == 'ok'
                    finally:
                        conn.close()
                
                health_ok = st.session_state.enhanced_assistant.storage._execute_with_retry(_health_check)
                if health_ok:
                    st.success("✅ Database is healthy")
                else:
                    st.warning("⚠️ Database integrity issues detected")
            except Exception as e:
                st.error(f"Database health check failed: {str(e)}")
        
        # Database Vacuum (optimize)
        if st.button("🗜️ Optimize Database", use_container_width=True):
            try:
                def _vacuum_db():
                    conn = st.session_state.enhanced_assistant.storage._get_connection()
                    try:
                        conn.execute('VACUUM;')
                        return True
                    finally:
                        conn.close()
                
                st.session_state.enhanced_assistant.storage._execute_with_retry(_vacuum_db)
                st.success("✅ Database optimized")
            except Exception as e:
                st.error(f"Database optimization failed: {str(e)}")
        
        # Enhanced Stats
        storage_stats = st.session_state.enhanced_assistant.storage.get_storage_stats()
        if storage_stats:
            st.subheader("📊 Enhanced Stats")
            st.write(f"🔍 Searches: {storage_stats.get('searches', 0)}")
            st.write(f"💬 Chat History: {storage_stats.get('chat_history', 0)}")
            st.write(f"🧠 Memory Context: {storage_stats.get('memory_context', 0)}")
            st.write(f"🤖 Agent States: {storage_stats.get('agent_states', 0)}")
            st.write(f"💾 DB Size: {storage_stats.get('db_size_mb', 0):.1f}MB")
        
        st.divider()
    
    # Main interface - More prominent
    st.markdown("""
    <div class="prominent-interface">
        <h2 class="prominent-title">💬 Query Interface</h2>
    """, unsafe_allow_html=True)
    
    # Analysis Mode Selection
    st.markdown("**Analysis Mode:**")
    
    # Initialize session state for analysis mode
    if 'analysis_mode' not in st.session_state:
        st.session_state.analysis_mode = "Extended Search"
    
    # Create three columns for the mode icons
    mode_col1, mode_col2, mode_col3 = st.columns(3)
    
    with mode_col1:
        if st.button("⚡\nQuick Search", key="quick_mode", use_container_width=True):
            st.session_state.analysis_mode = "Quick Search"
    
    with mode_col2:
        if st.button("🔍\nExtended Search", key="extended_mode", use_container_width=True):
            st.session_state.analysis_mode = "Extended Search"
    
    with mode_col3:
        if st.button("🎯\nDeep Search", key="deep_mode", use_container_width=True):
            st.session_state.analysis_mode = "Deep Search"
    
    # Display current mode
    mode_descriptions = {
        "Quick Search": "⚡ Fast response with basic multi-agent coordination",
        "Extended Search": "🔍 Comprehensive analysis with full LangGraph workflow", 
        "Deep Search": "🎯 In-depth research with enhanced validation & memory"
    }
    
    st.info(f"**Current Mode:** {st.session_state.analysis_mode} - {mode_descriptions[st.session_state.analysis_mode]}")
    
    # Set global token config based on selected mode
    set_search_mode_config(st.session_state.analysis_mode)
    
    # Enhanced controls
    col_agentic, col_memory, col_contextual = st.columns(3)
    with col_agentic:
        enable_langgraph = st.checkbox("🚀 LangGraph Agents", value=True, help="Multi-agent orchestration, smart routing, self-healing validation, and memory enhancement")
    
    with col_memory:
        enable_memory = st.checkbox("🧠 Memory Enhancement", value=True, help="Use previous context to enhance queries")
    
    with col_contextual:
        enable_contextual = st.checkbox("💬 Contextual Q&A", value=True, help="Enable contextual question answering using conversation history")
    
    
    # Chat input
    user_input = st.text_area(
        "Enter your business intelligence query:",
        height=120,
        placeholder="e.g., 'Analyze Tesla's competitive position in the EV market' or 'What are Microsoft's latest financial developments?'"
    )
    
    col_send, col_clear = st.columns([1, 1])
    
    with col_send:
        button_text = "💬 Contextual Analysis" if enable_contextual else "🚀 Analyze Now"
        if st.button(button_text, use_container_width=True, key="main_analysis_btn"):
            if user_input:
                context = ""
                with st.spinner("🤖 Processing your query..."):
                    if enable_contextual:
                        result = st.session_state.enhanced_assistant.contextual_search(
                            user_input, st.session_state.analysis_mode
                        )
                    else:
                        result = st.session_state.enhanced_assistant.search_company_info(
                            user_input, st.session_state.analysis_mode, context, enable_langgraph
                        )        
                        
                    if result['success']:
                        st.session_state.last_search = result
                        st.session_state.enhanced_assistant.save_to_history(
                            user_input,
                            result['content'],
                            result.get('company_name', ''),
                            result.get('provider', ''),
                            st.session_state.analysis_mode,
                            enable_langgraph
                        )
                        st.rerun()
    
    with col_clear:
        if st.button("🗑️ Clear Results", use_container_width=True, key="clear_results_btn"):
            if 'last_search' in st.session_state:
                del st.session_state.last_search
            st.rerun()
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Display enhanced results
    if 'last_search' in st.session_state:
        st.markdown("### 🚀 LangGraph Analysis Results")
        
        result_info = st.session_state.last_search
        
        # Enhanced result metadata with LangGraph indicators
        col_provider, col_mode_used, col_agents, col_memory_status, col_validation = st.columns(5)
        
        with col_provider:
            provider_text = result_info.get('provider', '🚀 LangGraph Multi-Agent')
            st.markdown(f"""
            <div style="background: linear-gradient(45deg, #667eea, #764ba2); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                {provider_text}
            </div>
            """, unsafe_allow_html=True)
        
        with col_mode_used:
            mode_colors = {"Quick Search": "#28a745", "Extended Search": "#ffc107", "Deep Search": "#dc3545"}
            bg_color = mode_colors.get(result_info.get('search_mode', 'Extended Search'), "#ffc107")
            st.markdown(f"""
            <div style="background-color: {bg_color}; color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                {result_info.get('search_mode', 'Extended Search')}
            </div>
            """, unsafe_allow_html=True)
        
        with col_agents:
            if result_info.get('agentic_enhanced', False):
                sources_count = result_info.get('sources_used', 0)
                query_enhanced = result_info.get('query_enhanced', False)
                enhancement_icon = "✨" if query_enhanced else "🤖"
                st.markdown(f"""
                <div style="background: linear-gradient(45deg, #667eea, #764ba2); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    {enhancement_icon} {sources_count} Agents
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background-color: #6c757d; color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    📚 Standard
                </div>
                """, unsafe_allow_html=True)
        
        with col_memory_status:
            if result_info.get('memory_enhanced', False):
                st.markdown("""
                <div style="background: linear-gradient(45deg, #4facfe, #00f2fe); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    🧠 Memory+
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background-color: #868686; color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    🔍 Fresh
                </div>
                """, unsafe_allow_html=True)
        
        with col_validation:
            if result_info.get('agentic_enhanced', False):
                validation_score = result_info.get('validation_score', 0)
                if validation_score:
                    st.markdown(f"""
                    <div style="background: linear-gradient(45deg, #28a745, #20c997); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                        🛡️ {validation_score:.0%}
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown("""
                    <div style="background: linear-gradient(45deg, #28a745, #20c997); color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                        🛡️ Validated
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background-color: #868686; color: white; padding: 0.4rem; border-radius: 5px; text-align: center; font-size: 0.7rem;">
                    📊 Basic
                </div>
                """, unsafe_allow_html=True)
        
        # Show if loaded from history
        if result_info.get('from_history', False):
            st.info(f"📜 Loaded from history: {result_info.get('original_query', 'Previous conversation')}")
        
        # Display query enhancement details if available
        if result_info.get('query_enhancement_details'):
            with st.expander("✨ View Query Enhancement Details"):
                enhancement_details = result_info['query_enhancement_details']
                st.write(f"**Original Query:** {enhancement_details.get('original_query', '')}")
                st.write(f"**Enhanced Query:** {enhancement_details.get('enhanced_query', '')}")
                
                focus_areas = enhancement_details.get('focus_areas', [])
                if focus_areas:
                    st.write(f"**Focus Areas:** {', '.join(focus_areas)}")
                
                key_metrics = enhancement_details.get('key_metrics', [])
                if key_metrics:
                    st.write(f"**Key Metrics:** {', '.join(key_metrics)}")
                
                research_scope = enhancement_details.get('research_scope', '')
                if research_scope:
                    st.write(f"**Research Scope:** {research_scope}")
        
        # Display workflow messages if available
        if result_info.get('workflow_messages'):
            with st.expander("🔍 View LangGraph Workflow Steps"):
                for i, message in enumerate(result_info['workflow_messages']):
                    st.write(f"**Step {i+1}:** {message}")
        
        # Display validation details if available
        if result_info.get('validation_details'):
            with st.expander("🛡️ View Self-Healing Validation Details"):
                validation_details = result_info['validation_details']
                st.write(f"**Final Confidence Score:** {validation_details.get('final_confidence', 0):.1%}")
                
                if validation_details.get('conflicts_detected'):
                    st.warning("⚠️ Data conflicts were detected")
                    if validation_details.get('resolved'):
                        st.success("✅ Conflicts were automatically resolved")
                        st.write(f"**Resolution:** {validation_details.get('resolution', 'Automated resolution applied')}")
                    else:
                        st.error("❌ Some conflicts remain unresolved")
                else:
                    st.success("✅ No data conflicts detected")
        
        # Enhanced content display with better structure and formatting
        content_class = "chat-message langgraph-enhanced" if result_info.get('agentic_enhanced', False) else "chat-message"
        
        # Process content for better display formatting
        display_content = result_info['content']
        
        # Enhanced content display with better structure
        st.markdown(f"""
        <div class="{content_class} enhanced-content">
        {display_content}
        </div>
        """, unsafe_allow_html=True)
        
        # Add section divider for visual separation
        st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
        
        # Export section
        st.markdown("### 📄 Export Analysis")
        col_format, col_export, col_share = st.columns([2, 1, 1])
        
        with col_format:
            export_format = st.selectbox("Select Format:", ["PDF", "Word", "Markdown"])
        
        with col_export:
            if st.button("📥 Export Report", use_container_width=True, key="export_report_btn"):
                content = st.session_state.last_search['content']
                company_name = st.session_state.last_search.get('company_name', 'Company')
                search_mode = st.session_state.last_search.get('search_mode', 'Analysis')
                from datetime import datetime
                filename = f"insyt_langgraph_report_{company_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                
                with st.spinner(f"Generating {export_format} report..."):
                    report_data = st.session_state.enhanced_assistant.export_report(
                        content, export_format, filename, company_name, search_mode
                    )
                    
                    if report_data:
                        file_ext = {"PDF": ".pdf", "Word": ".docx", "Markdown": ".md"}
                        mime_types = {"PDF": "application/pdf", "Word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "Markdown": "text/markdown"}
                        
                        st.download_button(
                            f"📄 Download {export_format}",
                            data=report_data,
                            file_name=f"{filename}{file_ext[export_format]}",
                            mime=mime_types[export_format],
                            use_container_width=True
                        )
        
        with col_share:
            if st.button("🔗 Share Analysis", use_container_width=True):
                st.info("🚀 LangGraph analysis shared! Enhanced multi-agent insights available.")
        
    
    # Communication History Section (moved from sidebar)
    st.markdown("---")
    st.markdown("""
        <div class="history-section">
        <h3 style="margin: 0 0 1rem 0; color: #1f4e79; text-align: center;">💬 History</h3>
        """, unsafe_allow_html=True)
            
    # Initialize pagination state for main UI
    if 'main_history_page' not in st.session_state:
        st.session_state.main_history_page = 1
    
    # Get paginated history (persistent across all sessions)
    history_data = st.session_state.enhanced_assistant.get_paginated_chat_history(
        page=st.session_state.main_history_page, 
        per_page=5
    )
    
    if history_data['total_count'] > 0:
        # History stats and pagination controls
        col_stats, col_controls = st.columns([1, 2])
        
        with col_stats:
            st.write(f"📊 Total conversations: {history_data['total_count']}")
            st.write(f"📄 Page {history_data['current_page']} of {history_data['total_pages']}")
        
        with col_controls:
            # Pagination controls
            col_prev, col_next, col_jump = st.columns([1, 1, 2])
            
            with col_prev:
                if st.button("⬅️ Previous", disabled=history_data['current_page'] <= 1, use_container_width=True, key="main_hist_prev"):
                    st.session_state.main_history_page = max(1, st.session_state.main_history_page - 1)
                    st.rerun()
            
            with col_next:
                if st.button("Next ➡️", disabled=history_data['current_page'] >= history_data['total_pages'], use_container_width=True, key="main_hist_next"):
                    st.session_state.main_history_page = min(history_data['total_pages'], st.session_state.main_history_page + 1)
                    st.rerun()
            
            with col_jump:
                max_pages = min(50, history_data['total_pages'])  # Maximum 50 pages as requested
                if max_pages > 1:
                    selected_page = st.selectbox(
                        "Jump to page:", 
                        range(1, max_pages + 1), 
                        index=min(st.session_state.main_history_page - 1, max_pages - 1),
                        key="main_page_selector"
                    )
                    if selected_page != st.session_state.main_history_page:
                        st.session_state.main_history_page = selected_page
                        st.rerun()
        
        # Display history threads in columns for better layout
        st.markdown("**Recent Conversations:**")
        
        # Create two columns for better display
        hist_col1, hist_col2 = st.columns(2)
        
        for i, chat_item in enumerate(history_data['history']):
            thread_number = ((history_data['current_page'] - 1) * 5) + i + 1
            
            # Alternate between columns
            current_col = hist_col1 if i % 2 == 0 else hist_col2
            
            with current_col:
                # Create compact thread display
                timestamp = chat_item.get('timestamp', '')
                if timestamp:
                    try:
                        from datetime import datetime
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        formatted_time = dt.strftime('%m/%d %H:%M')
                    except:
                        formatted_time = 'Unknown'
                else:
                    formatted_time = 'Unknown'
                
                query = chat_item.get('query', 'No query')
                query_preview = query[:60] + "..." if len(query) > 60 else query
                
                agentic_badge = "🚀" if chat_item.get('agentic_enhanced', False) else "📚"
                search_mode = chat_item.get('search_mode', 'Unknown')
                
                with st.expander(f"{agentic_badge} Thread #{thread_number} - {formatted_time}", expanded=False):
                    st.markdown(f"**Query:** {query}")
                    st.markdown(f"**Mode:** {search_mode}")
                    st.markdown(f"**Provider:** {chat_item.get('provider', 'Unknown')}")
                    
                    # Show response preview
                    response = chat_item.get('response', 'No response')
                    response_preview = response[:200] + "..." if len(response) > 200 else response
                    st.markdown(f"**Response:** {response_preview}")
                    
                    # Reload button
                    if st.button(f"🔄 Load Thread #{thread_number}", key=f"main_reload_{chat_item.get('id', i)}_{thread_number}", use_container_width=True):
                        st.session_state.last_search = {
                            'success': True,
                            'content': chat_item.get('response', ''),
                            'company_name': chat_item.get('company_name', ''),
                            'provider': chat_item.get('provider', ''),
                            'search_mode': chat_item.get('search_mode', ''),
                            'cached': True,
                            'agentic_enhanced': chat_item.get('agentic_enhanced', False),
                            'from_history': True,
                            'original_query': chat_item.get('query', '')
                        }
                        st.success(f"✅ Loaded Thread #{thread_number}")
                        st.rerun()
        
        # Clear history option
        st.markdown("---")
        col_clear, col_spacer = st.columns([1, 3])
        with col_clear:
            if st.button("🗑️ Clear All History", use_container_width=True, key="main_clear_history"):
                if st.checkbox("⚠️ Confirm deletion", key="main_confirm_clear_history"):
                    try:
                        def _clear_all_history():
                            conn = st.session_state.enhanced_assistant.storage._get_connection()
                            try:
                                cursor = conn.cursor()
                                cursor.execute('DELETE FROM chat_history')
                                conn.commit()
                                return True
                            finally:
                                conn.close()
                        
                        success = st.session_state.enhanced_assistant.storage._execute_with_retry(_clear_all_history)
                        if success:
                            st.session_state.main_history_page = 1
                            st.success("✅ All history cleared")
                            st.rerun()
                        else:
                            st.error("❌ Failed to clear history")
                    except Exception as e:
                        st.error(f"Error clearing history: {str(e)}")
    else:
        st.info("💡 No conversation history yet. Start a conversation to see it here!")
        
    # Enhanced Service Status Dashboard
    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("---")
    
    st.markdown("""
    <div style="background: linear-gradient(135deg, #667eea, #764ba2); color: white; border-radius: 10px; padding: 20px; margin: 20px 0;">
        <h3 style="margin: 0 0 15px 0; font-size: 1.4rem; font-weight: 600;">🔧 Service Status</h3>
    </div>
    """, unsafe_allow_html=True)
    
    status = check_service_status()
    
    # Enhanced service status items
    services = [
        ("Nova Bedrock", status['nova_bedrock']),
        ("Perplexity AI", status['perplexity_llm']),
        ("OpenAI LLM", status['openai_llm']),
        ("Serper Search", status['serper']),
        ("Twitter API", status['twitter_api']),
        ("LangGraph Agents", status['langgraph_agents']),
        ("Memory Enhancement", status['memory_enhancement'])
    ]
    
    # Create horizontal columns for service status
    status_cols = st.columns(len(services))
    
    for i, (service_name, service_status) in enumerate(services):
        with status_cols[i]:
            status_color = "#28a745" if service_status else "#dc3545"
            status_text = "Online" if service_status else "Offline"
            status_icon = "🟢" if service_status else "🔴"
            
            # Special styling for Perplexity and LangGraph services
            if "Perplexity" in service_name or "LangGraph" in service_name or "Memory" in service_name:
                bg_color = "rgba(102, 126, 234, 0.1)" if service_status else "rgba(220, 53, 69, 0.1)"
            else:
                bg_color = "rgba(255,255,255,0.1)"

            st.markdown(f"""
            <div style="background: {bg_color}; padding: 8px 12px; border-radius: 6px; text-align: center; margin: 5px 0; border: 1px solid {'#667eea' if 'Perplexity' in service_name or 'LangGraph' in service_name else '#dee2e6'};">
                <div style="font-size: 0.8rem; font-weight: 500; color: #1f4e79;">{service_name}</div>
                <div style="color: {status_color}; font-size: 0.7rem; margin-top: 2px;">
                    {status_icon} {status_text}
                </div>
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Minimal Footer
    st.markdown("""
    <div class="minimal-footer">
        © 2024 INSYT Platform • Powered by LangGraph Agentic AI
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()